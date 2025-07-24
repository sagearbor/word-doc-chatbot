#!/usr/bin/env python3
"""
Create a sample fallback document for testing
Usage: python test_fallback_sample.py
"""

from docx import Document
import os

def create_sample_fallback_document():
    """Create a sample fallback document with legal requirements"""
    
    # Create new document
    doc = Document()
    
    # Add title
    title = doc.add_paragraph()
    title_run = title.add_run("SAMPLE CONTRACT REQUIREMENTS")
    title_run.bold = True
    title.alignment = 1  # Center alignment
    
    doc.add_paragraph("")  # Empty line
    
    # Add requirements sections
    doc.add_paragraph("1. GENERAL REQUIREMENTS")
    
    doc.add_paragraph("1.1 The Contractor must provide all services in accordance with professional standards.")
    doc.add_paragraph("1.2 All deliverables shall be submitted within the agreed timeline.")
    doc.add_paragraph("1.3 The Contractor is required to maintain confidentiality of all project information.")
    
    doc.add_paragraph("")
    
    doc.add_paragraph("2. QUALITY STANDARDS")
    
    doc.add_paragraph("2.1 All work must meet industry best practices and standards.")
    doc.add_paragraph("2.2 The Contractor shall provide regular progress reports.")
    doc.add_paragraph("2.3 Any defects must be corrected within 5 business days.")
    
    doc.add_paragraph("")
    
    doc.add_paragraph("3. COMPLIANCE REQUIREMENTS")
    
    doc.add_paragraph("3.1 The Contractor must comply with all applicable laws and regulations.")
    doc.add_paragraph("3.2 Safety protocols are required to be followed at all times.")
    doc.add_paragraph("3.3 Documentation shall be maintained for audit purposes.")
    
    doc.add_paragraph("")
    
    doc.add_paragraph("4. PROHIBITED ACTIVITIES")
    
    doc.add_paragraph("4.1 Subcontracting is prohibited without written approval.")
    doc.add_paragraph("4.2 The Contractor must not disclose confidential information.")
    doc.add_paragraph("4.3 Use of project resources for personal purposes is not permitted.")
    
    # Save document
    filename = "sample_fallback_contract.docx"
    doc.save(filename)
    
    print(f"✅ Created sample fallback document: {filename}")
    print(f"📊 Document contains:")
    print(f"   - 4 main sections")
    print(f"   - 12 numbered subsections")
    print(f"   - Requirements using 'must', 'shall', 'required', 'prohibited'")
    print(f"   - Hierarchical numbering (1.1, 1.2, etc.)")
    print("")
    print(f"🔍 Test it with: python debug_fallback.py {filename}")
    
    return filename

def main():
    """Main function"""
    sample_file = create_sample_fallback_document()
    
    # Also provide usage instructions
    print("")
    print("📝 How to test:")
    print(f"1. Run: python debug_fallback.py {sample_file}")
    print("2. Upload this file to your Streamlit app as a fallback document")
    print("3. Check if requirements are properly detected")
    print("")
    print("💡 If this sample works but your document doesn't:")
    print("   - Make sure your document uses requirement language (must, shall, required)")
    print("   - Use numbered sections (1.1, 1.2, etc.)")
    print("   - Include clear legal/contractual language")

if __name__ == "__main__":
    main()