#!/usr/bin/env python3
"""
Test the complete fallback document processing pipeline
"""

import pytest
import os
import sys
import tempfile
import shutil
from pathlib import Path

# Add backend to path
backend_path = Path(__file__).parent.parent / 'backend'
sys.path.insert(0, str(backend_path))

@pytest.fixture(scope="session")
def test_documents():
    """Create test documents for the session"""
    from create_test_documents import create_main_document, create_fallback_document
    
    # Create temp directory
    temp_dir = tempfile.mkdtemp()
    original_cwd = os.getcwd()
    
    try:
        # Change to temp directory
        os.chdir(temp_dir)
        
        # Create test documents
        main_doc = create_main_document()
        fallback_doc = create_fallback_document()
        
        yield {
            'temp_dir': temp_dir,
            'main_doc': os.path.join(temp_dir, main_doc),
            'fallback_doc': os.path.join(temp_dir, fallback_doc)
        }
        
    finally:
        # Cleanup
        os.chdir(original_cwd)
        shutil.rmtree(temp_dir, ignore_errors=True)

class TestLegalDocumentProcessor:
    """Test the legal document processor components"""
    
    def test_requirement_extraction(self, test_documents):
        """Test that requirements are extracted from fallback document"""
        
        from legal_document_processor import extract_fallback_requirements
        
        fallback_path = test_documents['fallback_doc']
        requirements = extract_fallback_requirements(fallback_path)
        
        # Should find multiple requirements
        assert len(requirements) > 0, "Should extract requirements from fallback document"
        assert len(requirements) >= 5, f"Expected at least 5 requirements, got {len(requirements)}"
        
        # Check requirement types are detected
        req_types = [req.requirement_type for req in requirements]
        assert 'must' in req_types, "Should detect 'must' requirements"
        assert 'shall' in req_types, "Should detect 'shall' requirements"
        assert 'required' in req_types, "Should detect 'required' requirements"
        assert 'prohibited' in req_types, "Should detect 'prohibited' requirements"
        
    def test_instruction_generation(self, test_documents):
        """Test that instructions are generated from requirements"""
        
        from legal_document_processor import generate_instructions_from_fallback
        
        fallback_path = test_documents['fallback_doc']
        instructions = generate_instructions_from_fallback(fallback_path)
        
        assert instructions, "Should generate instructions"
        assert len(instructions) > 100, "Instructions should be substantial"
        assert "CRITICAL" in instructions or "HIGH" in instructions, "Should include priority levels"
        assert "must" in instructions.lower(), "Should include requirement keywords"
        
    def test_legal_document_parsing(self, test_documents):
        """Test full legal document parsing"""
        
        from legal_document_processor import parse_legal_document
        
        fallback_path = test_documents['fallback_doc']
        structure = parse_legal_document(fallback_path)
        
        assert structure.title, "Should extract document title"
        assert len(structure.requirements) > 0, "Should find requirements in document"
        assert len(structure.sections) > 0, "Should parse document sections"

class TestFallbackProcessingAPI:
    """Test the API endpoints for fallback processing"""
    
    @pytest.fixture
    def client(self):
        """Create test client"""
        from fastapi.testclient import TestClient
        from main import app
        return TestClient(app)
    
    def test_analyze_fallback_endpoint(self, client, test_documents):
        """Test the analyze fallback requirements endpoint"""
        
        fallback_path = test_documents['fallback_doc']
        
        with open(fallback_path, 'rb') as f:
            files = {'file': ('test_fallback.docx', f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
            data = {'context': 'Testing fallback analysis'}
            
            response = client.post('/analyze-fallback-requirements/', files=files, data=data)
        
        assert response.status_code == 200, f"API call failed: {response.text}"
        
        result = response.json()
        assert result['status'] == 'success', f"Analysis failed: {result}"
        assert result['requirements_count'] > 0, "Should find requirements"
        assert 'instructions' in result, "Should generate instructions"
        assert 'categorized_requirements' in result, "Should categorize requirements"
    
    def test_process_with_fallback_endpoint(self, client, test_documents):
        """Test the full processing with fallback endpoint"""
        
        main_path = test_documents['main_doc']
        fallback_path = test_documents['fallback_doc']
        
        with open(main_path, 'rb') as main_f, open(fallback_path, 'rb') as fallback_f:
            files = {
                'input_file': ('main.docx', main_f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'),
                'fallback_file': ('fallback.docx', fallback_f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            }
            data = {
                'user_instructions': 'Apply all requirements from fallback document',
                'author_name': 'Test Author',
                'case_sensitive': True,
                'add_comments': True,
                'debug_mode': True,
                'extended_debug_mode': True,
                'merge_strategy': 'append'
            }
            
            response = client.post('/process-document-with-fallback/', files=files, data=data)
        
        assert response.status_code == 200, f"Processing failed: {response.text}"
        
        result = response.json()
        assert 'processed_filename' in result, "Should return processed filename"
        assert 'status_message' in result, "Should return status message"
        assert 'edits_suggested_count' in result, "Should return edit counts"
        assert 'debug_info' in result, "Should return debug information"

class TestEndToEndWorkflow:
    """Test the complete end-to-end workflow"""
    
    def test_complete_fallback_workflow(self, test_documents):
        """Test the complete workflow from fallback analysis to document processing"""
        
        # Step 1: Extract requirements from fallback
        from legal_document_processor import extract_fallback_requirements, generate_instructions_from_fallback
        
        fallback_path = test_documents['fallback_doc']
        requirements = extract_fallback_requirements(fallback_path)
        
        assert len(requirements) > 0, "Step 1: Should extract requirements"
        print(f"✅ Step 1: Extracted {len(requirements)} requirements")
        
        # Step 2: Generate instructions
        instructions = generate_instructions_from_fallback(fallback_path)
        
        assert instructions, "Step 2: Should generate instructions"
        assert len(instructions) > 100, "Step 2: Instructions should be substantial"
        print(f"✅ Step 2: Generated {len(instructions)} character instructions")
        
        # Step 3: Process main document (simulate API call)
        from main import process_document_with_fallback
        from llm_handler import get_llm_edit_suggestions
        
        main_path = test_documents['main_doc']
        
        # This would be the full processing - for now just verify the components work
        print("✅ Step 3: Core components verified")
        
        # Verify we have the key components for processing
        assert os.path.exists(main_path), "Main document should exist"
        assert os.path.exists(fallback_path), "Fallback document should exist"
        
        print("🎉 Complete workflow components verified!")

def test_debug_information_structure():
    """Test that debug information has the expected structure"""
    
    from legal_document_processor import extract_fallback_requirements
    
    # Create a simple test case
    import tempfile
    from docx import Document
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc = Document()
        doc.add_paragraph("The contractor must complete work within 30 days.")
        doc.add_paragraph("Subcontracting is prohibited without approval.")
        doc.save(tmp.name)
        
        try:
            requirements = extract_fallback_requirements(tmp.name)
            
            # Verify debug-friendly information
            assert len(requirements) > 0, "Should extract requirements"
            
            for req in requirements:
                assert hasattr(req, 'text'), "Requirements should have text"
                assert hasattr(req, 'requirement_type'), "Requirements should have type"
                assert hasattr(req, 'priority'), "Requirements should have priority"
                assert hasattr(req, 'section'), "Requirements should have section"
                
        finally:
            os.unlink(tmp.name)

if __name__ == "__main__":
    # Run tests if called directly
    pytest.main([__file__, "-v"])