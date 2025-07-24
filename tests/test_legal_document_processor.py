"""
Comprehensive test suite for Legal Document Processor (Phase 1.1)
Tests document parsing, requirement extraction, and legal structure analysis
"""

import pytest
from unittest.mock import Mock, patch, MagicMock
from pathlib import Path
import tempfile
from docx import Document

# Import the module under test
import sys
import os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))

from backend.legal_document_processor import (
    LegalDocumentStructure,
    LegalRequirement,
    LegalDocumentParser,
    parse_legal_document,
    extract_fallback_requirements,
    generate_instructions_from_fallback
)


class TestLegalDocumentStructure:
    """Test the LegalDocumentStructure dataclass"""
    
    def test_structure_creation(self):
        """Test creating a legal document structure"""
        structure = LegalDocumentStructure(
            title="Test Agreement",
            sections=[],
            requirements=[],
            whereas_clauses=["Whereas party A", "Whereas party B"],
            definitions={},
            authors=["John Doe"],
            metadata={}
        )
        
        assert structure.title == "Test Agreement"
        assert len(structure.whereas_clauses) == 2
        assert "John Doe" in structure.authors


class TestLegalRequirement:
    """Test the LegalRequirement dataclass"""
    
    def test_requirement_creation(self):
        """Test creating a legal requirement"""
        req = LegalRequirement(
            text="The Contractor shall deliver services",
            requirement_type="shall",
            priority=1,
            section="2.1",
            context="Within the scope of services",
            formatting={"bold": True}
        )
        
        assert req.requirement_type == "shall"
        assert req.priority == 1
        assert req.formatting["bold"] is True


class TestLegalDocumentParser:
    """Test the main LegalDocumentParser class"""
    
    @pytest.fixture
    def parser(self):
        """Create a parser instance"""
        return LegalDocumentParser()
    
    @pytest.fixture
    def mock_document(self):
        """Create a mock Word document"""
        doc = Mock(spec=Document)
        
        # Mock paragraphs
        p1 = Mock()
        p1.text = "MASTER SERVICE AGREEMENT"
        p1.style.name = "Title"
        p1.runs = []
        
        p2 = Mock()
        p2.text = "1.1 The Contractor shall provide services"
        p2.style.name = "Normal"
        p2.runs = []
        
        p3 = Mock()
        p3.text = "2.1 The Client must pay within 30 days"
        p3.style.name = "Normal"
        p3.runs = []
        
        p4 = Mock()
        p4.text = "The parties shall not disclose confidential information"
        p4.style.name = "Normal"
        p4.runs = []
        
        doc.paragraphs = [p1, p2, p3, p4]
        doc.core_properties = Mock(author="Test Author")
        
        return doc
    
    def test_extract_title(self, parser, mock_document):
        """Test title extraction"""
        title = parser._extract_title(mock_document)
        assert title == "MASTER SERVICE AGREEMENT"
    
    def test_extract_sections(self, parser, mock_document):
        """Test section extraction with hierarchical numbering"""
        sections = parser._extract_sections(mock_document.paragraphs)
        assert len(sections) == 2
        assert sections[0]["number"] == "1.1"
        assert sections[1]["number"] == "2.1"
    
    def test_extract_requirements(self, parser, mock_document):
        """Test requirement extraction"""
        requirements = parser._extract_requirements(mock_document.paragraphs)
        assert len(requirements) >= 3  # Should find "shall", "must", and "shall not"
        
        # Check requirement types
        req_types = [req.requirement_type for req in requirements]
        assert "shall" in req_types
        assert "must" in req_types
        assert "prohibited" in req_types
    
    def test_parse_document(self, parser, mock_document):
        """Test full document parsing"""
        with patch('backend.legal_document_processor.Document', return_value=mock_document):
            structure = parser.parse("test.docx")
            
            assert structure.title == "MASTER SERVICE AGREEMENT"
            assert len(structure.sections) == 2
            assert len(structure.requirements) >= 3
            assert "Test Author" in structure.authors
    
    def test_requirement_priority(self, parser):
        """Test requirement priority assignment"""
        text1 = "The Contractor shall immediately cease all work"
        text2 = "The parties should consider mediation"
        text3 = "The Client may request additional services"
        
        req1 = parser._create_requirement(text1, "shall", "1.1", "")
        req2 = parser._create_requirement(text2, "should", "1.2", "")
        req3 = parser._create_requirement(text3, "may", "1.3", "")
        
        assert req1.priority < req2.priority  # "shall" + "immediately" = higher priority
        assert req2.priority < req3.priority  # "should" < "may"
    
    def test_cross_reference_detection(self, parser):
        """Test detection of cross-references"""
        text = "As specified in Section 2.3, the payment terms shall apply"
        refs = parser._extract_cross_references(text)
        assert "Section 2.3" in refs
        
        text2 = "Subject to Clause 5.1 and Appendix A requirements"
        refs2 = parser._extract_cross_references(text2)
        assert "Clause 5.1" in refs2
        assert "Appendix A" in refs2


class TestModuleFunctions:
    """Test module-level convenience functions"""
    
    @pytest.fixture
    def temp_docx(self):
        """Create a temporary DOCX file"""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc = Document()
            doc.add_heading("Test Document", 0)
            doc.add_paragraph("1. The Contractor shall deliver on time")
            doc.add_paragraph("2. Payment must be made within NET 30")
            doc.save(f.name)
            yield f.name
        os.unlink(f.name)
    
    def test_parse_legal_document(self, temp_docx):
        """Test the parse_legal_document function"""
        structure = parse_legal_document(temp_docx)
        
        assert isinstance(structure, LegalDocumentStructure)
        assert structure.title == "Test Document"
        assert len(structure.requirements) >= 2
    
    def test_extract_fallback_requirements(self, temp_docx):
        """Test requirement extraction function"""
        requirements = extract_fallback_requirements(temp_docx)
        
        assert isinstance(requirements, list)
        assert all(isinstance(req, LegalRequirement) for req in requirements)
        assert len(requirements) >= 2
    
    def test_generate_instructions_from_fallback(self, temp_docx):
        """Test instruction generation"""
        instructions = generate_instructions_from_fallback(temp_docx)
        
        assert isinstance(instructions, str)
        assert "shall" in instructions or "must" in instructions
        assert "legal document" in instructions.lower()
    
    def test_generate_instructions_with_context(self, temp_docx):
        """Test instruction generation with context"""
        context = "Processing clinical trial agreement"
        instructions = generate_instructions_from_fallback(temp_docx, context)
        
        assert context in instructions
        assert len(instructions) > 100  # Should be substantial


class TestEdgeCases:
    """Test edge cases and error handling"""
    
    def test_empty_document(self):
        """Test parsing empty document"""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc = Document()
            doc.save(f.name)
            
            try:
                structure = parse_legal_document(f.name)
                assert structure.title == ""
                assert len(structure.requirements) == 0
            finally:
                os.unlink(f.name)
    
    def test_invalid_file_path(self):
        """Test handling of invalid file path"""
        with pytest.raises(Exception):
            parse_legal_document("/nonexistent/file.docx")
    
    def test_malformed_numbering(self):
        """Test handling of malformed section numbers"""
        parser = LegalDocumentParser()
        
        # Should handle various numbering formats
        sections = parser._extract_sections([
            Mock(text="1.a.i The first section", style=Mock(name="Normal")),
            Mock(text="II.B Payment terms", style=Mock(name="Normal")),
            Mock(text="Article 3: Termination", style=Mock(name="Normal"))
        ])
        
        assert len(sections) >= 1  # Should extract something
    
    def test_unicode_handling(self):
        """Test handling of unicode characters"""
        parser = LegalDocumentParser()
        text = "The Contractor shall pay €1000 within 30 days"
        
        req = parser._create_requirement(text, "shall", "1.1", "")
        assert "€" in req.text
        assert req.requirement_type == "shall"


class TestComplexDocuments:
    """Test parsing of complex legal documents"""
    
    def test_nested_requirements(self):
        """Test extraction of nested requirements"""
        parser = LegalDocumentParser()
        
        paragraphs = [
            Mock(text="The Contractor shall:", style=Mock(name="Normal")),
            Mock(text="a) Provide monthly reports", style=Mock(name="Normal")),
            Mock(text="b) Must maintain insurance", style=Mock(name="Normal")),
            Mock(text="c) Shall not subcontract without approval", style=Mock(name="Normal"))
        ]
        
        requirements = parser._extract_requirements(paragraphs)
        assert len(requirements) >= 3
        
        # Check that sub-items are captured
        req_texts = [req.text for req in requirements]
        assert any("monthly reports" in text for text in req_texts)
        assert any("insurance" in text for text in req_texts)
    
    def test_whereas_clause_extraction(self):
        """Test extraction of whereas clauses"""
        parser = LegalDocumentParser()
        
        doc = Mock()
        doc.paragraphs = [
            Mock(text="WHEREAS, Party A is engaged in software development", style=Mock(name="Normal")),
            Mock(text="WHEREAS, Party B requires custom software", style=Mock(name="Normal")),
            Mock(text="NOW, THEREFORE, the parties agree", style=Mock(name="Normal"))
        ]
        
        whereas_clauses = parser._extract_whereas_clauses(doc)
        assert len(whereas_clauses) == 2
        assert "Party A" in whereas_clauses[0]
        assert "Party B" in whereas_clauses[1]


@pytest.mark.integration
class TestIntegration:
    """Integration tests with real document creation"""
    
    def test_full_document_processing(self):
        """Test processing a complete legal document"""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            # Create a realistic legal document
            doc = Document()
            doc.add_heading("SOFTWARE DEVELOPMENT AGREEMENT", 0)
            
            doc.add_paragraph("WHEREAS, Developer has expertise in software development;")
            doc.add_paragraph("WHEREAS, Client desires to engage Developer;")
            
            doc.add_heading("1. SERVICES", 1)
            doc.add_paragraph("1.1 The Developer shall provide software development services as specified in Exhibit A.")
            doc.add_paragraph("1.2 The Developer must complete all deliverables by the deadlines specified.")
            
            doc.add_heading("2. PAYMENT", 1)
            doc.add_paragraph("2.1 Client shall pay Developer within thirty (30) days of invoice.")
            doc.add_paragraph("2.2 Late payments shall incur interest at 1.5% per month.")
            
            doc.add_heading("3. CONFIDENTIALITY", 1)
            doc.add_paragraph("3.1 The parties shall not disclose any confidential information.")
            
            doc.save(f.name)
            
            try:
                # Parse the document
                structure = parse_legal_document(f.name)
                
                # Verify structure
                assert "SOFTWARE DEVELOPMENT AGREEMENT" in structure.title
                assert len(structure.whereas_clauses) == 2
                assert len(structure.sections) >= 3
                assert len(structure.requirements) >= 5
                
                # Generate instructions
                instructions = generate_instructions_from_fallback(f.name)
                assert len(instructions) > 200
                assert "shall" in instructions
                assert "must" in instructions
                
            finally:
                os.unlink(f.name)


if __name__ == "__main__":
    pytest.main([__file__, "-v"])