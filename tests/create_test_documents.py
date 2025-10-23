#!/usr/bin/env python3
"""
Create test Word documents for debugging the fallback document processing
"""

from docx import Document
from docx.shared import Inches
import os

def create_main_document():
    """Create a simple main document that can be edited"""
    
    doc = Document()
    
    # Add title
    title = doc.add_heading('SERVICE AGREEMENT', 0)
    
    # Add main content with items that should be changed based on fallback requirements
    doc.add_heading('1. GENERAL PROVISIONS', level=1)
    
    p1 = doc.add_paragraph(
        "The Contractor will provide services to the Client. "
        "All work should be completed in a reasonable timeframe. "
        "The Contractor may use subcontractors at their discretion. "
        "Payment terms are flexible and can be negotiated."
    )
    
    doc.add_heading('2. DELIVERABLES', level=1)
    
    p2 = doc.add_paragraph(
        "The Contractor will deliver work products as agreed. "
        "Quality standards will be maintained. "
        "Documentation may be provided if requested. "
        "The Client can review work at any time."
    )
    
    doc.add_heading('3. CONFIDENTIALITY', level=1)
    
    p3 = doc.add_paragraph(
        "Both parties should maintain confidentiality of sensitive information. "
        "Information can be shared with authorized personnel. "
        "Confidential data will be protected as appropriate."
    )
    
    doc.add_heading('4. TERMINATION', level=1)
    
    p4 = doc.add_paragraph(
        "Either party may terminate this agreement with notice. "
        "Upon termination, work will be concluded promptly. "
        "Final payments will be made as agreed."
    )
    
    # Save the document
    filename = "test_main_document.docx"
    doc.save(filename)
    print(f"✅ Created main document: {filename}")
    return filename

def create_fallback_document():
    """Create a fallback document with clear requirements"""
    
    doc = Document()
    
    # Add title
    title = doc.add_heading('EDITING REQUIREMENTS FOR SERVICE AGREEMENTS', 0)
    
    # Add requirements that should be detected by the legal processor
    doc.add_heading('MANDATORY REQUIREMENTS', level=1)
    
    p1 = doc.add_paragraph(
        "1.1 The Contractor must complete all work within 30 business days of project start. "
        "This requirement ensures timely delivery of services."
    )
    
    p2 = doc.add_paragraph(
        "1.2 All deliverables shall be reviewed and approved by the Client before final acceptance. "
        "This ensures quality control and client satisfaction."
    )
    
    p3 = doc.add_paragraph(
        "1.3 The Contractor is required to provide weekly progress reports to the Client. "
        "Regular communication is essential for project success."
    )
    
    doc.add_heading('PROHIBITED ACTIVITIES', level=1)
    
    p4 = doc.add_paragraph(
        "2.1 Subcontracting is prohibited without prior written approval from the Client. "
        "This maintains control over project quality and security."
    )
    
    p5 = doc.add_paragraph(
        "2.2 The Contractor must not share confidential information with unauthorized parties. "
        "Protection of sensitive data is critical."
    )
    
    p6 = doc.add_paragraph(
        "2.3 Use of project resources for personal purposes is not permitted. "
        "All resources must be dedicated to the contracted work."
    )
    
    doc.add_heading('ADDITIONAL REQUIREMENTS', level=1)
    
    p7 = doc.add_paragraph(
        "3.1 All work must meet industry best practices and professional standards. "
        "Quality is non-negotiable."
    )
    
    p8 = doc.add_paragraph(
        "3.2 Payment shall be made within 15 days of invoice submission. "
        "Prompt payment ensures continued service delivery."
    )
    
    p9 = doc.add_paragraph(
        "3.3 The agreement must include a clause for dispute resolution through mediation. "
        "This provides a structured approach to resolving conflicts."
    )
    
    # Save the document
    filename = "test_fallback_requirements.docx"
    doc.save(filename)
    print(f"✅ Created fallback document: {filename}")
    return filename

def create_complex_fallback_document():
    """Create a more complex fallback document with various requirement types"""
    
    doc = Document()
    
    # Add title
    title = doc.add_heading('COMPREHENSIVE CONTRACT EDITING GUIDELINES', 0)
    
    # Add preamble
    doc.add_paragraph(
        "WHEREAS, the parties desire to establish clear requirements for contract modifications; and "
        "WHEREAS, standardization improves legal compliance and reduces disputes;"
    )
    
    doc.add_heading('1. CRITICAL REQUIREMENTS', level=1)
    
    doc.add_paragraph(
        "1.1 The service provider must obtain professional liability insurance of at least $1,000,000. "
        "This requirement is mandatory for all professional service contracts."
    )
    
    doc.add_paragraph(
        "1.2 All intellectual property created during the engagement shall be owned by the Client. "
        "This ensures proper ownership of work products."
    )
    
    doc.add_paragraph(
        "1.3 The contractor is required to maintain all records for a minimum of seven years. "
        "Record retention supports audit and compliance requirements."
    )
    
    doc.add_heading('2. PERFORMANCE STANDARDS', level=1)
    
    doc.add_paragraph(
        "2.1 All deliverables must be completed according to the project timeline. "
        "Timely completion is essential for project success."
    )
    
    doc.add_paragraph(
        "2.2 The service provider shall provide monthly status reports. "
        "Regular reporting ensures transparency and accountability."
    )
    
    doc.add_paragraph(
        "2.3 Quality assurance testing is required for all software deliverables. "
        "Testing ensures deliverables meet functional requirements."
    )
    
    doc.add_heading('3. RESTRICTIONS AND PROHIBITIONS', level=1)
    
    doc.add_paragraph(
        "3.1 Disclosure of confidential information is prohibited except as specifically authorized. "
        "Confidentiality protection is critical for business operations."
    )
    
    doc.add_paragraph(
        "3.2 The contractor must not engage in any activities that create a conflict of interest. "
        "Avoiding conflicts maintains professional integrity."
    )
    
    doc.add_paragraph(
        "3.3 Subcontracting of core services is not permitted without written consent. "
        "Direct performance ensures quality and accountability."
    )
    
    # Save the document
    filename = "test_complex_fallback.docx"
    doc.save(filename)
    print(f"✅ Created complex fallback document: {filename}")
    return filename

def main():
    """Create all test documents"""
    
    print("📄 Creating test Word documents...")
    print("=" * 50)
    
    try:
        # Create the documents
        main_doc = create_main_document()
        fallback_doc = create_fallback_document()
        complex_fallback = create_complex_fallback_document()
        
        print("\n" + "=" * 50)
        print("📊 Test Documents Created:")
        print(f"   📄 Main document: {main_doc}")
        print(f"   📋 Fallback requirements: {fallback_doc}")
        print(f"   📋 Complex fallback: {complex_fallback}")
        
        print("\n💡 Usage:")
        print("1. Start backend and frontend")
        print(f"2. Upload '{main_doc}' as the main document")
        print(f"3. Upload '{fallback_doc}' as the fallback document")
        print("4. Test the fallback processing with extended debugging")
        
        print(f"\n🧪 For more complex testing, use '{complex_fallback}' as the fallback document")
        
        # Verify files exist
        for filename in [main_doc, fallback_doc, complex_fallback]:
            if os.path.exists(filename):
                size_kb = os.path.getsize(filename) // 1024
                print(f"   ✅ {filename} ({size_kb} KB)")
            else:
                print(f"   ❌ {filename} - File not created!")
                
    except Exception as e:
        print(f"❌ Error creating test documents: {e}")
        import traceback
        print(traceback.format_exc())

if __name__ == "__main__":
    main()