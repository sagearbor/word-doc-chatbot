#!/usr/bin/env python3
import copy
import datetime
import re
import json
import os # Added for path manipulation
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from typing import List, Dict, Tuple, Optional # For type hinting

# --- Global Configuration Flags ---
# These will be set by the main function call
DEBUG_MODE = False
CASE_SENSITIVE_SEARCH = True
ADD_COMMENTS_TO_CHANGES = True
DEFAULT_AUTHOR_NAME = "LLM Editor" # Default author for changes

# --- Constants ---
ERROR_LOG_FILENAME_BASE = "change_log" # Base name for the log file

# --- XML Helper Functions (largely unchanged) ---
def create_del_element(author="Python Program", date_time=None, del_id="0"):
    if date_time is None:
        date_time = datetime.datetime.now(datetime.timezone.utc)
    del_el = OxmlElement('w:del')
    del_el.set(qn('w:author'), author)
    del_el.set(qn('w:id'), str(del_id))
    del_el.set(qn('w:date'), date_time.strftime("%Y-%m-%dT%H:%M:%SZ"))
    return del_el

def create_ins_element(author="Python Program", date_time=None, ins_id="0"):
    if date_time is None:
        date_time = datetime.datetime.now(datetime.timezone.utc)
    ins_el = OxmlElement('w:ins')
    ins_el.set(qn('w:author'), author)
    ins_el.set(qn('w:id'), str(ins_id))
    ins_el.set(qn('w:date'), date_time.strftime("%Y-%m-%dT%H:%M:%SZ"))
    return ins_el

def create_run_element_with_text(text_content, template_run_r=None, is_del_text=False):
    new_r = OxmlElement('w:r')
    if template_run_r is not None:
        rPr = template_run_r.find(qn('w:rPr'))
        if rPr is not None:
            new_r.append(copy.deepcopy(rPr))
    text_element_tag = 'w:delText' if is_del_text else 'w:t'
    text_el = OxmlElement(text_element_tag)
    text_el.set(qn('xml:space'), 'preserve')
    text_el.text = text_content
    new_r.append(text_el)
    return new_r

def log_debug(message):
    if DEBUG_MODE:
        print(f"DEBUG (word_processor): {message}")

def replace_text_in_paragraph_with_tracked_change(
        paragraph_idx, paragraph,
        contextual_old_text, specific_old_text, specific_new_text, reason_for_change,
        author, case_sensitive_flag, add_comments_flag, ambiguous_or_failed_changes_log):
    """
    Finds specific_old_text within a unique contextual_old_text in a paragraph
    and replaces it with specific_new_text, marking the change as tracked and adding a comment.
    Logs issues to ambiguous_or_failed_changes_log.
    Returns a status string: SUCCESS, CONTEXT_NOT_FOUND, CONTEXT_AMBIGUOUS, etc.
    """
    log_debug(f"P{paragraph_idx+1}: Attempting to change '{specific_old_text}' to '{specific_new_text}' within context '{contextual_old_text[:30]}...'")

    elements_contributing_to_visible_text = []
    current_visible_text_parts = []
    for idx, p_child_element in enumerate(paragraph._p):
        text_contribution = ""
        if p_child_element.tag == qn("w:r"):
            if p_child_element.find(qn('w:delText')) is None: # Ignore already deleted text
                for t_node in p_child_element.findall(qn('w:t')):
                    if t_node.text: text_contribution += t_node.text
                if p_child_element.find(qn('w:tab')) is not None: text_contribution += '\t'
        elif p_child_element.tag == qn("w:ins"): # Consider inserted text as visible
            for r_in_ins in p_child_element.findall(qn('w:r')):
                for t_in_ins in r_in_ins.findall(qn('w:t')):
                    if t_in_ins.text: text_contribution += t_in_ins.text
                if r_in_ins.find(qn('w:tab')) is not None: text_contribution += '\t'
        elif p_child_element.tag == qn("w:hyperlink"):
             for r_in_hyperlink in p_child_element.findall(qn('w:r')):
                if r_in_hyperlink.find(qn('w:delText')) is None:
                    for t_in_hyperlink in r_in_hyperlink.findall(qn('w:t')):
                        if t_in_hyperlink.text: text_contribution += t_in_hyperlink.text
                    if r_in_hyperlink.find(qn('w:tab')) is not None: text_contribution += '\t'

        if text_contribution:
            elements_contributing_to_visible_text.append({'el': p_child_element, 'text': text_contribution, 'p_child_idx': idx})
            current_visible_text_parts.append(text_contribution)
    
    visible_paragraph_text_original_case = "".join(current_visible_text_parts)
    search_text_in_doc = visible_paragraph_text_original_case if case_sensitive_flag else visible_paragraph_text_original_case.lower()
    search_context_from_llm = contextual_old_text if case_sensitive_flag else contextual_old_text.lower()

    log_debug(f"P{paragraph_idx+1}: Visible text (len {len(visible_paragraph_text_original_case)}): '{visible_paragraph_text_original_case[:100]}{'...' if len(visible_paragraph_text_original_case)>100 else ''}'")

    occurrences = []
    try:
        for match in re.finditer(re.escape(search_context_from_llm), search_text_in_doc):
            occurrences.append(match)
    except re.error as e:
        log_message = f"P{paragraph_idx+1}: Regex error searching for context '{search_context_from_llm[:50]}...': {e}. Skipping. LLM Reason: {reason_for_change}"
        log_debug(log_message)
        ambiguous_or_failed_changes_log.append({
            "paragraph_index": paragraph_idx + 1, "visible_text_snippet": visible_paragraph_text_original_case[:100],
            "contextual_old_text": contextual_old_text, "specific_old_text": specific_old_text,
            "llm_reason": reason_for_change, "issue": f"Regex error: {e}"})
        return "REGEX_ERROR_IN_CONTEXT_SEARCH"

    if not occurrences:
        log_debug(f"P{paragraph_idx+1}: Context '{contextual_old_text[:30]}...' not found.")
        # No need to add to log here, it's a common case if the edit isn't for this paragraph.
        return "CONTEXT_NOT_FOUND"

    if len(occurrences) > 1:
        log_message = (f"P{paragraph_idx+1}: Context '{contextual_old_text[:50]}...' is AMBIGUOUS. Found {len(occurrences)} times. "
                       f"Skipping change of '{specific_old_text}' to '{specific_new_text}'. LLM Reason: {reason_for_change}")
        log_debug(log_message)
        ambiguous_or_failed_changes_log.append({
            "paragraph_index": paragraph_idx + 1, "visible_text_snippet": visible_paragraph_text_original_case[:100],
            "contextual_old_text": contextual_old_text, "specific_old_text": specific_old_text,
            "specific_new_text": specific_new_text, "llm_reason": reason_for_change,
            "issue": f"Ambiguous context: Found {len(occurrences)} occurrences."})
        return "CONTEXT_AMBIGUOUS"

    match_context = occurrences[0]
    actual_context_found_in_doc = visible_paragraph_text_original_case[match_context.start() : match_context.end()]
    text_to_search_within_context = actual_context_found_in_doc if case_sensitive_flag else actual_context_found_in_doc.lower()
    specific_text_to_find_in_context = specific_old_text if case_sensitive_flag else specific_old_text.lower()

    try:
        specific_start_in_context = text_to_search_within_context.index(specific_text_to_find_in_context)
    except ValueError:
        log_message = (f"P{paragraph_idx+1}: Specific text '{specific_old_text}' NOT FOUND within the identified unique context '{actual_context_found_in_doc}'. "
                       f"Skipping. LLM Reason: {reason_for_change}")
        log_debug(log_message)
        ambiguous_or_failed_changes_log.append({
            "paragraph_index": paragraph_idx + 1, "visible_text_snippet": visible_paragraph_text_original_case[:100],
            "contextual_old_text": contextual_old_text, "specific_old_text": specific_old_text,
            "specific_new_text": specific_new_text, "llm_reason": reason_for_change,
            "issue": "Specific text not found within unique context." })
        return "SPECIFIC_TEXT_NOT_IN_CONTEXT"

    # Determine the actual casing of the specific old text from the document
    actual_specific_old_text_to_delete = actual_context_found_in_doc[specific_start_in_context : specific_start_in_context + len(specific_old_text)]
    
    global_specific_start = match_context.start() + specific_start_in_context # Start of specific_old_text in the whole paragraph text
    log_debug(f"P{paragraph_idx+1}: Unique context found. Specific text '{actual_specific_old_text_to_delete}' (len {len(actual_specific_old_text_to_delete)}) located for replacement with '{specific_new_text}'.")


    target_p_child_indices = [] # Indices of paragraph._p elements to be replaced
    first_involved_r_element_for_style = None # To copy style from
    current_pos_in_visible_text = 0
    specific_end_index_in_visible_text = global_specific_start + len(actual_specific_old_text_to_delete)

    # Map the string indices (global_specific_start, specific_end_index_in_visible_text)
    # to the paragraph's child elements that contain this text.
    for item in elements_contributing_to_visible_text:
        el_visible_text = item['text']
        el_start_pos = current_pos_in_visible_text
        el_end_pos = current_pos_in_visible_text + len(el_visible_text)

        # Check for overlap: max(start1, start2) < min(end1, end2)
        if max(el_start_pos, global_specific_start) < min(el_end_pos, specific_end_index_in_visible_text):
            target_p_child_indices.append(item['p_child_idx'])
            if first_involved_r_element_for_style is None:
                # Try to get a w:r element for style properties
                if item['el'].tag == qn('w:r'):
                    first_involved_r_element_for_style = item['el']
                elif item['el'].tag in [qn('w:ins'), qn('w:hyperlink')]: # If it's inside an <ins> or <hyperlink>
                    r_child = item['el'].find(qn('w:r'))
                    if r_child is not None:
                        first_involved_r_element_for_style = r_child
        
        current_pos_in_visible_text = el_end_pos
        if current_pos_in_visible_text >= specific_end_index_in_visible_text and target_p_child_indices:
            # We've passed the end of the text to be replaced and found some elements
            break
    
    if not target_p_child_indices:
        log_message = f"P{paragraph_idx+1}: Could not map the specific text '{actual_specific_old_text_to_delete}' to underlying XML elements. This might happen with complex formatting or overlapping changes. Skipping. LLM Reason: {reason_for_change}"
        log_debug(log_message)
        ambiguous_or_failed_changes_log.append({
            "paragraph_index": paragraph_idx + 1, 
            "issue": "XML mapping failed for specific text.", 
            "contextual_old_text": contextual_old_text, 
            "specific_old_text": actual_specific_old_text_to_delete, 
            "llm_reason": reason_for_change,
            "visible_text_snippet": visible_paragraph_text_original_case[:100]
        })
        return "XML_MAPPING_FAILED"

    # Simplification: If the change spans multiple elements OR is only a part of a single element,
    # it's complex. The original script had logic for partial match in a single run but it was
    # marked as "too complex". We'll treat any multi-element or partial single-element change as complex for now.
    # This means the *entire text* of the involved elements must match actual_specific_old_text_to_delete.

    text_from_target_elements = ""
    for p_idx in target_p_child_indices:
        item = next(i for i in elements_contributing_to_visible_text if i['p_child_idx'] == p_idx)
        text_from_target_elements += item['text']
    
    compare_text_from_target_elements = text_from_target_elements if case_sensitive_flag else text_from_target_elements.lower()
    compare_actual_specific_old_text_to_delete = actual_specific_old_text_to_delete if case_sensitive_flag else actual_specific_old_text_to_delete.lower()

    if compare_text_from_target_elements != compare_actual_specific_old_text_to_delete:
        log_message = (f"P{paragraph_idx+1}: The text '{actual_specific_old_text_to_delete}' (from string search) does not perfectly match the concatenated text "
                       f"'{text_from_target_elements}' from the identified XML element(s). This indicates a partial match or complex structure. Skipping. LLM Reason: {reason_for_change}")
        log_debug(log_message)
        ambiguous_or_failed_changes_log.append({
            "paragraph_index": paragraph_idx + 1,
            "issue": "Partial match or complex structure spanning elements.",
            "contextual_old_text": contextual_old_text,
            "specific_old_text": actual_specific_old_text_to_delete,
            "text_from_elements": text_from_target_elements,
            "llm_reason": reason_for_change
        })
        return "PARTIAL_MATCH_COMPLEX"


    log_debug(f"P{paragraph_idx+1}: XML element indices for '{actual_specific_old_text_to_delete}': {target_p_child_indices}. Matched text from elements: '{text_from_target_elements}'")

    # If we reach here, the text_from_target_elements exactly matches actual_specific_old_text_to_delete.
    # We can proceed to replace these elements.

    change_time = datetime.datetime.now(datetime.timezone.utc)
    del_obj = create_del_element(author=author, date_time=change_time)
    # Use actual_specific_old_text_to_delete to preserve original casing in the <w:delText>
    del_run_el = create_run_element_with_text(actual_specific_old_text_to_delete, first_involved_r_element_for_style, is_del_text=True)
    del_obj.append(del_run_el)

    ins_obj = create_ins_element(author=author, date_time=change_time + datetime.timedelta(seconds=1)) # Ensure ins is later
    ins_run_el = create_run_element_with_text(specific_new_text, first_involved_r_element_for_style, is_del_text=False)
    ins_obj.append(ins_run_el)
    
    # The elements to be removed are paragraph._p[p_idx] for p_idx in target_p_child_indices
    # We need to insert del_obj and ins_obj at the position of the first element being removed.
    insertion_point_in_p = min(target_p_child_indices) if target_p_child_indices else 0

    # Remove elements in reverse order of their indices to avoid shifting issues
    for p_idx_to_remove in sorted(target_p_child_indices, reverse=True):
        try:
            element_to_remove = paragraph._p[p_idx_to_remove]
            paragraph._p.remove(element_to_remove)
            log_debug(f"P{paragraph_idx+1}: Removed element at original p_child_idx {p_idx_to_remove}")
        except IndexError:
            log_message = f"P{paragraph_idx+1}: Error removing XML element at index {p_idx_to_remove}. This shouldn't happen if mapping was correct. Skipping change."
            log_debug(log_message)
            ambiguous_or_failed_changes_log.append({
                "paragraph_index": paragraph_idx + 1,
                "issue": f"XML element removal error at index {p_idx_to_remove}.",
                "contextual_old_text": contextual_old_text,
                "specific_old_text": actual_specific_old_text_to_delete,
                "llm_reason": reason_for_change
            })
            # This is a critical error for this change, might need to restore paragraph state or just log and skip.
            return "XML_REMOVAL_ERROR"
            
    paragraph._p.insert(insertion_point_in_p, del_obj)
    # If specific_new_text is empty, only deletion is performed.
    if specific_new_text:
        paragraph._p.insert(insertion_point_in_p + 1, ins_obj)
        log_debug(f"P{paragraph_idx+1}: Inserted <w:del> for '{actual_specific_old_text_to_delete}' and <w:ins> for '{specific_new_text}'.")
    else:
        log_debug(f"P{paragraph_idx+1}: Inserted <w:del> for '{actual_specific_old_text_to_delete}'. No insertion as new text is empty.")


    if add_comments_flag and reason_for_change and specific_new_text: # Only add comment if text was inserted
        try:
            comment_author_name = f"{author} (LLM)"
            # Create initials from author name (e.g., "Mark Editson (LLM)" -> "ME")
            name_parts = [word for word in comment_author_name.replace("(", "").replace(")", "").split() if word]
            comment_initials = (name_parts[0][0] + name_parts[1][0]).upper() if len(name_parts) >= 2 else (name_parts[0][:2].upper() if name_parts else "LS")
            
            # Add comment to the paragraph. The API adds it near the last run of the paragraph.
            # If specific_new_text is now part of the paragraph, it might be associated there.
            # For more precise comment anchoring to the inserted text, docx library might need direct XML manipulation for comments.
            # The current paragraph.add_comment() is a high-level API.
            comment = paragraph.add_comment(reason_for_change, author=comment_author_name, initials=comment_initials)
            log_debug(f"P{paragraph_idx+1}: Added comment: '{reason_for_change[:30]}...' for new text '{specific_new_text}'.")
        except Exception as e:
            log_debug(f"P{paragraph_idx+1}: WARNING - Could not add comment for new text '{specific_new_text}'. Error: {e}")
            ambiguous_or_failed_changes_log.append({
                "paragraph_index": paragraph_idx + 1, "issue": f"Comment addition failed: {e}",
                "contextual_old_text": contextual_old_text, "specific_old_text": actual_specific_old_text_to_delete,
                "specific_new_text": specific_new_text, "llm_reason": reason_for_change, "type": "Warning"})
    elif add_comments_flag and reason_for_change and not specific_new_text: # Deletion only
         try:
            comment_author_name = f"{author} (LLM)"
            name_parts = [word for word in comment_author_name.replace("(", "").replace(")", "").split() if word]
            comment_initials = (name_parts[0][0] + name_parts[1][0]).upper() if len(name_parts) >= 2 else (name_parts[0][:2].upper() if name_parts else "LS")
            comment = paragraph.add_comment(f"Deleted: '{actual_specific_old_text_to_delete}'. Reason: {reason_for_change}", author=comment_author_name, initials=comment_initials)
            log_debug(f"P{paragraph_idx+1}: Added comment for deletion: '{reason_for_change[:30]}...'.")
         except Exception as e:
            log_debug(f"P{paragraph_idx+1}: WARNING - Could not add comment for deletion of '{actual_specific_old_text_to_delete}'. Error: {e}")
            # Log this warning as well
            ambiguous_or_failed_changes_log.append({
                "paragraph_index": paragraph_idx + 1, "issue": f"Comment addition failed for deletion: {e}",
                "contextual_old_text": contextual_old_text, "specific_old_text": actual_specific_old_text_to_delete,
                "llm_reason": reason_for_change, "type": "Warning"})


    return "SUCCESS"


def process_document_with_edits(
    input_docx_path: str,
    output_docx_path: str,
    edits_to_make: List[Dict],
    author_name: str = DEFAULT_AUTHOR_NAME,
    debug_mode_flag: bool = False,
    case_sensitive_flag: bool = True,
    add_comments_flag: bool = True
) -> Tuple[bool, Optional[str], List[Dict]]:
    """
    Main processing function, adapted from the original main().
    Takes input/output paths and edits as a list of dicts.
    Returns a tuple: (success_status, error_log_file_path_or_None, ambiguous_or_failed_changes_log_list)
    """
    global DEBUG_MODE, CASE_SENSITIVE_SEARCH, ADD_COMMENTS_TO_CHANGES
    DEBUG_MODE = debug_mode_flag
    CASE_SENSITIVE_SEARCH = case_sensitive_flag
    ADD_COMMENTS_TO_CHANGES = add_comments_flag

    log_debug(f"Script starting. Input: {input_docx_path}, Output: {output_docx_path}")
    log_debug(f"Debug: {DEBUG_MODE}, CaseSensitive: {CASE_SENSITIVE_SEARCH}, AddComments: {ADD_COMMENTS_TO_CHANGES}")
    log_debug(f"Author for changes: {author_name}")
    log_debug(f"Number of edits to attempt: {len(edits_to_make)}")

    if not isinstance(edits_to_make, list) or not all(isinstance(item, dict) for item in edits_to_make):
        log_debug("FATAL: Edits must be a list of dictionaries.")
        # This kind of fatal error should ideally be caught before calling this function
        return False, None, [{"issue": "FATAL: Edits must be a list of dictionaries."}]

    ambiguous_or_failed_changes_log = []
    try:
        doc = Document(input_docx_path)
        log_debug(f"Successfully opened Word document: '{input_docx_path}'")
    except Exception as e:
        log_debug(f"FATAL: Error opening Word document '{input_docx_path}': {e}")
        return False, None, [{"issue": f"FATAL: Error opening Word document '{input_docx_path}': {e}"}]

    for para_idx, paragraph in enumerate(doc.paragraphs):
        log_debug(f"\n--- Processing Paragraph {para_idx+1} ---")
        # Create a copy of edits_to_make to iterate over, as successful edits might be removed
        # to prevent re-application attempts or to manage context for subsequent edits in the same paragraph.
        # However, the current logic applies each edit independently, so iterating the original list is fine.
        for edit_item_idx, edit_item in enumerate(edits_to_make):
            log_debug(f"Attempting edit item {edit_item_idx+1} in P{para_idx+1}")
            required_keys = ["contextual_old_text", "specific_old_text", "specific_new_text", "reason_for_change"]
            if not all(key in edit_item for key in required_keys):
                log_message = f"P{para_idx+1}, EditItem{edit_item_idx+1}: Invalid edit item structure (missing keys): {edit_item}. Skipping this item."
                log_debug(log_message)
                # Ensure this log item has enough context if it's the only thing reported
                ambiguous_or_failed_changes_log.append({
                    "paragraph_index": para_idx + 1, 
                    "edit_item_index": edit_item_idx +1,
                    "issue": "Invalid edit item structure.", 
                    "edit_item_snippet": str(edit_item)[:100]
                })
                continue

            # Ensure specific_new_text exists, even if empty (for deletions)
            specific_new_text = edit_item.get("specific_new_text", "")

            status = replace_text_in_paragraph_with_tracked_change(
                para_idx, paragraph,
                edit_item["contextual_old_text"],
                edit_item["specific_old_text"],
                specific_new_text, # Use the potentially defaulted empty string
                edit_item["reason_for_change"],
                author_name,
                CASE_SENSITIVE_SEARCH,
                ADD_COMMENTS_TO_CHANGES,
                ambiguous_or_failed_changes_log # This list is appended to by the function
            )
            if status == "SUCCESS":
                success_msg = f"SUCCESS: P{para_idx+1}: Applied change for context '{edit_item['contextual_old_text'][:30]}...'. Reason: {edit_item['reason_for_change']}"
                print(success_msg) # Also print to console for direct script use
                log_debug(success_msg)
                # Optional: If an edit is successful, you might want to avoid trying other edits
                # in the same paragraph if the text has significantly changed.
                # For now, we continue processing other edits in the same paragraph.
            elif status not in ["CONTEXT_NOT_FOUND", "REGEX_ERROR_IN_CONTEXT_SEARCH"]: # Log other non-success statuses that are not simple misses
                info_msg = f"INFO: P{para_idx+1}: Edit for context '{edit_item['contextual_old_text'][:30]}...' resulted in status: {status}."
                print(info_msg)
                log_debug(info_msg)
                # The `ambiguous_or_failed_changes_log` is already updated by the called function for these cases.


    try:
        doc.save(output_docx_path)
        print(f"\nProcessed document saved to '{output_docx_path}'")
        log_debug(f"Processed document saved to '{output_docx_path}'")
    except Exception as e:
        log_debug(f"FATAL: Error saving document to '{output_docx_path}': {e}")
        ambiguous_or_failed_changes_log.append({"issue": f"FATAL: Error saving document to '{output_docx_path}': {e}"})
        return False, None, ambiguous_or_failed_changes_log

    error_log_file_path = None
    if ambiguous_or_failed_changes_log:
        # Create a unique log filename if multiple calls are made
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        log_filename_with_ts = f"{ERROR_LOG_FILENAME_BASE}_{timestamp}.txt"
        
        # Ensure the output directory exists (e.g., if output_docx_path includes subdirs)
        output_dir = os.path.dirname(output_docx_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True) # Create if not exists
        
        error_log_file_path = os.path.join(output_dir, log_filename_with_ts) if output_dir else log_filename_with_ts


        try:
            with open(error_log_file_path, "w", encoding="utf-8") as f:
                f.write(f"--- LOG OF CHANGES NOT MADE, AMBIGUITIES, OR WARNINGS ({datetime.datetime.now()}) ---\n")
                f.write(f"Input DOCX: {os.path.basename(input_docx_path)}\n")
                f.write(f"Output DOCX: {os.path.basename(output_docx_path)}\n")
                f.write(f"Total Edits Attempted: {len(edits_to_make)}\n")
                f.write(f"Log Items: {len(ambiguous_or_failed_changes_log)}\n\n")

                for log_entry in ambiguous_or_failed_changes_log:
                    f.write("-----------------------------------------\n")
                    f.write(f"Paragraph Index (0-based): {log_entry.get('paragraph_index', 'N/A')}\n")
                    f.write(f"Visible Text Snippet: {log_entry.get('visible_text_snippet', 'N/A')}\n")
                    f.write(f"Context Searched: '{log_entry.get('contextual_old_text', '')}'\n")
                    f.write(f"Specific Old Text: '{log_entry.get('specific_old_text', '')}'\n")
                    f.write(f"Specific New Text: '{log_entry.get('specific_new_text', 'N/A')}'\n")
                    f.write(f"LLM Reason for Change: {log_entry.get('llm_reason', 'N/A')}'\n")
                    f.write(f"Issue: {log_entry.get('issue', 'Unknown')}\n")
                    if 'text_from_elements' in log_entry:
                         f.write(f"Text from XML elements: '{log_entry.get('text_from_elements')}'\n")
                    if log_entry.get('type') == "Warning": f.write(f"Type: Warning\n")
                f.write("-----------------------------------------\n")
            print(f"Log file with details on {len(ambiguous_or_failed_changes_log)} items saved to: '{error_log_file_path}'")
            log_debug(f"Log file saved to: '{error_log_file_path}'")
        except Exception as e:
            log_debug(f"ERROR: Could not write to log file '{error_log_file_path}': {e}")
            # If logging fails, the path might be invalid or permissions issue.
            # The log content is still in ambiguous_or_failed_changes_log list.
            error_log_file_path = None # Indicate logging to file failed.
    else:
        print(f"No ambiguous changes, critical errors, or warnings logged that required file logging.")
        log_debug("No items for error log file.")

    return True, error_log_file_path, ambiguous_or_failed_changes_log

if __name__ == '__main__':
    import argparse
    # This part is for standalone script execution, not directly used by the FastAPI app
    # but useful for testing the word_processor independently.

    DEFAULT_EDITS_FILE_PATH = "edits_to_apply.json"
    DEFAULT_INPUT_DOCX_PATH = "simpleDoc_1editor.docx" # Provide a sample for testing
    DEFAULT_OUTPUT_DOCX_PATH = "simpleDoc_1editor_changed.docx"

    parser = argparse.ArgumentParser(description="Apply tracked changes to a Word document based on an edits file or direct JSON string.")
    parser.add_argument("--input", default=DEFAULT_INPUT_DOCX_PATH, help=f"Path to the input DOCX file (default: {DEFAULT_INPUT_DOCX_PATH})")
    parser.add_argument("--output", default=DEFAULT_OUTPUT_DOCX_PATH, help=f"Path for the output DOCX file (default: {DEFAULT_OUTPUT_DOCX_PATH})")
    parser.add_argument("--editsfile", default=DEFAULT_EDITS_FILE_PATH, help=f"Path to the JSON file with edits (default: {DEFAULT_EDITS_FILE_PATH})")
    parser.add_argument("--editsjson", type=str, help="JSON string containing the edits directly.")
    parser.add_argument("--author", default=DEFAULT_AUTHOR_NAME, help=f"Author name for tracked changes (default: {DEFAULT_AUTHOR_NAME}).")
    parser.add_argument("--debug", action="store_true", help="Enable debug mode for verbose console output.")
    parser.add_argument("--caseinsensitive", action="store_false", dest="case_sensitive", help="Perform case-insensitive search (default is case-sensitive).")
    parser.add_argument("--nocase", action="store_false", dest="case_sensitive", help="Alias for --caseinsensitive.")
    parser.add_argument("--nocomments", action="store_false", dest="add_comments", help="Disable adding comments to tracked changes (default is to add comments).")
    
    parser.set_defaults(case_sensitive=True, add_comments=True)
    args = parser.parse_args()

    edits_data = []
    if args.editsjson:
        try:
            edits_data = json.loads(args.editsjson)
            print(f"Loaded edits from --editsjson argument.")
        except json.JSONDecodeError as e:
            print(f"FATAL: Error decoding JSON from --editsjson: {e}. Exiting.")
            exit(1)
    elif args.editsfile:
        try:
            with open(args.editsfile, 'r', encoding='utf-8') as f:
                edits_data = json.load(f)
            print(f"Successfully loaded {len(edits_data)} edits from '{args.editsfile}'.")
        except FileNotFoundError:
            print(f"FATAL: Edits file '{args.editsfile}' not found. Please create it or check the path. Exiting.")
            exit(1)
        except json.JSONDecodeError as e:
            print(f"FATAL: Error decoding JSON from '{args.editsfile}': {e}. Exiting.")
            exit(1)
        except Exception as e:
            print(f"FATAL: An unexpected error occurred while loading '{args.editsfile}': {e}. Exiting.")
            exit(1)
    else:
        print("FATAL: No edits provided. Use --editsfile or --editsjson. Exiting.")
        exit(1)

    if not os.path.exists(args.input):
        print(f"FATAL: Input file '{args.input}' not found. Exiting.")
        # Create a dummy input file for a quick test if it doesn't exist and no edits are specified.
        if args.input == DEFAULT_INPUT_DOCX_PATH and not edits_data:
            try:
                doc = Document()
                doc.add_paragraph("This is a sample paragraph for testing. The old text needs to be changed.")
                doc.add_paragraph("Another paragraph with some example text. And another example text.")
                doc.save(DEFAULT_INPUT_DOCX_PATH)
                print(f"Created a dummy input file: {DEFAULT_INPUT_DOCX_PATH}")
                # Create a dummy edits file too for this basic test
                if args.editsfile == DEFAULT_EDITS_FILE_PATH:
                    dummy_edits = [
                        {
                            "contextual_old_text": "The old text needs to be changed",
                            "specific_old_text": "old text",
                            "specific_new_text": "new shiny text",
                            "reason_for_change": "Update terminology as per LLM suggestion."
                        },
                        {
                            "contextual_old_text": "example text. And another example text.",
                            "specific_old_text": "example text", # This is ambiguous in the second paragraph without better context
                            "specific_new_text": "illustrative piece",
                            "reason_for_change": "Clarification by LLM."
                        }
                    ]
                    with open(DEFAULT_EDITS_FILE_PATH, 'w', encoding='utf-8') as f_edits:
                        json.dump(dummy_edits, f_edits, indent=2)
                    print(f"Created a dummy edits file: {DEFAULT_EDITS_FILE_PATH}")
                    edits_data = dummy_edits # Load them for the run
            except Exception as e:
                print(f"Could not create dummy input file: {e}")
                exit(1)
        else:
            exit(1)


    process_document_with_edits(
        args.input,
        args.output,
        edits_data,
        args.author,
        args.debug,
        args.case_sensitive,
        args.add_comments
    )
