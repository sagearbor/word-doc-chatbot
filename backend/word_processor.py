#!/usr/bin/env python3
import copy
import datetime
import re
import json
import os
import traceback
import zipfile # Added for raw XML extraction
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from typing import List, Dict, Tuple, Optional, Any
from difflib import SequenceMatcher
from dataclasses import dataclass

# --- Global Configuration Flags ---
DEBUG_MODE = False
EXTENDED_DEBUG_MODE = False
CASE_SENSITIVE_SEARCH = True
ADD_COMMENTS_TO_CHANGES = False
DEFAULT_AUTHOR_NAME = "LLM Editor"
FUZZY_MATCHING_ENABLED = True
FUZZY_MATCHING_THRESHOLD = 0.85  # 85% similarity required for fuzzy match

# --- Constants ---
ERROR_LOG_FILENAME_BASE = "change_log"
HIGHLIGHT_COLOR_AMBIGUOUS_SKIPPED = "orange"
ALLOWED_POST_BOUNDARY_PUNCTUATION = {',', ';', '.', ':', '!', '?', ')', ']', '}', '"', "'"}

# --- Data Structures ---
@dataclass
class TrackedChange:
    """Represents a tracked change extracted from a document"""
    change_type: str  # "insertion", "deletion", or "substitution"
    old_text: str  # Text that was deleted (empty for pure insertion)
    new_text: str  # Text that was inserted (empty for pure deletion)
    author: str
    date: str
    paragraph_index: int
    context_before: str = ""  # Text before the change for context
    context_after: str = ""   # Text after the change for context

    def to_edit_dict(self) -> Dict[str, str]:
        """Convert tracked change to edit format expected by process_document_with_edits"""
        # For substitutions (delete + insert), we have both old and new text
        # For pure insertions, old_text is empty
        # For pure deletions, new_text is empty

        # Get context for better matching
        contextual_old = f"{self.context_before}{self.old_text}{self.context_after}".strip()
        contextual_new = f"{self.context_before}{self.new_text}{self.context_after}".strip()

        return {
            "contextual_old_text": contextual_old if contextual_old else self.old_text,
            "specific_old_text": self.old_text,
            "specific_new_text": self.new_text,
            "reason_for_change": f"Based on tracked change by {self.author} in fallback document{f' (dated {self.date})' if self.date else ''}"
        }

# --- Fuzzy Matching Functions ---

def fuzzy_search_best_match(target_text: str, search_text: str, threshold: float = FUZZY_MATCHING_THRESHOLD) -> Optional[Dict]:
    """
    Find the best fuzzy match for target_text within search_text
    Returns dict with start, end, similarity, matched_text or None if no good match
    """
    if not FUZZY_MATCHING_ENABLED or len(target_text) < 3:
        return None
    
    best_match = None
    target_len = len(target_text)
    
    # Try different window sizes around the target length
    for window_size in [target_len, target_len + 5, target_len - 5, target_len + 10, target_len - 10]:
        if window_size < 3:
            continue
            
        for i in range(len(search_text) - window_size + 1):
            candidate = search_text[i:i + window_size]
            similarity = SequenceMatcher(None, target_text.lower(), candidate.lower()).ratio()
            
            if similarity >= threshold:
                if best_match is None or similarity > best_match['similarity']:
                    best_match = {
                        'start': i,
                        'end': i + window_size,
                        'similarity': similarity,
                        'matched_text': candidate
                    }
    
    return best_match

def fuzzy_find_text_in_context(specific_text: str, context_text: str, context_start_in_doc: int) -> Optional[Dict]:
    """
    Find specific_text within context_text using fuzzy matching
    Returns dict with global start/end positions or None
    """
    if not FUZZY_MATCHING_ENABLED:
        return None
    
    fuzzy_match = fuzzy_search_best_match(specific_text, context_text)
    if fuzzy_match:
        return {
            'global_start': context_start_in_doc + fuzzy_match['start'],
            'global_end': context_start_in_doc + fuzzy_match['end'],
            'matched_text': fuzzy_match['matched_text'],
            'similarity': fuzzy_match['similarity']
        }
    return None

def is_boundary_valid_fuzzy(text: str, start_pos: int, end_pos: int, full_text: str, similarity: float) -> bool:
    """
    More flexible boundary checking for fuzzy matches
    Allows for slight punctuation mismatches if similarity is high
    """
    if similarity >= 0.95:  # Very high confidence - be more lenient
        return True
        
    # Standard boundary check
    char_before = full_text[start_pos - 1] if start_pos > 0 else None
    char_after = full_text[end_pos] if end_pos < len(full_text) else None
    
    is_start_ok = (start_pos == 0 or (char_before and char_before.isspace()))
    is_end_ok = (end_pos == len(full_text) or (char_after and (char_after.isspace() or char_after in ALLOWED_POST_BOUNDARY_PUNCTUATION)))
    
    if is_start_ok and is_end_ok:
        return True
    
    # For medium confidence fuzzy matches, be slightly more lenient with punctuation
    if similarity >= 0.90:
        # Allow common punctuation mismatches
        if char_before and char_before in ['"', "'", '(', '[', '{']:
            is_start_ok = True
        if char_after and char_after in ['"', "'", ')', ']', '}', ',', '.', ';', ':', ' ']:
            is_end_ok = True
    
    return is_start_ok and is_end_ok

# --- XML Helper Functions ---
# ... (keep existing create_del_element, create_ins_element, create_run_element_with_text) ...
def create_del_element(author="Python Program", date_time=None):
    if date_time is None: date_time = datetime.datetime.now(datetime.timezone.utc)
    del_el = OxmlElement('w:del')
    del_el.set(qn('w:author'), author)
    del_el.set(qn('w:date'), date_time.strftime("%Y-%m-%dT%H:%M:%SZ"))
    return del_el

def create_ins_element(author="Python Program", date_time=None):
    if date_time is None: date_time = datetime.datetime.now(datetime.timezone.utc)
    ins_el = OxmlElement('w:ins')
    ins_el.set(qn('w:author'), author)
    ins_el.set(qn('w:date'), date_time.strftime("%Y-%m-%dT%H:%M:%SZ"))
    return ins_el

def create_run_element_with_text(text_content, template_run_r=None, is_del_text=False, highlight_color: Optional[str] = None):
    new_r = OxmlElement('w:r')
    rPr_element = None
    if template_run_r is not None:
        rPr_template = template_run_r.find(qn('w:rPr'))
        if rPr_template is not None: rPr_element = copy.deepcopy(rPr_template)

    if rPr_element is None and highlight_color: rPr_element = OxmlElement('w:rPr')

    if highlight_color and rPr_element is not None:
        for highlight_node in rPr_element.findall(qn('w:highlight')): rPr_element.remove(highlight_node)
        highlight_el = OxmlElement('w:highlight'); highlight_el.set(qn('w:val'), highlight_color)
        rPr_element.append(highlight_el)

    if rPr_element is not None and len(rPr_element) > 0: new_r.append(rPr_element)

    text_el_tag = 'w:delText' if is_del_text else 'w:t'
    text_el = OxmlElement(text_el_tag)
    text_el.set(qn('xml:space'), 'preserve'); text_el.text = text_content
    new_r.append(text_el)
    return new_r

def log_debug(message):
    if DEBUG_MODE: print(f"DEBUG (word_processor): {message}")

def _get_element_style_template_run(element_info_item_el, fallback_style_run):
    if element_info_item_el is not None and element_info_item_el.tag == qn('w:r'): return element_info_item_el
    if element_info_item_el is not None:
        r_child = element_info_item_el.find(qn('w:r'))
        if r_child is not None: return r_child
    return fallback_style_run if fallback_style_run is not None else OxmlElement('w:r')


def _build_visible_text_map(paragraph) -> Tuple[str, List[Dict[str, Any]]]:
    # ... (keep existing _build_visible_text_map) ...
    elements_map = []
    current_text_offset = 0
    for p_child_idx_loop, p_child_element in enumerate(list(paragraph._p)):
        text_contribution = ""
        element_type = "other"
        if p_child_element.tag == qn("w:r"):
            element_type = "r"
            if p_child_element.find(qn('w:delText')) is None:
                for t_node in p_child_element.findall(qn('w:t')):
                    if t_node.text is not None: text_contribution += t_node.text
                if p_child_element.find(qn('w:tab')) is not None: text_contribution += '\t'
        elif p_child_element.tag == qn("w:ins"):
            element_type = "ins"
            for r_in_ins in p_child_element.findall(qn('w:r')):
                if r_in_ins.find(qn('w:delText')) is None:
                    for t_in_ins in r_in_ins.findall(qn('w:t')):
                        if t_in_ins.text is not None: text_contribution += t_in_ins.text
                    if r_in_ins.find(qn('w:tab')) is not None: text_contribution += '\t'
        elif p_child_element.tag == qn("w:del"):
            element_type = "del"
            pass
        elif p_child_element.tag == qn("w:hyperlink"):
             element_type = "hyperlink"
             for r_in_hyperlink in p_child_element.findall(qn('w:r')):
                if r_in_hyperlink.find(qn('w:delText')) is None:
                    for t_in_hyperlink in r_in_hyperlink.findall(qn('w:t')):
                        if t_in_hyperlink.text is not None: text_contribution += t_in_hyperlink.text
                    if r_in_hyperlink.find(qn('w:tab')) is not None: text_contribution += '\t'
        if text_contribution:
            elements_map.append({
                'el': p_child_element, 'text': text_contribution,
                'p_child_idx': p_child_idx_loop, 'doc_start_offset': current_text_offset,
                'type': element_type,
                'original_author': p_child_element.get(qn('w:author')) if element_type in ["ins", "del"] else None,
                'original_date': p_child_element.get(qn('w:date')) if element_type in ["ins", "del"] else None
            })
            current_text_offset += len(text_contribution)
    final_text = "".join(item['text'] for item in elements_map)
    return final_text, elements_map


# --- NEW/MODIFIED FUNCTIONS FOR ANALYSIS ---

def extract_tracked_changes_as_text(input_docx_path: str) -> str:
    """
    Extracts a textual summary of tracked insertions and deletions from a DOCX file.
    (For Approach 1: Concise summary to LLM)
    """
    try:
        doc = Document(input_docx_path)
    except Exception as e:
        print(f"Error opening Word document '{input_docx_path}' for change extraction: {e}")
        return f"Error_Internal: Could not open document {os.path.basename(input_docx_path)} to extract changes. Details: {str(e)}"

    changes_summary_lines = ["Summary of Tracked Changes Found in the Document:\n"]
    changes_found_in_doc = False

    for para_idx, paragraph in enumerate(doc.paragraphs):
        para_has_changes = False
        current_para_changes_lines = []

        for p_child_element in paragraph._p: # Iterate over direct children of <w:p>
            if p_child_element.tag == qn("w:ins"):
                inserted_text_parts = []
                author = p_child_element.get(qn('w:author'), "Unknown Author")
                date_str = p_child_element.get(qn('w:date'), "")
                if date_str and 'T' in date_str: date_str = date_str.split('T')[0] # Simplify date

                for r_in_ins in p_child_element.findall(qn('w:r')):
                    for t_in_ins in r_in_ins.findall(qn('w:t')):
                        if t_in_ins.text: inserted_text_parts.append(t_in_ins.text)
                if inserted_text_parts:
                    current_para_changes_lines.append(f"  - Inserted: \"{''.join(inserted_text_parts)}\" (By: {author}{f', around {date_str}' if date_str else ''})")
                    para_has_changes = True
                    changes_found_in_doc = True

            elif p_child_element.tag == qn("w:del"):
                deleted_text_parts = []
                author = p_child_element.get(qn('w:author'), "Unknown Author")
                date_str = p_child_element.get(qn('w:date'), "")
                if date_str and 'T' in date_str: date_str = date_str.split('T')[0] # Simplify date
                    
                for r_in_del in p_child_element.findall(qn('w:r')):
                    for del_text_node in r_in_del.findall(qn('w:delText')):
                        if del_text_node.text: deleted_text_parts.append(del_text_node.text)
                if deleted_text_parts:
                    current_para_changes_lines.append(f"  - Deleted: \"{''.join(deleted_text_parts)}\" (By: {author}{f', around {date_str}' if date_str else ''})")
                    para_has_changes = True
                    changes_found_in_doc = True
        
        if para_has_changes:
            changes_summary_lines.append(f"In Paragraph {para_idx + 1}:")
            changes_summary_lines.extend(current_para_changes_lines)
            changes_summary_lines.append("")

    if not changes_found_in_doc:
        return "No tracked insertions or deletions were found in this document."
    return "\n".join(changes_summary_lines)

def extract_tracked_changes_structured(input_docx_path: str, context_chars: int = 50) -> List[TrackedChange]:
    """
    Extracts tracked changes from a DOCX file as structured data.

    This function processes a Word document with tracked changes and returns a list of
    TrackedChange objects that can be used to apply similar changes to another document.

    Args:
        input_docx_path: Path to the input DOCX file with tracked changes
        context_chars: Number of characters before/after to include as context

    Returns:
        List of TrackedChange objects representing all tracked changes in the document
    """
    try:
        doc = Document(input_docx_path)
    except Exception as e:
        print(f"Error opening Word document '{input_docx_path}' for structured change extraction: {e}")
        return []

    tracked_changes = []

    for para_idx, paragraph in enumerate(doc.paragraphs):
        # Build a map of the paragraph content to get context
        para_text = paragraph.text

        # Track position in paragraph for context extraction
        current_pos = 0
        pending_deletion = None  # Store deletion to pair with following insertion

        for p_child_element in paragraph._p:
            if p_child_element.tag == qn("w:ins"):
                # Extraction (insertion)
                inserted_text_parts = []
                author = p_child_element.get(qn('w:author'), "Unknown Author")
                date_str = p_child_element.get(qn('w:date'), "")
                if date_str and 'T' in date_str:
                    date_str = date_str.split('T')[0]

                for r_in_ins in p_child_element.findall(qn('w:r')):
                    for t_in_ins in r_in_ins.findall(qn('w:t')):
                        if t_in_ins.text:
                            inserted_text_parts.append(t_in_ins.text)

                if inserted_text_parts:
                    new_text = ''.join(inserted_text_parts)

                    # Get context
                    context_before = para_text[max(0, current_pos - context_chars):current_pos]
                    context_after = para_text[current_pos:current_pos + context_chars]

                    # Check if there's a pending deletion (substitution case)
                    if pending_deletion:
                        # This is a substitution (delete + insert)
                        change = TrackedChange(
                            change_type="substitution",
                            old_text=pending_deletion['text'],
                            new_text=new_text,
                            author=author,
                            date=date_str,
                            paragraph_index=para_idx,
                            context_before=context_before,
                            context_after=context_after
                        )
                        pending_deletion = None  # Clear the pending deletion
                    else:
                        # Pure insertion
                        change = TrackedChange(
                            change_type="insertion",
                            old_text="",
                            new_text=new_text,
                            author=author,
                            date=date_str,
                            paragraph_index=para_idx,
                            context_before=context_before,
                            context_after=context_after
                        )

                    tracked_changes.append(change)
                    current_pos += len(new_text)

            elif p_child_element.tag == qn("w:del"):
                # Deletion
                deleted_text_parts = []
                author = p_child_element.get(qn('w:author'), "Unknown Author")
                date_str = p_child_element.get(qn('w:date'), "")
                if date_str and 'T' in date_str:
                    date_str = date_str.split('T')[0]

                for r_in_del in p_child_element.findall(qn('w:r')):
                    for del_text_node in r_in_del.findall(qn('w:delText')):
                        if del_text_node.text:
                            deleted_text_parts.append(del_text_node.text)

                if deleted_text_parts:
                    deleted_text = ''.join(deleted_text_parts)

                    # Store this deletion in case the next element is an insertion (substitution)
                    pending_deletion = {
                        'text': deleted_text,
                        'author': author,
                        'date': date_str,
                        'pos': current_pos
                    }

            else:
                # Regular text - add any pending deletion as a pure deletion
                if pending_deletion:
                    context_before = para_text[max(0, pending_deletion['pos'] - context_chars):pending_deletion['pos']]
                    context_after = para_text[pending_deletion['pos']:pending_deletion['pos'] + context_chars]

                    change = TrackedChange(
                        change_type="deletion",
                        old_text=pending_deletion['text'],
                        new_text="",
                        author=pending_deletion['author'],
                        date=pending_deletion['date'],
                        paragraph_index=para_idx,
                        context_before=context_before,
                        context_after=context_after
                    )
                    tracked_changes.append(change)
                    pending_deletion = None

                # Update position for regular text
                if hasattr(p_child_element, 'text') and p_child_element.text:
                    current_pos += len(p_child_element.text)

        # Handle any remaining pending deletion at end of paragraph
        if pending_deletion:
            context_before = para_text[max(0, pending_deletion['pos'] - context_chars):pending_deletion['pos']]
            context_after = para_text[pending_deletion['pos']:pending_deletion['pos'] + context_chars]

            change = TrackedChange(
                change_type="deletion",
                old_text=pending_deletion['text'],
                new_text="",
                author=pending_deletion['author'],
                date=pending_deletion['date'],
                paragraph_index=para_idx,
                context_before=context_before,
                context_after=context_after
            )
            tracked_changes.append(change)

    return tracked_changes

def convert_tracked_changes_to_edits(tracked_changes: List[TrackedChange]) -> List[Dict[str, str]]:
    """
    Convert a list of TrackedChange objects to edit dictionaries for processing.

    This function takes tracked changes extracted from a fallback document and converts
    them to the format expected by process_document_with_edits().

    Args:
        tracked_changes: List of TrackedChange objects

    Returns:
        List of edit dictionaries with contextual_old_text, specific_old_text,
        specific_new_text, and reason_for_change fields
    """
    edits = []

    for change in tracked_changes:
        # Skip pure deletions with no new text (those would just delete content)
        # Only process changes that have new text or are substitutions
        if change.new_text or change.old_text:
            edit_dict = change.to_edit_dict()
            edits.append(edit_dict)

    return edits

def get_document_xml_raw_text(docx_path: str) -> str:
    """
    Extracts the raw content of word/document.xml from a DOCX file.
    (For Approach 2: Raw XML to LLM)
    """
    try:
        with zipfile.ZipFile(docx_path, 'r') as docx_zip:
            if 'word/document.xml' in docx_zip.namelist():
                with docx_zip.open('word/document.xml') as xml_file:
                    xml_content_bytes = xml_file.read()
                    return xml_content_bytes.decode('utf-8', errors='ignore') # Ignore decoding errors for robustness
            else:
                return "Error_Internal: word/document.xml not found in the DOCX archive."
    except Exception as e:
        return f"Error_Internal: Exception while processing DOCX to get raw XML: {str(e)}"

# --- END NEW/MODIFIED FUNCTIONS FOR ANALYSIS ---


def _add_comment_to_paragraph(doc, paragraph, current_para_idx: int, comment_text: str, author_name: str,
                             ambiguous_or_failed_changes_log: List[Dict],
                             edit_item_details_for_log: Optional[Dict] = None):
    # ... (keep existing _add_comment_to_paragraph) ...
    if not ADD_COMMENTS_TO_CHANGES or not comment_text: 
        log_debug(f"P{current_para_idx+1}: Skipping comment addition - ADD_COMMENTS_TO_CHANGES={ADD_COMMENTS_TO_CHANGES}, comment_text_exists={bool(comment_text)}")
        return
    log_ctx = {"paragraph_index": current_para_idx, **(edit_item_details_for_log or {})}
    try:
        author_display = f"{author_name} (LLM)"
        name_parts = [w for w in author_display.replace("(", "").replace(")", "").split() if w]
        initials = (name_parts[0][0] + name_parts[1][0]).upper() if len(name_parts) >= 2 else (name_parts[0][:2].upper() if name_parts else "AI")
        # Fix: Use document.add_comment() with paragraph.runs instead of paragraph.add_comment()
        log_debug(f"P{current_para_idx+1}: Debug - doc type: {type(doc)}, has add_comment: {hasattr(doc, 'add_comment')}")
        if paragraph.runs:  # Only add comment if paragraph has runs
            log_debug(f"P{current_para_idx+1}: Attempting to add comment with {len(paragraph.runs)} runs...")
            doc.add_comment(paragraph.runs, text=comment_text, author=author_display, initials=initials)
            log_debug(f"P{current_para_idx+1}: Successfully added comment: '{comment_text[:30]}...'.")
        else:
            log_debug(f"P{current_para_idx+1}: Cannot add comment - paragraph has no runs.")
    except AttributeError as e_attr:
        log_message = f"Comment addition failed for P{current_para_idx+1} (AttributeError): {e_attr}. Object type: {type(paragraph)}. Comment: '{comment_text[:50]}...'"
        log_debug(f"CRITICAL_WARNING - {log_message}")
        ambiguous_or_failed_changes_log.append({**log_ctx, "issue": log_message, "type": "CriticalWarning"})
    except Exception as e_gen:
        log_message = f"Comment addition failed for P{current_para_idx+1} (General Error): {e_gen}. Comment: '{comment_text[:50]}...'"
        log_debug(f"WARNING - {log_message}")
        ambiguous_or_failed_changes_log.append({**log_ctx, "issue": log_message, "type": "Warning"})


def _apply_markup_to_span(
        paragraph_obj, current_para_idx: int, global_start: int, global_end: int, text_to_markup: str,
        highlight_color: str, comment_text: str, author_name: str,
        initial_fallback_style_run: Optional[OxmlElement],
        ambiguous_or_failed_changes_log: List[Dict], edit_item_details: Dict
    ) -> bool:
    # ... (keep existing _apply_markup_to_span) ...
    log_debug(f"P{current_para_idx+1}: Applying markup: Highlight '{text_to_markup}' ({highlight_color}) at {global_start}-{global_end}, Comment: '{comment_text[:30]}...'")
    current_visible_text, current_elements_map = _build_visible_text_map(paragraph_obj)
    if not (0 <= global_start < global_end <= len(current_visible_text)):
        log_debug(f"P{current_para_idx+1}: Invalid span {global_start}-{global_end} for markup in text of len {len(current_visible_text)}. Skipping markup.")
        ambiguous_or_failed_changes_log.append({"paragraph_index": current_para_idx, "issue": f"Markup skipped: Invalid span {global_start}-{global_end} for text len {len(current_visible_text)}", **edit_item_details, "type": "MarkupError"})
        return False
    text_actually_at_span = current_visible_text[global_start:global_end]
    if text_actually_at_span != text_to_markup:
        log_debug(f"P{current_para_idx+1}: Markup text mismatch. Expected '{text_to_markup}', found '{text_actually_at_span}'. Skipping markup.")
        ambiguous_or_failed_changes_log.append({"paragraph_index": current_para_idx, "issue": f"Markup skipped: Text mismatch. Expected '{text_to_markup}', found '{text_actually_at_span}'", **edit_item_details, "type": "MarkupError"})
        return False
    involved_span_element_infos = []
    span_first_style_run = initial_fallback_style_run if initial_fallback_style_run is not None else OxmlElement('w:r')
    for item in current_elements_map:
        item_doc_end_offset = item['doc_start_offset'] + len(item['text'])
        if max(item['doc_start_offset'], global_start) < min(item_doc_end_offset, global_end):
            involved_span_element_infos.append(item)
            if span_first_style_run == initial_fallback_style_run or (span_first_style_run is not None and span_first_style_run.tag != qn('w:r')):
                 current_el_style_run = _get_element_style_template_run(item['el'], initial_fallback_style_run)
                 if current_el_style_run is not None : span_first_style_run = current_el_style_run
    if not involved_span_element_infos:
        log_debug(f"P{current_para_idx+1}: Markup failed. No XML elements identified for span {global_start}-{global_end} ('{text_to_markup}').")
        ambiguous_or_failed_changes_log.append({"paragraph_index": current_para_idx, "issue": f"Markup skipped: No XML elements for span '{text_to_markup}'", **edit_item_details, "type": "MarkupError"})
        return False
    if span_first_style_run is None or span_first_style_run.tag != qn('w:r'):
        span_first_style_run = OxmlElement('w:r')
        log_debug(f"P{current_para_idx+1}: Warning - span_first_style_run for markup was not a w:r after element search, using new default.")
    new_xml_sequence = []
    first_item_markup = involved_span_element_infos[0]
    prefix_len_markup = global_start - first_item_markup['doc_start_offset']
    prefix_text_markup = first_item_markup['text'][:prefix_len_markup] if prefix_len_markup > 0 else ""
    if prefix_text_markup:
        new_xml_sequence.append(create_run_element_with_text(prefix_text_markup, _get_element_style_template_run(first_item_markup['el'], span_first_style_run)))
    new_xml_sequence.append(create_run_element_with_text(text_to_markup, span_first_style_run, highlight_color=highlight_color))
    last_item_markup = involved_span_element_infos[-1]
    suffix_start_markup = global_end - last_item_markup['doc_start_offset']
    suffix_text_markup = last_item_markup['text'][suffix_start_markup:] if suffix_start_markup < len(last_item_markup['text']) else ""
    if suffix_text_markup:
        new_xml_sequence.append(create_run_element_with_text(suffix_text_markup, _get_element_style_template_run(last_item_markup['el'], span_first_style_run)))
    p_child_indices_to_remove_markup = sorted(list(set(item['p_child_idx'] for item in involved_span_element_infos)), reverse=True)
    if not p_child_indices_to_remove_markup:
        log_debug(f"P{current_para_idx+1}: Markup failed. No XML parent indices to remove for '{text_to_markup}'.")
        return False
    insertion_point_markup = p_child_indices_to_remove_markup[-1]
    for p_idx_remove in p_child_indices_to_remove_markup:
        try:
            element_to_remove = paragraph_obj._p[p_idx_remove]
            paragraph_obj._p.remove(element_to_remove)
        except (ValueError, IndexError) as e:
            log_debug(f"P{current_para_idx+1}: Error removing element at original index {p_idx_remove} for markup: {e}. List size: {len(paragraph_obj._p)}");
            ambiguous_or_failed_changes_log.append({"paragraph_index": current_para_idx, "issue": f"Markup error during XML removal: {e}", **edit_item_details, "type": "CriticalMarkupError"})
            return False
    for i, new_el in enumerate(new_xml_sequence):
        paragraph_obj._p.insert(insertion_point_markup + i, new_el)
    log_debug(f"P{current_para_idx+1}: Markup applied for '{text_to_markup}'.")
    _add_comment_to_paragraph(doc, paragraph_obj, current_para_idx, comment_text, author_name, ambiguous_or_failed_changes_log, edit_item_details)
    return True


def replace_text_in_paragraph_with_tracked_change(
        doc, current_para_idx: int, paragraph,
        contextual_old_text_llm, specific_old_text_llm, specific_new_text,
        reason_for_change, author, case_sensitive_flag,
        ambiguous_or_failed_changes_log) -> Tuple[str, Optional[List[Dict[str, Any]]]]:
    # ... (keep existing replace_text_in_paragraph_with_tracked_change) ...
    # Debug output disabled
    if DEBUG_MODE:
      print(f"Attempting in P{current_para_idx+1}: Context='{contextual_old_text_llm[:50]}...', SpecificOld='{specific_old_text_llm}', New='{specific_new_text}'")
    if EXTENDED_DEBUG_MODE:
        log_debug(f"P{current_para_idx+1}: Full LLM Context: '{contextual_old_text_llm}'")
        log_debug(f"P{current_para_idx+1}: Full LLM Specific Old: '{specific_old_text_llm}'")
    log_debug(f"P{current_para_idx+1}: Attempting change '{specific_old_text_llm}' to '{specific_new_text}' within LLM context '{contextual_old_text_llm[:30]}...'")
    visible_paragraph_text_original_case, elements_map = _build_visible_text_map(paragraph)
    if DEBUG_MODE or EXTENDED_DEBUG_MODE:
        pass  # Debug output disabled
    edit_details_for_log = {
        "contextual_old_text": contextual_old_text_llm,
        "specific_old_text": specific_old_text_llm,
        "specific_new_text": specific_new_text,
        "llm_reason": reason_for_change,
        "visible_text_snippet": visible_paragraph_text_original_case[:100]
    }
    if not visible_paragraph_text_original_case and (contextual_old_text_llm or specific_old_text_llm):
        log_debug(f"P{current_para_idx+1}: Paragraph is empty or yields no visible text, but an edit was provided. Skipping.");
        return "CONTEXT_NOT_FOUND", None
    search_text_in_doc = visible_paragraph_text_original_case if case_sensitive_flag else visible_paragraph_text_original_case.lower()
    search_context_from_llm_processed = contextual_old_text_llm if case_sensitive_flag else contextual_old_text_llm.lower()
    log_debug(f"P{current_para_idx+1}: Current visible paragraph text (len {len(visible_paragraph_text_original_case)}): '{visible_paragraph_text_original_case[:60]}{'...' if len(visible_paragraph_text_original_case)>60 else ''}'")
    occurrences_of_context = []
    try:
        for match in re.finditer(re.escape(search_context_from_llm_processed), search_text_in_doc):
            occurrences_of_context.append({"start": match.start(), "end": match.end(), "text": visible_paragraph_text_original_case[match.start():match.end()]})
    except re.error as e:
        ambiguous_or_failed_changes_log.append({"paragraph_index": current_para_idx, "issue": f"Regex error searching for context: {e}", **edit_details_for_log});
        return "REGEX_ERROR_IN_CONTEXT_SEARCH", None
    if not occurrences_of_context:
        log_debug(f"P{current_para_idx+1}: LLM Context '{contextual_old_text_llm[:30]}...' not found in paragraph text.");
        return "CONTEXT_NOT_FOUND", None
    if len(occurrences_of_context) > 1:
        log_debug(f"P{current_para_idx+1}: LLM Context '{contextual_old_text_llm[:30]}...' is AMBIGUOUS ({len(occurrences_of_context)} found in paragraph).")
        return "CONTEXT_AMBIGUOUS", occurrences_of_context
    unique_context_match_info = occurrences_of_context[0]
    actual_context_found_in_doc_str = unique_context_match_info["text"]
    prefix_display = visible_paragraph_text_original_case[max(0, unique_context_match_info['start']-10) : unique_context_match_info['start']]
    suffix_display = visible_paragraph_text_original_case[unique_context_match_info['end'] : min(len(visible_paragraph_text_original_case), unique_context_match_info['end']+10)]
    log_debug(f"P{current_para_idx+1}: Unique LLM context found: '...{prefix_display}[{actual_context_found_in_doc_str}]{suffix_display}...' at {unique_context_match_info['start']}-{unique_context_match_info['end']}")
    text_to_search_specific_within = actual_context_found_in_doc_str if case_sensitive_flag else actual_context_found_in_doc_str.lower()
    specific_text_to_find_llm_processed = specific_old_text_llm if case_sensitive_flag else specific_old_text_llm.lower()
    try:
        specific_start_in_context = text_to_search_specific_within.index(specific_text_to_find_llm_processed)
        actual_specific_old_text_to_delete = actual_context_found_in_doc_str[specific_start_in_context : specific_start_in_context + len(specific_old_text_llm)]
        fuzzy_match_used = False
        fuzzy_similarity = 1.0
    except ValueError:
        # Try fuzzy matching as fallback
        fuzzy_match = fuzzy_search_best_match(specific_text_to_find_llm_processed, text_to_search_specific_within)
        if fuzzy_match:
            specific_start_in_context = fuzzy_match['start']
            actual_specific_old_text_to_delete = fuzzy_match['matched_text']
            fuzzy_match_used = True
            fuzzy_similarity = fuzzy_match['similarity']
            log_debug(f"P{current_para_idx+1}: Exact match failed, using fuzzy match: '{actual_specific_old_text_to_delete}' (similarity: {fuzzy_similarity:.2f})")
        else:
            log_debug(f"P{current_para_idx+1}: Specific text '{specific_old_text_llm}' NOT FOUND within the unique context '{actual_context_found_in_doc_str}'. Change skipped.")
            ambiguous_or_failed_changes_log.append({"paragraph_index": current_para_idx, "issue": "Specific text not found within unique context (exact or fuzzy).", **edit_details_for_log});
            return "SPECIFIC_TEXT_NOT_IN_CONTEXT", None
    global_specific_start_offset = unique_context_match_info['start'] + specific_start_in_context
    global_specific_end_offset = global_specific_start_offset + len(actual_specific_old_text_to_delete)
    log_debug(f"P{current_para_idx+1}: LLM specific_old_text: '{specific_old_text_llm}' (len {len(specific_old_text_llm)})")
    log_debug(f"P{current_para_idx+1}: Actual specific_old_text_to_delete (from doc): '{actual_specific_old_text_to_delete}' (len {len(actual_specific_old_text_to_delete)})")
    log_debug(f"P{current_para_idx+1}: Global offsets for specific text: {global_specific_start_offset}-{global_specific_end_offset}")
    if len(actual_specific_old_text_to_delete) != len(specific_old_text_llm):
        log_debug(f"P{current_para_idx+1}: CRITICAL WARNING! Length mismatch between LLM specific_old_text ('{specific_old_text_llm}') and actual text found in doc to delete ('{actual_specific_old_text_to_delete}'). This will likely cause incorrect changes.")
        ambiguous_or_failed_changes_log.append({"paragraph_index": current_para_idx, "issue": "Length mismatch between LLM specific_old_text and actual text to delete.", "type":"CriticalWarning", **edit_details_for_log})
    char_before_specific = visible_paragraph_text_original_case[global_specific_start_offset - 1] if global_specific_start_offset > 0 else None
    char_after_specific = visible_paragraph_text_original_case[global_specific_end_offset] if global_specific_end_offset < len(visible_paragraph_text_original_case) else None
    
    # Use fuzzy boundary checking if fuzzy match was used
    if fuzzy_match_used:
        boundary_valid = is_boundary_valid_fuzzy(actual_specific_old_text_to_delete, global_specific_start_offset, global_specific_end_offset, visible_paragraph_text_original_case, fuzzy_similarity)
        boundary_type = "fuzzy"
    else:
        is_start_boundary_ok = (global_specific_start_offset == 0 or (char_before_specific is not None and char_before_specific.isspace()))
        is_end_boundary_ok = (global_specific_end_offset == len(visible_paragraph_text_original_case) or (char_after_specific is not None and (char_after_specific.isspace() or char_after_specific in ALLOWED_POST_BOUNDARY_PUNCTUATION)))
        boundary_valid = is_start_boundary_ok and is_end_boundary_ok
        boundary_type = "exact"
    
    if not boundary_valid:
        match_info = f" (fuzzy match, similarity: {fuzzy_similarity:.2f})" if fuzzy_match_used else ""
        log_message = (f"P{current_para_idx+1}: Specific text '{actual_specific_old_text_to_delete}' is NOT validly bounded{match_info}. " f"Preceded by: '{char_before_specific}', Followed by: '{char_after_specific}'. Change skipped.")
        log_debug(log_message)
        ambiguous_or_failed_changes_log.append({"paragraph_index": current_para_idx, "issue": f"Skipped: Specific text not validly bounded (Prev: '{char_before_specific}', Next: '{char_after_specific}') [{boundary_type} match]", **edit_details_for_log})
        return "SKIPPED_INVALID_BOUNDARY", None
    
    if fuzzy_match_used:
        log_debug(f"P{current_para_idx+1}: ✅ Fuzzy match boundary validation passed (similarity: {fuzzy_similarity:.2f})")
        print(f"SUCCESS: P{current_para_idx+1}: Using fuzzy match (similarity: {fuzzy_similarity:.2f}) for '{actual_specific_old_text_to_delete}'")
    involved_element_infos = []
    first_involved_r_element_for_style = None
    for item_map_entry in elements_map:
        item_doc_end_offset = item_map_entry['doc_start_offset'] + len(item_map_entry['text'])
        if max(item_map_entry['doc_start_offset'], global_specific_start_offset) < min(item_doc_end_offset, global_specific_end_offset):
            involved_element_infos.append(item_map_entry)
            if first_involved_r_element_for_style is None:
                first_involved_r_element_for_style = _get_element_style_template_run(item_map_entry['el'], None)
    if not involved_element_infos:
        log_debug(f"P{current_para_idx+1}: XML_MAPPING_FAILED for '{actual_specific_old_text_to_delete}'. No XML elements correspond to the identified text span.")
        ambiguous_or_failed_changes_log.append({"paragraph_index": current_para_idx, "issue": "XML mapping failed for specific text.", **edit_details_for_log});
        return "XML_MAPPING_FAILED", None
    if first_involved_r_element_for_style is None:
        first_involved_r_element_for_style = OxmlElement('w:r')
        log_debug(f"P{current_para_idx+1}: WARNING - No template <w:r> for styling the change. Using a new default <w:r>.")
    log_debug(f"P{current_para_idx+1}: Modifying {len(involved_element_infos)} raw XML elements for the change of '{actual_specific_old_text_to_delete}' (validly bounded).")
    new_xml_elements_for_paragraph = []
    first_involved_item_details = involved_element_infos[0]
    prefix_len_in_first_item = global_specific_start_offset - first_involved_item_details['doc_start_offset']
    prefix_text_content = ""
    if prefix_len_in_first_item < 0:
        log_debug(f"P{current_para_idx+1}: WARNING: prefix_len_in_first_item ({prefix_len_in_first_item}) is negative. Clamping to 0. This might indicate an offset issue.")
        prefix_len_in_first_item = 0
    if prefix_len_in_first_item > 0 :
        prefix_text_content = first_involved_item_details['text'][:prefix_len_in_first_item]
    log_debug(f"P{current_para_idx+1}: First involved item for change: text='{first_involved_item_details['text']}', doc_start_offset={first_involved_item_details['doc_start_offset']}, type='{first_involved_item_details['type']}'")
    log_debug(f"P{current_para_idx+1}: Calculated prefix_len_in_first_item: {prefix_len_in_first_item}, resulting prefix_text_content: '{prefix_text_content}'")
    if prefix_text_content:
        style_run_for_prefix = _get_element_style_template_run(first_involved_item_details['el'], first_involved_r_element_for_style)
        if first_involved_item_details['type'] == 'ins':
            original_ins_el = create_ins_element(author=first_involved_item_details['original_author'], date_time=None)
            if first_involved_item_details['original_date'] is not None: original_ins_el.set(qn('w:date'), first_involved_item_details['original_date'])
            original_ins_el.append(create_run_element_with_text(prefix_text_content, style_run_for_prefix))
            new_xml_elements_for_paragraph.append(original_ins_el)
        else:
            new_xml_elements_for_paragraph.append(create_run_element_with_text(prefix_text_content, style_run_for_prefix))
        log_debug(f"P{current_para_idx+1}: Added prefix '{prefix_text_content}' from first element (type: {first_involved_item_details['type']}).")
    change_time = datetime.datetime.now(datetime.timezone.utc)
    del_obj = create_del_element(author=author, date_time=change_time)
    del_run_el = create_run_element_with_text(actual_specific_old_text_to_delete, first_involved_r_element_for_style, is_del_text=True)
    del_obj.append(del_run_el)
    new_xml_elements_for_paragraph.append(del_obj)
    log_debug(f"P{current_para_idx+1}: Added <w:del> for '{actual_specific_old_text_to_delete}'.")
    if specific_new_text:
        ins_obj = create_ins_element(author=author, date_time=change_time + datetime.timedelta(seconds=1))
        ins_run_el = create_run_element_with_text(specific_new_text, first_involved_r_element_for_style, is_del_text=False, highlight_color=None)
        ins_obj.append(ins_run_el)
        new_xml_elements_for_paragraph.append(ins_obj)
        log_debug(f"P{current_para_idx+1}: Added <w:ins> for '{specific_new_text}'.")
    last_involved_item_details = involved_element_infos[-1]
    suffix_start_in_last_item = global_specific_end_offset - last_involved_item_details['doc_start_offset']
    suffix_text_content = ""
    if suffix_start_in_last_item < 0:
        log_debug(f"P{current_para_idx+1}: WARNING: suffix_start_in_last_item ({suffix_start_in_last_item}) is negative. Clamping. This might indicate an offset issue.")
        suffix_start_in_last_item = len(last_involved_item_details['text'])
    if suffix_start_in_last_item < len(last_involved_item_details['text']):
        suffix_text_content = last_involved_item_details['text'][suffix_start_in_last_item:]
    log_debug(f"P{current_para_idx+1}: Last involved item for change: text='{last_involved_item_details['text']}', doc_start_offset={last_involved_item_details['doc_start_offset']}, type='{last_involved_item_details['type']}'")
    log_debug(f"P{current_para_idx+1}: Calculated suffix_start_in_last_item: {suffix_start_in_last_item}, resulting suffix_text_content: '{suffix_text_content}'")
    if suffix_text_content:
        style_run_for_suffix = _get_element_style_template_run(last_involved_item_details['el'], first_involved_r_element_for_style)
        if last_involved_item_details['type'] == 'ins':
            original_ins_el_suffix = create_ins_element(author=last_involved_item_details['original_author'], date_time=None)
            if last_involved_item_details['original_date'] is not None: original_ins_el_suffix.set(qn('w:date'), last_involved_item_details['original_date'])
            original_ins_el_suffix.append(create_run_element_with_text(suffix_text_content, style_run_for_suffix))
            new_xml_elements_for_paragraph.append(original_ins_el_suffix)
        else:
            new_xml_elements_for_paragraph.append(create_run_element_with_text(suffix_text_content, style_run_for_suffix))
        log_debug(f"P{current_para_idx+1}: Added suffix '{suffix_text_content}' from last element (type: {last_involved_item_details['type']}).")
    p_child_indices_to_remove = sorted(list(set(item['p_child_idx'] for item in involved_element_infos)), reverse=True)
    if not p_child_indices_to_remove:
        log_debug(f"P{current_para_idx+1}: XML_REMOVAL_ERROR_NO_INDICES. No paragraph child indices identified for removal. This change cannot be applied. For text '{actual_specific_old_text_to_delete}'.")
        ambiguous_or_failed_changes_log.append({"paragraph_index": current_para_idx, "issue": "XML_REMOVAL_ERROR_NO_INDICES: No elements to remove.", **edit_details_for_log});
        return "XML_REMOVAL_ERROR_NO_INDICES", None
    insertion_point_xml = p_child_indices_to_remove[-1]
    for p_idx_to_remove_loop in p_child_indices_to_remove:
        try:
            element_to_remove = paragraph._p[p_idx_to_remove_loop]
            paragraph._p.remove(element_to_remove)
            log_debug(f"P{current_para_idx+1}: Removed original XML element that was at original p_child_idx {p_idx_to_remove_loop}.")
        except (IndexError, ValueError) as e_remove:
            log_message = f"P{current_para_idx+1}: XML element removal error at original index {p_idx_to_remove_loop}. Error: {e_remove}. Change aborted for '{actual_specific_old_text_to_delete}'."
            log_debug(log_message)
            ambiguous_or_failed_changes_log.append({"paragraph_index": current_para_idx, "issue": f"XML element removal error at index {p_idx_to_remove_loop}: {e_remove}", "type": "CriticalError", **edit_details_for_log})
            return "XML_REMOVAL_ERROR", None
    for i, new_el in enumerate(new_xml_elements_for_paragraph):
        paragraph._p.insert(insertion_point_xml + i, new_el)
    log_debug(f"P{current_para_idx+1}: Inserted {len(new_xml_elements_for_paragraph)} new XML elements at original index {insertion_point_xml}.")
    comment_to_add = reason_for_change
    if not specific_new_text:
        comment_to_add = f"Deleted: '{actual_specific_old_text_to_delete}'. Reason: {reason_for_change}"
    _add_comment_to_paragraph(doc, paragraph, current_para_idx, comment_to_add, author, ambiguous_or_failed_changes_log, edit_details_for_log)
    return "SUCCESS", None


def process_document_with_edits(
    input_docx_path: str, output_docx_path: str, edits_to_make: List[Dict],
    author_name: str = DEFAULT_AUTHOR_NAME,
    debug_mode_flag: bool = False,
    extended_debug_mode_flag: bool = False,
    case_sensitive_flag: bool = True,
    add_comments_param: bool = True
) -> Tuple[bool, Optional[str], List[Dict], int]:
    # ... (keep existing process_document_with_edits, ensuring it uses the global DEBUG_MODE flags correctly) ...
    # Text extraction verification disabled
    # Text extraction verification disabled
    error_log_file_path: Optional[str] = None
    global DEBUG_MODE, EXTENDED_DEBUG_MODE, CASE_SENSITIVE_SEARCH, ADD_COMMENTS_TO_CHANGES
    DEBUG_MODE = debug_mode_flag  # Use parameter from caller
    EXTENDED_DEBUG_MODE = extended_debug_mode_flag  # Use parameter from caller
    CASE_SENSITIVE_SEARCH = case_sensitive_flag
    ADD_COMMENTS_TO_CHANGES = add_comments_param  # Use parameter from caller
    # Debug output disabled
    log_debug(f"Script starting. Input: {input_docx_path}, Output: {output_docx_path}")
    log_debug(f"Settings - Debug:{DEBUG_MODE}, ExtDebug:{EXTENDED_DEBUG_MODE}, CaseSensitive:{CASE_SENSITIVE_SEARCH}, AddComments:{ADD_COMMENTS_TO_CHANGES}, Author:{author_name}")
    log_debug(f"Number of edits to attempt: {len(edits_to_make)}")
    if not isinstance(edits_to_make, list) or not all(isinstance(item,dict) for item in edits_to_make):
        return False, error_log_file_path, [{"issue": "FATAL: Edits must be a list of dictionaries."}], 0
    ambiguous_or_failed_changes_log: List[Dict] = []
    try:
        doc = Document(input_docx_path)
        log_debug(f"Successfully opened '{input_docx_path}'")
    except Exception as e:
        return False, error_log_file_path, [{"issue": f"FATAL: Error opening Word document '{input_docx_path}': {e}"}], 0
    edits_successfully_applied_count = 0
    for para_idx, paragraph_obj in enumerate(doc.paragraphs):
        _initial_para_text_for_style, _initial_para_elements_map_for_style = _build_visible_text_map(paragraph_obj)
        fallback_style_run_for_markup = OxmlElement('w:r')
        if _initial_para_elements_map_for_style:
            first_r_el_in_para = next((item['el'] for item in _initial_para_elements_map_for_style if item['el'].tag == qn('w:r')), None)
            if first_r_el_in_para is not None:
                fallback_style_run_for_markup = _get_element_style_template_run(first_r_el_in_para, fallback_style_run_for_markup)
        log_debug(f"\n--- Processing Paragraph {para_idx+1} (Initial Text Snapshot: '{paragraph_obj.text[:60]}{'...' if len(paragraph_obj.text)>60 else ''}') ---")
        for edit_item_idx, edit_item in enumerate(list(edits_to_make)):
            log_debug(f"P{para_idx+1}: Attempting edit item {edit_item_idx+1} ('{edit_item.get('specific_old_text')}' -> '{edit_item.get('specific_new_text')}')")
            required_keys = ["contextual_old_text", "specific_old_text", "reason_for_change"]
            if not all(key in edit_item for key in required_keys):
                log_message = f"P{para_idx+1}, EditItem{edit_item_idx+1}: Invalid structure (missing keys): {str(edit_item)[:100]}. Skipping."
                log_debug(log_message)
                ambiguous_or_failed_changes_log.append({"paragraph_index": para_idx, "edit_item_index": edit_item_idx +1, "issue": "Invalid edit item structure.", "edit_item_snippet": str(edit_item)[:100]})
                continue
            status, data_from_replace = "INIT", None
            current_edit_details_for_log = {
                "contextual_old_text": edit_item["contextual_old_text"],
                "specific_old_text": edit_item["specific_old_text"],
                "specific_new_text": edit_item.get("specific_new_text",""),
                "llm_reason": edit_item["reason_for_change"]
            }
            try:
                status, data_from_replace = replace_text_in_paragraph_with_tracked_change(
                    doc, para_idx, paragraph_obj,
                    edit_item["contextual_old_text"], edit_item["specific_old_text"],
                    edit_item.get("specific_new_text",""),
                    edit_item["reason_for_change"],
                    author_name, CASE_SENSITIVE_SEARCH,
                    ambiguous_or_failed_changes_log
                )
            except Exception as e_replace_call:
                status = "EXCEPTION_IN_REPLACE_CALL"
                log_message = f"P{para_idx+1}, EditItem{edit_item_idx+1}: Unhandled EXCEPTION during *call* to replacement function for '{edit_item['specific_old_text']}'. Error: {type(e_replace_call).__name__}: {e_replace_call}"
                log_debug(log_message)
                if DEBUG_MODE: log_debug(traceback.format_exc())
                ambiguous_or_failed_changes_log.append({
                    "paragraph_index": para_idx, "edit_item_index": edit_item_idx +1,
                    "issue": f"Unhandled exception in replacement call: {e_replace_call}", "type": "CriticalError",
                    **current_edit_details_for_log
                })
            if status == "SUCCESS":
                success_msg = f"SUCCESS: P{para_idx+1}: Applied change for context '{edit_item['contextual_old_text'][:30]}...', specific '{edit_item['specific_old_text']}'. Reason: {edit_item['reason_for_change']}"
                print(success_msg)
                log_debug(success_msg)
                edits_successfully_applied_count += 1
            elif status == "CONTEXT_AMBIGUOUS":
                log_debug(f"P{para_idx+1}: Context '{edit_item['contextual_old_text'][:30]}...' was AMBIGUOUS for specific text '{edit_item['specific_old_text']}'. Attempting markup.")
                if data_from_replace and isinstance(data_from_replace, list):
                    spans_to_markup_this_edit_item = []
                    for ctx_occurrence in data_from_replace:
                        ctx_text_original_case = ctx_occurrence["text"]
                        ctx_start_global = ctx_occurrence["start"]
                        s_old_llm_val = edit_item["specific_old_text"]
                        s_old_search_val = s_old_llm_val if CASE_SENSITIVE_SEARCH else s_old_llm_val.lower()
                        ctx_text_search_val = ctx_text_original_case if CASE_SENSITIVE_SEARCH else ctx_text_original_case.lower()
                        current_offset_in_ctx = 0
                        while current_offset_in_ctx < len(ctx_text_search_val):
                            try:
                                found_idx_in_ctx = ctx_text_search_val.index(s_old_search_val, current_offset_in_ctx)
                                specific_text_to_markup_val = ctx_text_original_case[found_idx_in_ctx : found_idx_in_ctx + len(s_old_llm_val)]
                                specific_start_abs = ctx_start_global + found_idx_in_ctx
                                specific_end_abs = specific_start_abs + len(specific_text_to_markup_val)
                                spans_to_markup_this_edit_item.append({"start":specific_start_abs, "end":specific_end_abs, "text":specific_text_to_markup_val})
                                current_offset_in_ctx = found_idx_in_ctx + len(specific_text_to_markup_val)
                            except ValueError:
                                # Try fuzzy matching for ambiguous contexts too
                                remaining_text = ctx_text_search_val[current_offset_in_ctx:]
                                fuzzy_match = fuzzy_search_best_match(s_old_search_val, remaining_text)
                                if fuzzy_match:
                                    found_idx_in_ctx = current_offset_in_ctx + fuzzy_match['start']
                                    specific_text_to_markup_val = fuzzy_match['matched_text']
                                    specific_start_abs = ctx_start_global + found_idx_in_ctx
                                    specific_end_abs = specific_start_abs + len(specific_text_to_markup_val)
                                    spans_to_markup_this_edit_item.append({"start":specific_start_abs, "end":specific_end_abs, "text":specific_text_to_markup_val})
                                    current_offset_in_ctx = found_idx_in_ctx + len(specific_text_to_markup_val)
                                    log_debug(f"P{para_idx+1}: Used fuzzy match in ambiguous context: '{specific_text_to_markup_val}' (similarity: {fuzzy_match['similarity']:.2f})")
                                else:
                                    break
                    spans_to_markup_this_edit_item.sort(key=lambda x:x["start"], reverse=True)
                    applied_any_markup_for_this_ambiguity = False
                    for span_info in spans_to_markup_this_edit_item:
                        if _apply_markup_to_span(paragraph_obj, para_idx,
                                               span_info["start"], span_info["end"], span_info["text"],
                                               HIGHLIGHT_COLOR_AMBIGUOUS_SKIPPED,
                                               f"Skipped change: Ambiguous context for '{edit_item['specific_old_text']}'. This is one of multiple occurrences. LLM Reason for original change: {edit_item['reason_for_change']}",
                                               author_name, fallback_style_run_for_markup,
                                               ambiguous_or_failed_changes_log, current_edit_details_for_log):
                            applied_any_markup_for_this_ambiguity = True
                    if applied_any_markup_for_this_ambiguity:
                        ambiguous_or_failed_changes_log.append({"paragraph_index":para_idx, "issue": f"CONTEXT_AMBIGUOUS: Marked {len(spans_to_markup_this_edit_item)} instance(s) of '{edit_item['specific_old_text']}' with orange highlight.", "type": "Info", **current_edit_details_for_log})
                    else:
                        log_debug(f"P{para_idx+1}: CONTEXT_AMBIGUOUS status, but no specific text instances were marked up for '{edit_item['specific_old_text']}'.")
                        ambiguous_or_failed_changes_log.append({"paragraph_index":para_idx, "issue": "CONTEXT_AMBIGUOUS, but no markup applied (specific text not in contexts or markup failed).", "type": "Warning", **current_edit_details_for_log})
                else:
                    log_debug(f"P{para_idx+1}: CONTEXT_AMBIGUOUS status but no occurrence data was returned by replacement function. Edit item: {edit_item}")
                    ambiguous_or_failed_changes_log.append({"paragraph_index":para_idx, "issue": "CONTEXT_AMBIGUOUS but no occurrence data returned from replace function.", "type": "Warning", **current_edit_details_for_log})
            elif status == "CONTEXT_NOT_FOUND":
                log_debug(f"P{para_idx+1}: Edit skipped - CONTEXT_NOT_FOUND for '{edit_item['specific_old_text']}'")
                ambiguous_or_failed_changes_log.append({
                    "paragraph_index": para_idx,
                    "issue": "Context not found in paragraph text.",
                    "type": "Skipped",
                    **current_edit_details_for_log
                })
            elif status == "SPECIFIC_TEXT_NOT_IN_CONTEXT":
                # Already logged in replace_text_in_paragraph_with_tracked_change
                pass
            elif status == "SKIPPED_INVALID_BOUNDARY":
                # Already logged in replace_text_in_paragraph_with_tracked_change
                pass
            elif status == "XML_MAPPING_FAILED":
                # Already logged in replace_text_in_paragraph_with_tracked_change
                pass
            elif status == "REGEX_ERROR_IN_CONTEXT_SEARCH":
                # Already logged in replace_text_in_paragraph_with_tracked_change
                pass
            elif status == "XML_REMOVAL_ERROR_NO_INDICES":
                # Should be logged elsewhere if needed
                pass
            elif status not in ["INIT"]:
                info_msg = f"INFO: P{para_idx+1}: Edit for context '{edit_item['contextual_old_text'][:30]}...' specific '{edit_item['specific_old_text']}' resulted in status: {status}."
                print(info_msg)
                log_debug(info_msg)
            if status in ["XML_REMOVAL_ERROR", "XML_REMOVAL_ERROR_NO_INDICES", "EXCEPTION_IN_REPLACE_CALL"]:
                log_debug(f"P{para_idx+1}: Critical error status '{status}' encountered. Halting further edit attempts for THIS PARAGRAPH.")
                break
    try:
        doc.save(output_docx_path)
        print(f"\nProcessed document saved to '{output_docx_path}'")
        log_debug(f"Processed document saved to '{output_docx_path}'")
    except Exception as e:
        log_debug(f"FATAL: Error saving document to '{output_docx_path}': {e}")
        ambiguous_or_failed_changes_log.append({"issue": f"FATAL: Error saving document to '{output_docx_path}': {e}", "type":"FatalError"})
        return False, None, ambiguous_or_failed_changes_log, edits_successfully_applied_count
    if ambiguous_or_failed_changes_log:
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        log_filename_with_ts = f"{ERROR_LOG_FILENAME_BASE}_{timestamp}.txt"
        output_dir = os.path.dirname(os.path.abspath(output_docx_path))
        if output_dir and not os.path.exists(output_dir):
            try: os.makedirs(output_dir, exist_ok=True)
            except Exception as e_mkdir:
                log_debug(f"Could not create output directory '{output_dir}' for log file: {e_mkdir}. Log will be placed in script dir.")
                output_dir = ""
        error_log_file_path = os.path.join(output_dir, log_filename_with_ts) if output_dir else log_filename_with_ts
        try:
            with open(error_log_file_path, "w", encoding="utf-8") as f:
                f.write(f"--- LOG OF CHANGES NOT MADE, AMBIGUITIES, OR WARNINGS ({datetime.datetime.now()}) ---\n")
                f.write(f"Input DOCX: {os.path.basename(input_docx_path)}\n")
                f.write(f"Output DOCX: {os.path.basename(output_docx_path)}\n")
                f.write(f"Total Edit Instructions Provided: {len(edits_to_make)}\n")
                f.write(f"Edits Successfully Applied This Run: {edits_successfully_applied_count}\n")
                f.write(f"Log Items (Failures/Warnings/Errors/Info): {len(ambiguous_or_failed_changes_log)}\n\n")
                for log_entry in ambiguous_or_failed_changes_log:
                    f.write("-----------------------------------------\n")
                    para_display_index = log_entry.get('paragraph_index', -1)
                    if isinstance(para_display_index, int) and para_display_index >=0 : para_display_index +=1
                    else: para_display_index = 'N/A'
                    f.write(f"Paragraph Index (1-based): {para_display_index}\n")
                    f.write(f"Original Visible Text Snippet (at time of processing this item): {log_entry.get('visible_text_snippet', 'N/A')}\n")
                    f.write(f"LLM Context Searched: '{log_entry.get('contextual_old_text', 'N/A')}'\n")
                    f.write(f"LLM Specific Old Text: '{log_entry.get('specific_old_text', 'N/A')}'\n")
                    f.write(f"LLM Specific New Text: '{log_entry.get('specific_new_text', 'N/A')}'\n")
                    f.write(f"LLM Reason for Change: '{log_entry.get('llm_reason', 'N/A')}'\n")
                    f.write(f"Issue/Status: {log_entry.get('issue', 'Unknown')}\n")
                    log_type = log_entry.get('type', 'Log')
                    f.write(f"Log Entry Type: {log_type}\n")
                    if 'edit_item_index' in log_entry: f.write(f"Edit Item Index (0-based in list): {log_entry['edit_item_index']}\n")
                    if 'edit_item_snippet' in log_entry: f.write(f"Edit Item Snippet: {log_entry['edit_item_snippet']}\n")
                f.write("-----------------------------------------\n")
            print(f"Log file with {len(ambiguous_or_failed_changes_log)} items saved to: '{error_log_file_path}'")
            log_debug(f"Log file saved to: '{error_log_file_path}'")
        except Exception as e_log:
            log_debug(f"ERROR: Could not write to log file '{error_log_file_path}': {e_log}")
            error_log_file_path = None
    else:
        if not edits_to_make: print(f"No edits were provided to process.")
        elif edits_successfully_applied_count == len(edits_to_make) and len(edits_to_make) > 0 :
             print(f"All {edits_successfully_applied_count} targeted changes were applied successfully. No issues logged.")
        elif edits_successfully_applied_count < len(edits_to_make) and edits_successfully_applied_count > 0:
            print(f"{edits_successfully_applied_count} changes applied. Some of the {len(edits_to_make)} edits may not have found their context, were ambiguous, or skipped. No critical errors logged that prevented saving.")
        elif edits_successfully_applied_count == 0 and len(edits_to_make) > 0:
            print(f"No changes were applied out of {len(edits_to_make)} provided. Edits may not have found their context, were ambiguous, or skipped. No critical errors logged that prevented saving.")
        log_debug("No items for error log file as ambiguous_or_failed_changes_log was empty.")
    return True, error_log_file_path, ambiguous_or_failed_changes_log, edits_successfully_applied_count


if __name__ == '__main__':
    # ... (keep existing __main__ block) ...
    import argparse
    DEFAULT_EDITS_FILE_PATH = "edits_to_apply.json"
    DEFAULT_INPUT_DOCX_PATH = "sample_input.docx"
    DEFAULT_OUTPUT_DOCX_PATH = "sample_output_corrected_v3.docx"
    parser = argparse.ArgumentParser(description="Apply tracked changes to a Word document.")
    parser.add_argument("--input", default=DEFAULT_INPUT_DOCX_PATH, help=f"Input DOCX (default: {DEFAULT_INPUT_DOCX_PATH})")
    parser.add_argument("--output", default=DEFAULT_OUTPUT_DOCX_PATH, help=f"Output DOCX (default: {DEFAULT_OUTPUT_DOCX_PATH})")
    parser.add_argument("--editsfile", default=DEFAULT_EDITS_FILE_PATH, help=f"JSON edits file (default: {DEFAULT_EDITS_FILE_PATH})")
    parser.add_argument("--editsjson", type=str, help="JSON string of edits.")
    parser.add_argument("--author", default=DEFAULT_AUTHOR_NAME, help=f"Author for changes (default: {DEFAULT_AUTHOR_NAME}).")
    parser.add_argument("--debug", action="store_true", help="Enable debug mode (prints to console).")
    parser.add_argument("--extended_debug", action="store_true", help="Enable extended debug mode for very verbose context logging (prints to console).")
    parser.add_argument("--caseinsensitive", action="store_false", dest="case_sensitive", help="Perform case-insensitive search for context and specific text.")
    parser.set_defaults(case_sensitive=True)
    parser.add_argument("--nocomments", action="store_false", dest="add_comments", help="Disable adding comments alongside tracked changes.")
    parser.set_defaults(add_comments=True)
    args = parser.parse_args()
    if args.debug: DEBUG_MODE = True
    if args.extended_debug: EXTENDED_DEBUG_MODE = True
    edits_data = []
    if args.editsjson:
        try:
            edits_data = json.loads(args.editsjson)
            log_debug(f"Loaded {len(edits_data)} edits from --editsjson argument.")
        except json.JSONDecodeError as e:
            print(f"FATAL: Error decoding JSON from --editsjson: {e}. Exiting."); exit(1)
    elif args.editsfile:
        try:
            with open(args.editsfile, 'r', encoding='utf-8') as f: edits_data = json.load(f)
            log_debug(f"Successfully loaded {len(edits_data)} edits from '{args.editsfile}'.")
        except FileNotFoundError:
            print(f"FATAL: Edits file '{args.editsfile}' not found. Exiting. Create it or use --editsjson.")
            if DEBUG_MODE and args.editsfile == DEFAULT_EDITS_FILE_PATH:
                 print(f"Attempting to create dummy '{DEFAULT_EDITS_FILE_PATH}' for testing.")
                 try:
                    dummy_edits_for_file = [
                        {"contextual_old_text": "cost would be $101 , to a new number", "specific_old_text": "$101", "specific_new_text": "$777", "reason_for_change": "Test $101 from file, expect SUCCESS if input doc is correct."},
                        {"contextual_old_text": "last edited by MrArbor, but that name", "specific_old_text": "MrArbor", "specific_new_text": "DrArbor", "reason_for_change": "Test MrArbor from file."},
                        {"contextual_old_text": "Bob started the sentence", "specific_old_text": "Bob", "specific_new_text": "Robert", "reason_for_change": "Test Bob from file."}
                    ]
                    with open(DEFAULT_EDITS_FILE_PATH, 'w', encoding='utf-8') as df: json.dump(dummy_edits_for_file, df, indent=2)
                    print(f"Created dummy edits file '{DEFAULT_EDITS_FILE_PATH}'. Please re-run.")
                 except Exception as e_create_dummy: print(f"Could not create dummy edits file: {e_create_dummy}")
            exit(1)
        except json.JSONDecodeError as e: print(f"FATAL: Error decoding JSON from '{args.editsfile}': {e}. Exiting."); exit(1)
        except Exception as e: print(f"FATAL: An unexpected error occurred while loading '{args.editsfile}': {e}. Exiting."); exit(1)
    else:
        if not DEBUG_MODE: DEBUG_MODE = True
        # Debug output disabled
        log_debug("No edits provided via CLI. Using internal dummy edits for testing.")
        edits_data = [
            {"contextual_old_text": "cost would be $101 , to a new number", "specific_old_text": "$101", "specific_new_text": "$999", "reason_for_change": "Dummy change: Update cost from $101 to $999"},
            {"contextual_old_text": "last edited by MrArbor, but that name", "specific_old_text": "MrArbor", "specific_new_text": "ProfSage", "reason_for_change": "Dummy change: Update MrArbor to ProfSage"},
            {"contextual_old_text": "Bob started the sentence", "specific_old_text": "Bob", "specific_new_text": "Robert", "reason_for_change": "Dummy change: Update Bob to Robert (1st instance)."},
            {"contextual_old_text": "Bob was also in the middle", "specific_old_text": "Bob", "specific_new_text": "Robert", "reason_for_change": "Dummy change: Update Bob to Robert (2nd instance)."},
            {"contextual_old_text": "changed ok bob", "specific_old_text": "bob", "specific_new_text": "Robert", "reason_for_change": "Dummy change: Update 'bob' (lowercase) to Robert."}
        ]
        log_debug(f"Using {len(edits_data)} internal dummy edits for testing.")
        if not os.path.exists(DEFAULT_INPUT_DOCX_PATH):
            print(f"INFO: Dummy input file '{DEFAULT_INPUT_DOCX_PATH}' not found. Creating it for testing.")
            try:
                doc_dummy = Document()
                doc_dummy.add_paragraph("A simple file seeing if tracked changes program can work. It should change this sentence from saying the cost would be $101 , to a new number.")
                doc_dummy.add_paragraph("Bob started the sentence but Bob was also in the middle, and Bobby goes by Robert so Bob-words or bob-words or any$bob$word should be changed ok bob")
                doc_dummy.add_paragraph(" ") 
                doc_dummy.add_paragraph("Another line after blank lines. Lets count 1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11.")
                doc_dummy.add_paragraph(" ") 
                doc_dummy.add_paragraph("Here is a long long long long long long long long long long repetitive line with 9 longs.")
                doc_dummy.add_paragraph(" ")
                doc_dummy.add_paragraph("The last line was last edited by MrArbor, but that name can change.")
                doc_dummy.save(DEFAULT_INPUT_DOCX_PATH)
                print(f"Created dummy input file: '{DEFAULT_INPUT_DOCX_PATH}'")
            except Exception as e_doc: print(f"FATAL: Could not create dummy input file '{DEFAULT_INPUT_DOCX_PATH}': {e_doc}"); exit(1)
    if not os.path.exists(args.input): print(f"FATAL: Input file '{args.input}' not found. Exiting."); exit(1)
    process_document_with_edits(
        args.input, args.output, edits_data, args.author, args.debug,
        args.extended_debug, args.case_sensitive, args.add_comments
    )