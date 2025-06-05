#!/usr/bin/env python3
import copy
import datetime
import re
import json
import os
import traceback 
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from typing import List, Dict, Tuple, Optional, Any

# --- Global Configuration Flags ---
DEBUG_MODE = False
EXTENDED_DEBUG_MODE = False 
CASE_SENSITIVE_SEARCH = True
ADD_COMMENTS_TO_CHANGES = True 
DEFAULT_AUTHOR_NAME = "LLM Editor"

# --- Constants ---
ERROR_LOG_FILENAME_BASE = "change_log"
HIGHLIGHT_COLOR_AMBIGUOUS_SKIPPED = "orange"
ALLOWED_POST_BOUNDARY_PUNCTUATION = {',', ';', '.'}

# --- XML Helper Functions ---
def create_del_element(author="Python Program", date_time=None):
    if date_time is None: date_time = datetime.datetime.now(datetime.timezone.utc)
    del_el = OxmlElement('w:del')
    del_el.set(qn('w:author'), author)
    del_el.set(qn('w:date'), date_time.strftime("%Y-%m-%dT%H:%M:%SZ"))
    return del_el

def create_ins_element(author="Python Program", date_time=None):
    if date_time is None: date_time = datetime.datetime.now(datetime.timezone.utc)
    ins_el = OxmlElement('w:ins')
    ins_el.set(qn('w:author'), author)
    ins_el.set(qn('w:date'), date_time.strftime("%Y-%m-%dT%H:%M:%SZ"))
    return ins_el

def create_run_element_with_text(text_content, template_run_r=None, is_del_text=False, highlight_color: Optional[str] = None):
    new_r = OxmlElement('w:r')
    rPr_element = None
    if template_run_r is not None:
        rPr_template = template_run_r.find(qn('w:rPr'))
        if rPr_template is not None: rPr_element = copy.deepcopy(rPr_template)
    
    if rPr_element is None and highlight_color: rPr_element = OxmlElement('w:rPr')
    
    if highlight_color and rPr_element is not None:
        for highlight_node in rPr_element.findall(qn('w:highlight')): rPr_element.remove(highlight_node)
        highlight_el = OxmlElement('w:highlight'); highlight_el.set(qn('w:val'), highlight_color)
        rPr_element.append(highlight_el)
    
    if rPr_element is not None and len(rPr_element) > 0: new_r.append(rPr_element)

    text_el = OxmlElement('w:delText' if is_del_text else 'w:t')
    text_el.set(qn('xml:space'), 'preserve'); text_el.text = text_content
    new_r.append(text_el)
    return new_r

def log_debug(message):
    if DEBUG_MODE: print(f"DEBUG (word_processor): {message}")

def _get_element_style_template_run(element_info_item_el, fallback_style_run):
    if element_info_item_el is not None and element_info_item_el.tag == qn('w:r'): return element_info_item_el
    if element_info_item_el is not None:
        r_child = element_info_item_el.find(qn('w:r'))
        if r_child is not None: return r_child
    return fallback_style_run if fallback_style_run is not None else OxmlElement('w:r')

def _build_visible_text_map(paragraph) -> Tuple[str, List[Dict[str, Any]]]:
    elements_map = []
    current_text_offset = 0
    for p_child_idx_loop, p_child_element in enumerate(list(paragraph._p)): 
        text_contribution = ""; element_type = "other"
        if p_child_element.tag == qn("w:r"):
            element_type = "r"
            if p_child_element.find(qn('w:delText')) is None:
                for t_node in p_child_element.findall(qn('w:t')):
                    if t_node.text: text_contribution += t_node.text
                if p_child_element.find(qn('w:tab')) is not None: text_contribution += '\t'
        elif p_child_element.tag == qn("w:ins"):
            element_type = "ins"
            for r_in_ins in p_child_element.findall(qn('w:r')):
                if r_in_ins.find(qn('w:delText')) is None: 
                    for t_in_ins in r_in_ins.findall(qn('w:t')):
                        if t_in_ins.text: text_contribution += t_in_ins.text
                    if r_in_ins.find(qn('w:tab')) is not None: text_contribution += '\t'
        elif p_child_element.tag == qn("w:hyperlink"):
             element_type = "hyperlink"
             for r_in_hyperlink in p_child_element.findall(qn('w:r')):
                if r_in_hyperlink.find(qn('w:delText')) is None:
                    for t_in_hyperlink in r_in_hyperlink.findall(qn('w:t')):
                        if t_in_hyperlink.text: text_contribution += t_in_hyperlink.text
                    if r_in_hyperlink.find(qn('w:tab')) is not None: text_contribution += '\t'
        if text_contribution:
            elements_map.append({'el': p_child_element, 'text': text_contribution, 
                'p_child_idx': p_child_idx_loop, 'doc_start_offset': current_text_offset,
                'type': element_type,
                'original_author': p_child_element.get(qn('w:author')) if element_type == "ins" else None,
                'original_date': p_child_element.get(qn('w:date')) if element_type == "ins" else None})
            current_text_offset += len(text_contribution)
    return "".join(item['text'] for item in elements_map), elements_map

def _add_comment_to_paragraph(paragraph, current_para_idx: int, comment_text: str, author_name: str, 
                             ambiguous_or_failed_changes_log: List[Dict], 
                             edit_item_details_for_log: Optional[Dict] = None):
    if not ADD_COMMENTS_TO_CHANGES or not comment_text: return
    log_ctx = {"paragraph_index": current_para_idx, **(edit_item_details_for_log or {})}
    try:
        author_display = f"{author_name} (LLM)"
        name_parts = [w for w in author_display.replace("(", "").replace(")", "").split() if w]
        initials = (name_parts[0][0] + name_parts[1][0]).upper() if len(name_parts) >= 2 else (name_parts[0][:2].upper() if name_parts else "AI")
        paragraph.add_comment(comment_text, author=author_display, initials=initials)
        log_debug(f"P{current_para_idx+1}: Added comment: '{comment_text[:30]}...'.")
    except AttributeError as e_attr: 
        ambiguous_or_failed_changes_log.append({**log_ctx, "issue": f"Comment addition failed (AttributeError): {e_attr}. Comment: {comment_text}", "type": "CriticalWarning"})
        log_debug(f"P{current_para_idx+1}: CRITICAL_WARNING - Could not add comment ('{comment_text[:30]}...') due to AttributeError: {e_attr}.")
    except Exception as e_gen: 
        ambiguous_or_failed_changes_log.append({**log_ctx, "issue": f"Comment addition failed: {e_gen}. Comment: {comment_text}", "type": "Warning"})
        log_debug(f"P{current_para_idx+1}: WARNING - Could not add comment ('{comment_text[:30]}...'). Error: {e_gen}")

def _apply_markup_to_span(
        paragraph_obj, current_para_idx: int, global_start: int, global_end: int, text_to_markup: str,
        highlight_color: str, comment_text: str, author_name: str,
        initial_fallback_style_run: Optional[OxmlElement], add_comments_param: bool, # Renamed add_comments_flag
        ambiguous_or_failed_changes_log: List[Dict], edit_item_details: Dict # For logging
    ) -> bool:
    log_debug(f"P{current_para_idx+1}: Applying markup: Highlight '{text_to_markup}' ({highlight_color}) at {global_start}-{global_end}, Comment: '{comment_text[:30]}...'")
    current_visible_text, current_elements_map = _build_visible_text_map(paragraph_obj)
    if not (0 <= global_start < global_end <= len(current_visible_text)):
        log_debug(f"P{current_para_idx+1}: Invalid span {global_start}-{global_end} for markup in text of len {len(current_visible_text)}. Skipping.")
        return False
    text_actually_at_span = current_visible_text[global_start:global_end]
    if text_actually_at_span != text_to_markup:
        log_debug(f"P{current_para_idx+1}: Markup text mismatch. Expected '{text_to_markup}', found '{text_actually_at_span}'. Skipping.")
        return False

    involved_span_element_infos = []
    span_first_style_run = initial_fallback_style_run
    for item in current_elements_map:
        item_doc_end_offset = item['doc_start_offset'] + len(item['text'])
        if max(item['doc_start_offset'], global_start) < min(item_doc_end_offset, global_end):
            involved_span_element_infos.append(item)
            if span_first_style_run == initial_fallback_style_run or span_first_style_run is None :
                 current_el_style_run = _get_element_style_template_run(item['el'], initial_fallback_style_run)
                 if current_el_style_run is not None : span_first_style_run = current_el_style_run
    if not involved_span_element_infos: log_debug(f"P{current_para_idx+1}: Markup failed. No XML elements for span."); return False
    if span_first_style_run is None: span_first_style_run = OxmlElement('w:r')

    new_xml_sequence = []
    first_item_markup = involved_span_element_infos[0]
    prefix_len_markup = global_start - first_item_markup['doc_start_offset']
    prefix_text_markup = first_item_markup['text'][:prefix_len_markup] if prefix_len_markup > 0 else ""
    if prefix_text_markup:
        new_xml_sequence.append(create_run_element_with_text(prefix_text_markup, _get_element_style_template_run(first_item_markup['el'], span_first_style_run)))
    new_xml_sequence.append(create_run_element_with_text(text_to_markup, span_first_style_run, highlight_color=highlight_color))
    last_item_markup = involved_span_element_infos[-1]
    suffix_start_markup = global_end - last_item_markup['doc_start_offset']
    suffix_text_markup = last_item_markup['text'][suffix_start_markup:] if suffix_start_markup < len(last_item_markup['text']) else ""
    if suffix_text_markup:
        new_xml_sequence.append(create_run_element_with_text(suffix_text_markup, _get_element_style_template_run(last_item_markup['el'], span_first_style_run)))

    p_child_indices_to_remove_markup = sorted(list(set(item['p_child_idx'] for item in involved_span_element_infos)), reverse=True)
    if not p_child_indices_to_remove_markup: return False
    insertion_point_markup = p_child_indices_to_remove_markup[-1]
    for p_idx_remove in p_child_indices_to_remove_markup:
        try: paragraph_obj._p.remove(paragraph_obj._p[p_idx_remove])
        except (ValueError, IndexError) as e: log_debug(f"P{current_para_idx+1}: Error removing element for markup {p_idx_remove}: {e}"); return False
    for i, new_el in enumerate(new_xml_sequence): paragraph_obj._p.insert(insertion_point_markup + i, new_el)
    
    log_debug(f"P{current_para_idx+1}: Markup applied for '{text_to_markup}'.")
    if add_comments_param: # Use passed parameter
        _add_comment_to_paragraph(paragraph_obj, current_para_idx, comment_text, author_name, ambiguous_or_failed_changes_log, edit_item_details)
    return True

def replace_text_in_paragraph_with_tracked_change(
        current_para_idx: int, paragraph, contextual_old_text_llm, specific_old_text_llm, 
        specific_new_text, reason_for_change, author, case_sensitive_flag, 
        add_comments_main_param: bool, # Renamed to avoid conflict
        ambiguous_or_failed_changes_log) -> Tuple[str, Optional[List[Dict[str, Any]]]]:
    if EXTENDED_DEBUG_MODE:
        log_debug(f"P{current_para_idx+1}: Full LLM Context: '{contextual_old_text_llm}'")
        log_debug(f"P{current_para_idx+1}: Full LLM Specific Old: '{specific_old_text_llm}'")
    log_debug(f"P{current_para_idx+1}: Attempting change '{specific_old_text_llm}' to '{specific_new_text}' within LLM context '{contextual_old_text_llm[:30]}...'")
    
    visible_paragraph_text_original_case, elements_map = _build_visible_text_map(paragraph)
    edit_details_for_log = {"contextual_old_text": contextual_old_text_llm, "specific_old_text": specific_old_text_llm, "specific_new_text": specific_new_text, "llm_reason": reason_for_change}

    if not visible_paragraph_text_original_case and (contextual_old_text_llm or specific_old_text_llm):
        log_debug(f"P{current_para_idx+1}: Paragraph empty, edit provided. Skipping."); return "CONTEXT_NOT_FOUND", None
    search_text_in_doc = visible_paragraph_text_original_case if case_sensitive_flag else visible_paragraph_text_original_case.lower()
    search_context_from_llm_processed = contextual_old_text_llm if case_sensitive_flag else contextual_old_text_llm.lower()
    log_debug(f"P{current_para_idx+1}: Visible text (len {len(visible_paragraph_text_original_case)}): '{visible_paragraph_text_original_case[:50]}...'")

    occurrences_of_context = []
    try:
        for match in re.finditer(re.escape(search_context_from_llm_processed), search_text_in_doc): 
            occurrences_of_context.append({"start": match.start(), "end": match.end(), "text": match.group(0)}) # Use match.group(0) for actual matched text
    except re.error as e:
        ambiguous_or_failed_changes_log.append({"paragraph_index": current_para_idx, "issue": f"Regex error: {e}", **edit_details_for_log}); return "REGEX_ERROR_IN_CONTEXT_SEARCH", None
    if not occurrences_of_context: log_debug(f"P{current_para_idx+1}: LLM Context '{contextual_old_text_llm[:30]}...' not found."); return "CONTEXT_NOT_FOUND", None
    if len(occurrences_of_context) > 1:
        log_debug(f"P{current_para_idx+1}: LLM Context '{contextual_old_text_llm[:30]}...' AMBIGUOUS ({len(occurrences_of_context)} found).")
        return "CONTEXT_AMBIGUOUS", occurrences_of_context

    unique_context_match_info = occurrences_of_context[0]
    actual_context_found_in_doc_str = unique_context_match_info["text"]
    prefix_display = visible_paragraph_text_original_case[max(0, unique_context_match_info['start']-10) : unique_context_match_info['start']]
    suffix_display = visible_paragraph_text_original_case[unique_context_match_info['end'] : min(len(visible_paragraph_text_original_case), unique_context_match_info['end']+10)]
    log_debug(f"P{current_para_idx+1}: Unique LLM context found: '...{prefix_display}[{actual_context_found_in_doc_str}]{suffix_display}...'")

    text_to_search_specific_within = actual_context_found_in_doc_str if case_sensitive_flag else actual_context_found_in_doc_str.lower()
    specific_text_to_find_llm_processed = specific_old_text_llm if case_sensitive_flag else specific_old_text_llm.lower()
    try:
        specific_start_in_context = text_to_search_specific_within.index(specific_text_to_find_llm_processed)
    except ValueError:
        ambiguous_or_failed_changes_log.append({"paragraph_index": current_para_idx, "issue": "Specific text not found within unique context.", **edit_details_for_log}); return "SPECIFIC_TEXT_NOT_IN_CONTEXT", None
    
    actual_specific_old_text_to_delete = actual_context_found_in_doc_str[specific_start_in_context : specific_start_in_context + len(specific_old_text_llm)]
    global_specific_start_offset = unique_context_match_info['start'] + specific_start_in_context
    global_specific_end_offset = global_specific_start_offset + len(actual_specific_old_text_to_delete)
    log_debug(f"P{current_para_idx+1}: Specific text for replacement: '{actual_specific_old_text_to_delete}'. Global offsets: {global_specific_start_offset}-{global_specific_end_offset}")

    char_before_specific = visible_paragraph_text_original_case[global_specific_start_offset - 1] if global_specific_start_offset > 0 else None
    char_after_specific = visible_paragraph_text_original_case[global_specific_end_offset] if global_specific_end_offset < len(visible_paragraph_text_original_case) else None
    is_start_boundary_ok = (global_specific_start_offset == 0 or (char_before_specific is not None and char_before_specific.isspace()))
    is_end_boundary_ok = (global_specific_end_offset == len(visible_paragraph_text_original_case) or \
                          (char_after_specific is not None and (char_after_specific.isspace() or char_after_specific in ALLOWED_POST_BOUNDARY_PUNCTUATION)))
    if not (is_start_boundary_ok and is_end_boundary_ok):
        log_message = f"P{current_para_idx+1}: Specific text '{actual_specific_old_text_to_delete}' NOT validly bounded. Preceded: '{char_before_specific}', Followed: '{char_after_specific}'. Change skipped."
        log_debug(log_message)
        ambiguous_or_failed_changes_log.append({"paragraph_index": current_para_idx, "issue": f"Skipped: Not validly bounded (prev: '{char_before_specific}', next: '{char_after_specific}')", **edit_details_for_log})
        return "SKIPPED_INVALID_BOUNDARY", None

    involved_element_infos = []
    first_involved_r_element_for_style = None
    for item in elements_map:
        item_doc_end_offset = item['doc_start_offset'] + len(item['text'])
        if max(item['doc_start_offset'], global_specific_start_offset) < min(item_doc_end_offset, global_specific_end_offset):
            involved_element_infos.append(item)
            if first_involved_r_element_for_style is None: first_involved_r_element_for_style = _get_element_style_template_run(item['el'], None)
    if not involved_element_infos: 
        ambiguous_or_failed_changes_log.append({"paragraph_index": current_para_idx, "issue": "XML mapping failed.", **edit_details_for_log}); return "XML_MAPPING_FAILED", None
    if first_involved_r_element_for_style is None: first_involved_r_element_for_style = OxmlElement('w:r') 
    log_debug(f"P{current_para_idx+1}: Modifying {len(involved_element_infos)} XML elements for change (validly bounded).")
    
    new_xml_elements_for_paragraph = []
    first_item = involved_element_infos[0]
    prefix_len_in_first_item = global_specific_start_offset - first_item['doc_start_offset']
    prefix_text = first_item['text'][:prefix_len_in_first_item] if prefix_len_in_first_item > 0 else ""
    if prefix_text:
        style_run = _get_element_style_template_run(first_item['el'], first_involved_r_element_for_style)
        if first_item['type'] == 'ins': 
            orig_ins = create_ins_element(author=first_item['original_author'],date_time=None)
            if first_item['original_date']: orig_ins.set(qn('w:date'), first_item['original_date'])
            orig_ins.append(create_run_element_with_text(prefix_text, style_run)); new_xml_elements_for_paragraph.append(orig_ins)
        else: new_xml_elements_for_paragraph.append(create_run_element_with_text(prefix_text, style_run))
        log_debug(f"P{current_para_idx+1}: Added prefix '{prefix_text}' (type: {first_item['type']}).")

    change_time = datetime.datetime.now(datetime.timezone.utc)
    del_obj = create_del_element(author=author, date_time=change_time)
    del_obj.append(create_run_element_with_text(actual_specific_old_text_to_delete, first_involved_r_element_for_style, True))
    new_xml_elements_for_paragraph.append(del_obj)
    log_debug(f"P{current_para_idx+1}: Added <w:del> for '{actual_specific_old_text_to_delete}'.")

    if specific_new_text:
        ins_obj = create_ins_element(author=author, date_time=change_time + datetime.timedelta(seconds=1))
        ins_obj.append(create_run_element_with_text(specific_new_text, first_involved_r_element_for_style, False, None))
        new_xml_elements_for_paragraph.append(ins_obj)
        log_debug(f"P{current_para_idx+1}: Added <w:ins> for '{specific_new_text}'.")

    last_item = involved_element_infos[-1]
    suffix_start_in_last_item = global_specific_end_offset - last_item['doc_start_offset']
    suffix_text = last_item['text'][suffix_start_in_last_item:] if suffix_start_in_last_item < len(last_item['text']) else ""
    if suffix_text:
        style_run = _get_element_style_template_run(last_item['el'], first_involved_r_element_for_style)
        if last_item['type'] == 'ins': 
            orig_ins = create_ins_element(author=last_item['original_author'],date_time=None)
            if last_item['original_date']: orig_ins.set(qn('w:date'), last_item['original_date'])
            orig_ins.append(create_run_element_with_text(suffix_text, style_run)); new_xml_elements_for_paragraph.append(orig_ins)
        else: new_xml_elements_for_paragraph.append(create_run_element_with_text(suffix_text, style_run))
        log_debug(f"P{current_para_idx+1}: Added suffix '{suffix_text}' (type: {last_item['type']}).")

    p_child_indices_to_remove = sorted(list(set(item['p_child_idx'] for item in involved_element_infos)), reverse=True)
    if not p_child_indices_to_remove: 
        ambiguous_or_failed_changes_log.append({"paragraph_index": current_para_idx, "issue": "XML_REMOVAL_ERROR_NO_INDICES", **edit_details_for_log}); return "XML_REMOVAL_ERROR_NO_INDICES", None
    insertion_point_xml = p_child_indices_to_remove[-1] 
    for p_idx_remove in p_child_indices_to_remove:
        try: paragraph._p.remove(paragraph._p[p_idx_remove])
        except (IndexError, ValueError) as e: 
            ambiguous_or_failed_changes_log.append({"paragraph_index": current_para_idx, "issue": f"XML element removal error: {e}", **edit_details_for_log}); return "XML_REMOVAL_ERROR", None
    log_debug(f"P{current_para_idx+1}: Removed {len(p_child_indices_to_remove)} original XML elements.")
    for i,new_el in enumerate(new_xml_elements_for_paragraph): paragraph._p.insert(insertion_point_xml + i, new_el)
    log_debug(f"P{current_para_idx+1}: Inserted {len(new_xml_elements_for_paragraph)} new XML elements at index {insertion_point_xml}.")
    
    comment = reason_for_change if specific_new_text else f"Deleted: '{actual_specific_old_text_to_delete}'. Reason: {reason_for_change}"
    _add_comment_to_paragraph(paragraph, current_para_idx, comment, author, ambiguous_or_failed_changes_log, edit_details_for_log)
    return "SUCCESS", None

def process_document_with_edits(
    input_docx_path: str, output_docx_path: str, edits_to_make: List[Dict],
    author_name: str = DEFAULT_AUTHOR_NAME, # DEFAULT_AUTHOR_NAME should be defined in this file
    debug_mode_flag: bool = False,
    extended_debug_mode_flag: bool = False, 
    case_sensitive_flag: bool = True, 
    add_comments_param: bool = True # <--- Ensure this parameter name is used
) -> Tuple[bool, Optional[str], List[Dict], int]: # Added int for processed_edits_count
    error_log_file_path: Optional[str] = None 
    global DEBUG_MODE, EXTENDED_DEBUG_MODE, CASE_SENSITIVE_SEARCH, ADD_COMMENTS_TO_CHANGES
    DEBUG_MODE = debug_mode_flag
    EXTENDED_DEBUG_MODE = extended_debug_mode_flag # Set global flag
    CASE_SENSITIVE_SEARCH = case_sensitive_flag
    ADD_COMMENTS_TO_CHANGES = add_comments_param # Assign to global used by helper functions

    log_debug(f"Script starting. Input: {input_docx_path}, Output: {output_docx_path}")
    log_debug(f"Settings - Debug:{DEBUG_MODE}, ExtDebug:{EXTENDED_DEBUG_MODE}, CaseSensitive:{CASE_SENSITIVE_SEARCH}, AddComments:{ADD_COMMENTS_TO_CHANGES}, Author:{author_name}")
    log_debug(f"Number of edits to attempt: {len(edits_to_make)}")

    if not isinstance(edits_to_make, list) or not all(isinstance(item,dict) for item in edits_to_make):
        return False, error_log_file_path, [{"issue": "FATAL: Edits must be list of dicts."}], 0
    ambiguous_or_failed_changes_log: List[Dict] = []
    try: doc = Document(input_docx_path); log_debug(f"Opened '{input_docx_path}'")
    except Exception as e: return False, error_log_file_path, [{"issue": f"FATAL: Error opening Doc: {e}"}], 0

    edits_successfully_applied = 0
    for para_idx, paragraph_obj in enumerate(doc.paragraphs):
        paragraph_modified_in_pass = False 
        log_debug(f"\n--- Processing P{para_idx+1} ('{paragraph_obj.text[:50]}...') ---")
        for edit_item_idx, edit_item in enumerate(list(edits_to_make)): 
            if paragraph_modified_in_pass: break 
            log_debug(f"Attempting edit {edit_item_idx+1} ('{edit_item.get('specific_old_text')}' -> '{edit_item.get('specific_new_text')}') in P{para_idx+1}")
            required = ["contextual_old_text", "specific_old_text", "reason_for_change"]
            if not all(k in edit_item for k in required):
                ambiguous_or_failed_changes_log.append({"paragraph_index":para_idx, "edit_item_index":edit_item_idx+1, "issue":"Invalid structure.", "edit_item_snippet":str(edit_item)[:100]})
                continue
            status, data = "INIT", None
            try:
                status, data = replace_text_in_paragraph_with_tracked_change(
                    para_idx, paragraph_obj, edit_item["contextual_old_text"], edit_item["specific_old_text"],
                    edit_item.get("specific_new_text",""), edit_item["reason_for_change"],
                    author_name, CASE_SENSITIVE_SEARCH, ADD_COMMENTS_TO_CHANGES, ambiguous_or_failed_changes_log)
            except Exception as e_rep: 
                status = "EXCEPTION_IN_REPLACE" 
                log_debug(f"P{para_idx+1}, Edit{edit_item_idx+1}: EXCEPTION for '{edit_item['specific_old_text']}'. Err: {type(e_rep).__name__}: {e_rep}")
                if DEBUG_MODE: log_debug(traceback.format_exc()) 
                ambiguous_or_failed_changes_log.append({"paragraph_index":para_idx, "edit_item_index":edit_item_idx+1, "issue":f"EXCEPTION in replace: {e_rep}", "type":"CriticalError", **edit_item})
            
            log_entry_base = {"contextual_old_text": edit_item["contextual_old_text"], "specific_old_text": edit_item["specific_old_text"], "llm_reason": edit_item["reason_for_change"]}

            if status == "SUCCESS":
                print(f"SUCCESS: P{para_idx+1}: Applied change for '{edit_item['specific_old_text']}'.")
                edits_successfully_applied += 1; paragraph_modified_in_pass = True
            elif status == "CONTEXT_AMBIGUOUS":
                log_debug(f"P{para_idx+1}: Context '{edit_item['contextual_old_text'][:30]}...' AMBIGUOUS for '{edit_item['specific_old_text']}'.")
                if data: # data is list of context occurrences
                    # Build initial map once for this paragraph before any markups for THIS edit item
                    # _initial_visible_text_amb, _initial_elements_map_amb = _build_visible_text_map(paragraph_obj) # This is risky if prev edit_item changed para
                    
                    # For styling, try to get a representative run from the current paragraph state
                    current_vis_text, current_elem_map = _build_visible_text_map(paragraph_obj)
                    style_run_amb = OxmlElement('w:r') # Default
                    if current_elem_map:
                        first_r_from_map = next((item['el'] for item in current_elem_map if item['el'].tag == qn('w:r')), None)
                        if first_r_from_map: style_run_amb = first_r_from_map
                    
                    spans_to_markup_this_edit_item = []
                    for ctx_occurrence in data: # ctx_occurrence is {"start":, "end":, "text":} for contextual_old_text
                        ctx_text = ctx_occurrence["text"]
                        ctx_start_global = ctx_occurrence["start"]
                        s_old_llm = edit_item["specific_old_text"]
                        s_old_search = s_old_llm if CASE_SENSITIVE_SEARCH else s_old_llm.lower()
                        ctx_text_search = ctx_text if CASE_SENSITIVE_SEARCH else ctx_text.lower()
                        
                        offset_in_ctx = 0
                        while offset_in_ctx < len(ctx_text_search):
                            try:
                                found_idx = ctx_text_search.index(s_old_search, offset_in_ctx)
                                specific_text_val = ctx_text[found_idx : found_idx + len(s_old_llm)]
                                specific_start_abs = ctx_start_global + found_idx
                                specific_end_abs = specific_start_abs + len(specific_text_val)
                                spans_to_markup_this_edit_item.append({"start":specific_start_abs, "end":specific_end_abs, "text":specific_text_val})
                                offset_in_ctx = found_idx + len(specific_text_val)
                            except ValueError: break
                    
                    spans_to_markup_this_edit_item.sort(key=lambda x:x["start"], reverse=True) # Process last first

                    for span in spans_to_markup_this_edit_item:
                        if _apply_markup_to_span(paragraph_obj, para_idx, span["start"], span["end"], span["text"], 
                                               HIGHLIGHT_COLOR_AMBIGUOUS_SKIPPED, 
                                               f"Skipped: Ambiguous context for '{s_old_llm}'. LLM Reason: {edit_item['reason_for_change']}", 
                                               author_name, style_run_amb, ADD_COMMENTS_TO_CHANGES, 
                                               ambiguous_or_failed_changes_log, log_entry_base):
                            paragraph_modified_in_pass = True # If any markup applied
                else: # data (occurrences) was None for some reason
                    ambiguous_or_failed_changes_log.append({"paragraph_index":para_idx, "issue": "CONTEXT_AMBIGUOUS but no occurrence data returned.", "type": "Warning", **log_entry_base})
            
            elif status != "INIT" and status not in ["CONTEXT_NOT_FOUND", "REGEX_ERROR_IN_CONTEXT_SEARCH", "SKIPPED_INVALID_BOUNDARY", "SPECIFIC_TEXT_NOT_IN_CONTEXT", "XML_MAPPING_FAILED"]:
                log_debug(f"P{para_idx+1}: Edit for '{edit_item['specific_old_text']}' status: {status}.")
            
            if status in ["EXCEPTION_IN_REPLACE", "XML_REMOVAL_ERROR", "XML_REMOVAL_ERROR_NO_INDICES"]:
                log_debug(f"P{para_idx+1}: Critical error {status}. Halting edits for this paragraph.")
                paragraph_modified_in_pass = True # Treat as processed to stop further attempts

    try: doc.save(output_docx_path); print(f"\nSaved: '{output_docx_path}'")
    except Exception as e:
        ambiguous_or_failed_changes_log.append({"issue":f"FATAL: Save Error: {e}", "type":"FatalError"})
        return False, error_log_file_path, ambiguous_or_failed_changes_log, edits_successfully_applied

    if ambiguous_or_failed_changes_log:
        # ... (log file writing logic remains similar, ensure error_log_file_path is defined)
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        log_filename_with_ts = f"{ERROR_LOG_FILENAME_BASE}_{timestamp}.txt"
        try:
            output_dir = os.path.dirname(os.path.abspath(output_docx_path))
            if output_dir and not os.path.exists(output_dir): os.makedirs(output_dir, exist_ok=True)
            error_log_file_path = os.path.join(output_dir, log_filename_with_ts) if output_dir else log_filename_with_ts
            with open(error_log_file_path, "w", encoding="utf-8") as f:
                # ... (write log content) ...
                f.write(f"Total Edits: {len(edits_to_make)}, Applied: {edits_successfully_applied}\n")
                f.write(f"Log Items: {len(ambiguous_or_failed_changes_log)}\n\n")
                for entry in ambiguous_or_failed_changes_log: f.write(f"{json.dumps(entry)}\n")
            print(f"Log file with {len(ambiguous_or_failed_changes_log)} items: '{error_log_file_path}'")
        except Exception as e_log:
            log_debug(f"ERROR writing log file '{error_log_file_path}': {e_log}")
            error_log_file_path = None # Indicate failure
    else:
        # ... (summary messages based on edits_successfully_applied vs len(edits_to_make)) ...
        pass # Placeholder for refined summary messages
    return True, error_log_file_path, ambiguous_or_failed_changes_log, edits_successfully_applied

if __name__ == '__main__':
    # ... (argparse setup as before, including extended_debug) ...
    # Pass args.extended_debug to process_document_with_edits
    # ...
    parser = argparse.ArgumentParser(description="Apply tracked changes to a Word document.")
    parser.add_argument("--input", default=DEFAULT_INPUT_DOCX_PATH, help=f"Input DOCX (default: {DEFAULT_INPUT_DOCX_PATH})")
    parser.add_argument("--output", default=DEFAULT_OUTPUT_DOCX_PATH, help=f"Output DOCX (default: {DEFAULT_OUTPUT_DOCX_PATH})")
    parser.add_argument("--editsfile", default=DEFAULT_EDITS_FILE_PATH, help=f"JSON edits file (default: {DEFAULT_EDITS_FILE_PATH})")
    parser.add_argument("--editsjson", type=str, help="JSON string of edits.")
    parser.add_argument("--author", default=DEFAULT_AUTHOR_NAME, help=f"Author for changes (default: {DEFAULT_AUTHOR_NAME}).")
    parser.add_argument("--debug", action="store_true", help="Enable debug mode.")
    parser.add_argument("--extended_debug", action="store_true", help="Enable extended debug mode for very verbose context logging.") 
    parser.add_argument("--caseinsensitive", action="store_false", dest="case_sensitive", help="Case-insensitive search.")
    parser.set_defaults(case_sensitive=True)
    parser.add_argument("--nocomments", action="store_false", dest="add_comments", help="Disable comments.")
    parser.set_defaults(add_comments=True)
    
    args = parser.parse_args()
    if args.debug: DEBUG_MODE = True
    if args.extended_debug: EXTENDED_DEBUG_MODE = True 

    edits_data = []
    # ... (edits data loading unchanged from previous version, including dummy data if no file/json) ...
    if args.editsjson:
        try: edits_data = json.loads(args.editsjson); log_debug(f"Loaded edits from --editsjson argument.")
        except json.JSONDecodeError as e: print(f"FATAL: Error decoding JSON from --editsjson: {e}. Exiting."); exit(1)
    elif args.editsfile:
        try:
            with open(args.editsfile, 'r', encoding='utf-8') as f: edits_data = json.load(f)
            log_debug(f"Successfully loaded {len(edits_data)} edits from '{args.editsfile}'.")
        except FileNotFoundError: print(f"FATAL: Edits file '{args.editsfile}' not found. Exiting."); exit(1)
        except json.JSONDecodeError as e: print(f"FATAL: Error decoding JSON from '{args.editsfile}': {e}. Exiting."); exit(1)
        except Exception as e: print(f"FATAL: An unexpected error loading '{args.editsfile}': {e}. Exiting."); exit(1)
    else:
        if not DEBUG_MODE: DEBUG_MODE = True 
        log_debug("No edits provided. Using dummy edits for testing if input/output are defaults.")
        if args.input == DEFAULT_INPUT_DOCX_PATH and args.output == DEFAULT_OUTPUT_DOCX_PATH : # Ensure a specific output name for this test
            print("INFO: No edits specified. Using dummy edits for testing boundary conditions and ambiguity.")
            edits_data = [ # Dummy edits updated to test new boundary rules and ambiguity
                {"contextual_old_text": "cost would be $101 , to a new number", "specific_old_text": "$101", "specific_new_text": "$777","reason_for_change": "Test '$101' (boundary: space before, comma after -> Should be OK)"},
                {"contextual_old_text": "MrArbor, but that name", "specific_old_text": "MrArbor", "specific_new_text": "DrArbor","reason_for_change": "Test 'MrArbor' (boundary: space before, comma after -> Should be OK)"},
                {"contextual_old_text": "name can change", "specific_old_text": "name", "specific_new_text": "label","reason_for_change": "Test 'name' (space bounded -> SUCCESS)"},
                {"contextual_old_text": "with 9 longs.", "specific_old_text": "longs", "specific_new_text": "iterations","reason_for_change": "Test 'longs' (boundary: space before, period after -> Should be OK)"},
                {"contextual_old_text": "A simple file. It should", "specific_old_text": "file", "specific_new_text": "document","reason_for_change": "Test 'file' (boundary: space before, period after -> Should be OK)"},
                {"contextual_old_text": "Bob was also", "specific_old_text": "Bob", "specific_new_text": "Robert","reason_for_change": "Test one Bob. If this context is unique, it changes. If ambiguous, it should be orange marked."},
                {"contextual_old_text": "Bob", "specific_old_text": "Bob", "specific_new_text": "Robert","reason_for_change": "Test ambiguous Bob for orange marking."} # This one likely ambiguous
            ]
            if not os.path.exists(DEFAULT_INPUT_DOCX_PATH):
                try:
                    doc_dummy = Document()
                    doc_dummy.add_paragraph("A simple file. It should change this sentence from saying the cost would be $101 , to a new number.")
                    doc_dummy.add_paragraph("First Bob. Bob was also here. And Bob was there. This Bob is specific.")
                    doc_dummy.add_paragraph("Here is a long long long long long long long long long long repetitive line with 9 longs.")
                    doc_dummy.add_paragraph("The last line was last edited by MrArbor, but that name can change.")
                    doc_dummy.save(DEFAULT_INPUT_DOCX_PATH)
                    print(f"Created dummy input: {DEFAULT_INPUT_DOCX_PATH}")
                except Exception as e_doc: print(f"Could not create dummy input: {e_doc}"); exit(1)
        else: print("FATAL: No edits provided and not using default paths for dummy data. Use --editsfile or --editsjson. Exiting."); exit(1)

    if not os.path.exists(args.input): print(f"FATAL: Input file '{args.input}' not found. Exiting."); exit(1)

    process_document_with_edits(
        args.input, args.output, edits_data,
        args.author, args.debug, args.extended_debug, 
        args.case_sensitive, args.add_comments
    )