#!/usr/bin/env python3
import copy
import datetime
import re
import json
import os
# import traceback # For more detailed debugging if needed
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from typing import List, Dict, Tuple, Optional

# --- Global Configuration Flags ---
DEBUG_MODE = False
CASE_SENSITIVE_SEARCH = True
ADD_COMMENTS_TO_CHANGES = True
DEFAULT_AUTHOR_NAME = "LLM Editor"

# --- Constants ---
ERROR_LOG_FILENAME_BASE = "change_log"
HIGHLIGHT_COLOR_FOR_NON_WHITESPACE_BOUNDED_CHANGES = "yellow"

# --- XML Helper Functions ---
def create_del_element(author="Python Program", date_time=None):
    if date_time is None:
        date_time = datetime.datetime.now(datetime.timezone.utc)
    del_el = OxmlElement('w:del')
    del_el.set(qn('w:author'), author)
    del_el.set(qn('w:date'), date_time.strftime("%Y-%m-%dT%H:%M:%SZ"))
    return del_el

def create_ins_element(author="Python Program", date_time=None):
    if date_time is None:
        date_time = datetime.datetime.now(datetime.timezone.utc)
    ins_el = OxmlElement('w:ins')
    ins_el.set(qn('w:author'), author)
    ins_el.set(qn('w:date'), date_time.strftime("%Y-%m-%dT%H:%M:%SZ"))
    return ins_el

def create_run_element_with_text(text_content, template_run_r=None, is_del_text=False, highlight_color: Optional[str] = None):
    new_r = OxmlElement('w:r')
    rPr_element = None
    
    if template_run_r is not None:
        rPr_template = template_run_r.find(qn('w:rPr'))
        if rPr_template is not None:
            rPr_element = copy.deepcopy(rPr_template)
    
    if rPr_element is None and highlight_color:
        rPr_element = OxmlElement('w:rPr')
    
    if highlight_color and rPr_element is not None:
        highlight_el = OxmlElement('w:highlight')
        highlight_el.set(qn('w:val'), highlight_color)
        rPr_element.append(highlight_el)
    
    if rPr_element is not None and len(rPr_element) > 0: # Corrected check
        new_r.append(rPr_element)

    text_element_tag = 'w:delText' if is_del_text else 'w:t'
    text_el = OxmlElement(text_element_tag)
    text_el.set(qn('xml:space'), 'preserve')
    text_el.text = text_content
    new_r.append(text_el)
    return new_r

def log_debug(message):
    if DEBUG_MODE:
        print(f"DEBUG (word_processor): {message}")

def _get_element_style_template_run(element_info_item, fallback_style_run):
    if element_info_item['el'].tag == qn('w:r'):
        return element_info_item['el']
    r_child = element_info_item['el'].find(qn('w:r'))
    if r_child is not None:
        return r_child
    return fallback_style_run

def replace_text_in_paragraph_with_tracked_change(
        paragraph_idx, paragraph,
        contextual_old_text, specific_old_text, specific_new_text, reason_for_change,
        author, case_sensitive_flag, add_comments_flag, ambiguous_or_failed_changes_log):
    log_debug(f"P{paragraph_idx+1}: Attempting to change '{specific_old_text}' to '{specific_new_text}' within context '{contextual_old_text[:30]}...'")

    elements_contributing_to_visible_text = []
    current_text_offset = 0
    for p_child_idx_loop, p_child_element in enumerate(list(paragraph._p)): 
        text_contribution = ""
        element_type = "other"
        if p_child_element.tag == qn("w:r"):
            element_type = "r"
            if p_child_element.find(qn('w:delText')) is None:
                for t_node in p_child_element.findall(qn('w:t')):
                    if t_node.text: text_contribution += t_node.text
                if p_child_element.find(qn('w:tab')) is not None: text_contribution += '\t'
        elif p_child_element.tag == qn("w:ins"):
            element_type = "ins"
            for r_in_ins in p_child_element.findall(qn('w:r')):
                if r_in_ins.find(qn('w:delText')) is None: 
                    for t_in_ins in r_in_ins.findall(qn('w:t')):
                        if t_in_ins.text: text_contribution += t_in_ins.text
                    if r_in_ins.find(qn('w:tab')) is not None: text_contribution += '\t'
        elif p_child_element.tag == qn("w:hyperlink"):
             element_type = "hyperlink"
             for r_in_hyperlink in p_child_element.findall(qn('w:r')):
                if r_in_hyperlink.find(qn('w:delText')) is None:
                    for t_in_hyperlink in r_in_hyperlink.findall(qn('w:t')):
                        if t_in_hyperlink.text: text_contribution += t_in_hyperlink.text
                    if r_in_hyperlink.find(qn('w:tab')) is not None: text_contribution += '\t'
        
        if text_contribution:
            elements_contributing_to_visible_text.append({
                'el': p_child_element, 'text': text_contribution, 
                'p_child_idx': p_child_idx_loop, 'doc_start_offset': current_text_offset,
                'type': element_type,
                'original_author': p_child_element.get(qn('w:author')) if element_type == "ins" else None,
                'original_date': p_child_element.get(qn('w:date')) if element_type == "ins" else None
            })
            current_text_offset += len(text_contribution)
    
    visible_paragraph_text_original_case = "".join(item['text'] for item in elements_contributing_to_visible_text)
    search_text_in_doc = visible_paragraph_text_original_case if case_sensitive_flag else visible_paragraph_text_original_case.lower()
    search_context_from_llm = contextual_old_text if case_sensitive_flag else contextual_old_text.lower()

    log_debug(f"P{paragraph_idx+1}: Visible text (len {len(visible_paragraph_text_original_case)}): '{visible_paragraph_text_original_case[:100]}{'...' if len(visible_paragraph_text_original_case)>100 else ''}'")

    occurrences = []
    try:
        for match in re.finditer(re.escape(search_context_from_llm), search_text_in_doc): occurrences.append(match)
    except re.error as e:
        log_message = f"P{paragraph_idx+1}: Regex error for context '{search_context_from_llm[:50]}...': {e}. LLM Reason: {reason_for_change}"
        log_debug(log_message)
        ambiguous_or_failed_changes_log.append({"paragraph_index": paragraph_idx, "visible_text_snippet": visible_paragraph_text_original_case[:100],"contextual_old_text": contextual_old_text, "specific_old_text": specific_old_text,"llm_reason": reason_for_change, "issue": f"Regex error: {e}"})
        return "REGEX_ERROR_IN_CONTEXT_SEARCH"

    if not occurrences: return "CONTEXT_NOT_FOUND"
    if len(occurrences) > 1:
        log_message = (f"P{paragraph_idx+1}: Context '{contextual_old_text[:50]}...' AMBIGUOUS ({len(occurrences)} found). Change '{specific_old_text}' to '{specific_new_text}' skipped. LLM Reason: {reason_for_change}")
        log_debug(log_message)
        ambiguous_or_failed_changes_log.append({"paragraph_index": paragraph_idx, "visible_text_snippet": visible_paragraph_text_original_case[:100],"contextual_old_text": contextual_old_text, "specific_old_text": specific_old_text,"specific_new_text": specific_new_text, "llm_reason": reason_for_change,"issue": f"Ambiguous context: Found {len(occurrences)} occurrences."})
        return "CONTEXT_AMBIGUOUS"

    match_context = occurrences[0]
    actual_context_found_in_doc = visible_paragraph_text_original_case[match_context.start() : match_context.end()]
    text_to_search_within_context = actual_context_found_in_doc if case_sensitive_flag else actual_context_found_in_doc.lower()
    specific_text_to_find_in_context = specific_old_text if case_sensitive_flag else specific_old_text.lower()

    try:
        specific_start_in_context = text_to_search_within_context.index(specific_text_to_find_in_context)
    except ValueError:
        log_message = (f"P{paragraph_idx+1}: Specific text '{specific_old_text}' NOT IN CONTEXT '{actual_context_found_in_doc}'. Skipped. LLM Reason: {reason_for_change}")
        log_debug(log_message)
        ambiguous_or_failed_changes_log.append({"paragraph_index": paragraph_idx, "visible_text_snippet": visible_paragraph_text_original_case[:100],"contextual_old_text": contextual_old_text, "specific_old_text": specific_old_text,"specific_new_text": specific_new_text, "llm_reason": reason_for_change,"issue": "Specific text not found within unique context." })
        return "SPECIFIC_TEXT_NOT_IN_CONTEXT"
    
    actual_specific_old_text_to_delete = actual_context_found_in_doc[specific_start_in_context : specific_start_in_context + len(specific_old_text)]
    global_specific_start_offset = match_context.start() + specific_start_in_context
    global_specific_end_offset = global_specific_start_offset + len(actual_specific_old_text_to_delete)
    
    log_debug(f"P{paragraph_idx+1}: Unique context. Specific text '{actual_specific_old_text_to_delete}' (len {len(actual_specific_old_text_to_delete)}) to be replaced with '{specific_new_text}'. Global offsets: {global_specific_start_offset}-{global_specific_end_offset}")

    involved_element_infos = []
    first_involved_r_element_for_style = None
    for item in elements_contributing_to_visible_text:
        item_doc_end_offset = item['doc_start_offset'] + len(item['text'])
        if max(item['doc_start_offset'], global_specific_start_offset) < min(item_doc_end_offset, global_specific_end_offset):
            involved_element_infos.append(item)
            if first_involved_r_element_for_style is None:
                first_involved_r_element_for_style = _get_element_style_template_run(item, None)
    
    if not involved_element_infos:
        log_message = f"P{paragraph_idx+1}: XML_MAPPING_FAILED for '{actual_specific_old_text_to_delete}'. LLM Reason: {reason_for_change}"
        log_debug(log_message)
        ambiguous_or_failed_changes_log.append({"paragraph_index": paragraph_idx, "issue": "XML mapping failed for specific text.", "contextual_old_text": contextual_old_text, "specific_old_text": actual_specific_old_text_to_delete, "llm_reason": reason_for_change,"visible_text_snippet": visible_paragraph_text_original_case[:100]})
        return "XML_MAPPING_FAILED"

    if first_involved_r_element_for_style is None: 
        first_involved_r_element_for_style = OxmlElement('w:r') 
        log_debug(f"P{paragraph_idx+1}: WARNING - No template <w:r> for styling. Using default.")

    log_debug(f"P{paragraph_idx+1}: Modifying {len(involved_element_infos)} raw XML elements for the change.")

    is_start_boundary_space = (global_specific_start_offset == 0 or \
                               (global_specific_start_offset > 0 and visible_paragraph_text_original_case[global_specific_start_offset - 1].isspace()))
    is_end_boundary_space = (global_specific_end_offset == len(visible_paragraph_text_original_case) or \
                             (global_specific_end_offset < len(visible_paragraph_text_original_case) and visible_paragraph_text_original_case[global_specific_end_offset].isspace()))
    highlight_new_text_flag = not (is_start_boundary_space and is_end_boundary_space)
    log_debug(f"P{paragraph_idx+1}: Highlight check: StartSpace:{is_start_boundary_space}, EndSpace:{is_end_boundary_space}. Highlight:{highlight_new_text_flag}")

    new_xml_elements_for_paragraph = []
    first_item = involved_element_infos[0]
    prefix_len_in_first_item = global_specific_start_offset - first_item['doc_start_offset']
    if prefix_len_in_first_item < 0: prefix_len_in_first_item = 0
    prefix_text = first_item['text'][:prefix_len_in_first_item]

    if prefix_text:
        style_run_for_prefix = _get_element_style_template_run(first_item, first_involved_r_element_for_style)
        if first_item['type'] == 'ins':
            orig_ins = create_ins_element(author=first_item['original_author'], date_time=None)
            if first_item['original_date']: orig_ins.set(qn('w:date'), first_item['original_date'])
            orig_ins.append(create_run_element_with_text(prefix_text, style_run_for_prefix))
            new_xml_elements_for_paragraph.append(orig_ins)
        else:
            new_xml_elements_for_paragraph.append(create_run_element_with_text(prefix_text, style_run_for_prefix))
        log_debug(f"P{paragraph_idx+1}: Added prefix '{prefix_text}' from first element (type: {first_item['type']}).")

    change_time = datetime.datetime.now(datetime.timezone.utc)
    del_obj = create_del_element(author=author, date_time=change_time)
    del_obj.append(create_run_element_with_text(actual_specific_old_text_to_delete, first_involved_r_element_for_style, is_del_text=True))
    new_xml_elements_for_paragraph.append(del_obj)
    log_debug(f"P{paragraph_idx+1}: Added <w:del> for '{actual_specific_old_text_to_delete}'.")

    if specific_new_text:
        ins_obj = create_ins_element(author=author, date_time=change_time + datetime.timedelta(seconds=1))
        ins_run_el = create_run_element_with_text(
            specific_new_text, 
            first_involved_r_element_for_style, 
            is_del_text=False,
            highlight_color=HIGHLIGHT_COLOR_FOR_NON_WHITESPACE_BOUNDED_CHANGES if highlight_new_text_flag else None
        )
        ins_obj.append(ins_run_el)
        new_xml_elements_for_paragraph.append(ins_obj)
        log_debug(f"P{paragraph_idx+1}: Added <w:ins> for '{specific_new_text}'" + (" (highlighted)" if highlight_new_text_flag else ""))

    last_item = involved_element_infos[-1]
    suffix_start_in_last_item = global_specific_end_offset - last_item['doc_start_offset']
    suffix_text = ""
    if suffix_start_in_last_item < len(last_item['text']):
        suffix_text = last_item['text'][suffix_start_in_last_item:]
    if suffix_text:
        style_run_for_suffix = _get_element_style_template_run(last_item, first_involved_r_element_for_style)
        if last_item['type'] == 'ins':
            orig_ins = create_ins_element(author=last_item['original_author'], date_time=None)
            if last_item['original_date']: orig_ins.set(qn('w:date'), last_item['original_date'])
            orig_ins.append(create_run_element_with_text(suffix_text, style_run_for_suffix))
            new_xml_elements_for_paragraph.append(orig_ins)
        else:
            new_xml_elements_for_paragraph.append(create_run_element_with_text(suffix_text, style_run_for_suffix))
        log_debug(f"P{paragraph_idx+1}: Added suffix '{suffix_text}' from last element (type: {last_item['type']}).")

    p_child_indices_to_remove = sorted(list(set(item['p_child_idx'] for item in involved_element_infos)), reverse=True)
    
    if not p_child_indices_to_remove:
        log_debug(f"P{paragraph_idx+1}: No paragraph child indices identified for removal. Aborting this change.")
        # This case should ideally be caught by `if not involved_element_infos:` earlier.
        ambiguous_or_failed_changes_log.append({"paragraph_index": paragraph_idx, "issue": "XML_REMOVAL_ERROR_NO_INDICES: No elements to remove.", "contextual_old_text": contextual_old_text, "specific_old_text": actual_specific_old_text_to_delete,"llm_reason": reason_for_change})
        return "XML_REMOVAL_ERROR_NO_INDICES"

    insertion_point_xml = p_child_indices_to_remove[-1] 

    for p_idx_to_remove_loop in p_child_indices_to_remove:
        try:
            # Ensure index is still valid if paragraph._p was modified by a previous removal in this loop
            # This check is tricky because p_idx_to_remove_loop are original indices.
            # By removing in reverse sorted order, the indices for subsequent removals remain valid relative to the current state of paragraph._p
            element_to_remove = paragraph._p[p_idx_to_remove_loop]
            paragraph._p.remove(element_to_remove) # CORRECTED: Use remove(element)
            log_debug(f"P{paragraph_idx+1}: Removed original XML element at original p_child_idx {p_idx_to_remove_loop}.")
        except (IndexError, ValueError) as e_remove: # ValueError if element not found (should not happen if index is correct)
            log_message = f"P{paragraph_idx+1}: XML element removal error at original index {p_idx_to_remove_loop}. Error: {e_remove}. Change aborted. LLM Reason: {reason_for_change}"
            log_debug(log_message)
            ambiguous_or_failed_changes_log.append({"paragraph_index": paragraph_idx, "issue": f"XML element removal error at index {p_idx_to_remove_loop}: {e_remove}","contextual_old_text": contextual_old_text, "specific_old_text": actual_specific_old_text_to_delete,"llm_reason": reason_for_change})
            # This is a critical failure for this change. The paragraph might be in an inconsistent state.
            # It might be better to re-raise or return a very specific error to stop further processing on this paragraph.
            return "XML_REMOVAL_ERROR" 
            
    for i, new_el in enumerate(new_xml_elements_for_paragraph):
        paragraph._p.insert(insertion_point_xml + i, new_el)
    log_debug(f"P{paragraph_idx+1}: Inserted {len(new_xml_elements_for_paragraph)} new XML elements at original index {insertion_point_xml}.")

    if add_comments_flag and reason_for_change:
        comment_text_to_add = reason_for_change
        if not specific_new_text:
            comment_text_to_add = f"Deleted: '{actual_specific_old_text_to_delete}'. Reason: {reason_for_change}"
        try:
            comment_author_name = f"{author} (LLM)"
            name_parts = [word for word in comment_author_name.replace("(", "").replace(")", "").split() if word]
            comment_initials = (name_parts[0][0] + name_parts[1][0]).upper() if len(name_parts) >= 2 else (name_parts[0][:2].upper() if name_parts else "LS")
            paragraph.add_comment(comment_text_to_add, author=comment_author_name, initials=comment_initials)
            log_debug(f"P{paragraph_idx+1}: Added comment: '{comment_text_to_add[:30]}...'.")
        except AttributeError as e_attr: 
            log_debug(f"P{paragraph_idx+1}: CRITICAL_WARNING - Could not add comment due to AttributeError: {e_attr}. Paragraph object might be in an unstable state after XML manipulation.")
            ambiguous_or_failed_changes_log.append({"paragraph_index": paragraph_idx, "issue": f"Comment addition failed (AttributeError): {e_attr}. Paragraph object may be unstable.","contextual_old_text": contextual_old_text, "specific_old_text": actual_specific_old_text_to_delete,"specific_new_text": specific_new_text, "llm_reason": reason_for_change, "type": "CriticalWarning"})
        except Exception as e_gen: 
            log_debug(f"P{paragraph_idx+1}: WARNING - Could not add comment. Error: {e_gen}")
            ambiguous_or_failed_changes_log.append({"paragraph_index": paragraph_idx, "issue": f"Comment addition failed: {e_gen}","contextual_old_text": contextual_old_text, "specific_old_text": actual_specific_old_text_to_delete,"specific_new_text": specific_new_text, "llm_reason": reason_for_change, "type": "Warning"})
    return "SUCCESS"

def process_document_with_edits(
    input_docx_path: str, output_docx_path: str, edits_to_make: List[Dict],
    author_name: str = DEFAULT_AUTHOR_NAME, debug_mode_flag: bool = False,
    case_sensitive_flag: bool = True, add_comments_flag: bool = True
) -> Tuple[bool, Optional[str], List[Dict]]:
    global DEBUG_MODE, CASE_SENSITIVE_SEARCH, ADD_COMMENTS_TO_CHANGES
    DEBUG_MODE = debug_mode_flag
    CASE_SENSITIVE_SEARCH = case_sensitive_flag
    ADD_COMMENTS_TO_CHANGES = add_comments_flag

    log_debug(f"Script starting. Input: {input_docx_path}, Output: {output_docx_path}")
    log_debug(f"Debug:{DEBUG_MODE}, CaseSensitive:{CASE_SENSITIVE_SEARCH}, AddComments:{ADD_COMMENTS_TO_CHANGES}, Author:{author_name}")
    log_debug(f"Number of edits to attempt: {len(edits_to_make)}")

    if not isinstance(edits_to_make, list) or not all(isinstance(item, dict) for item in edits_to_make):
        log_debug("FATAL: Edits must be a list of dictionaries.")
        return False, None, [{"issue": "FATAL: Edits must be a list of dictionaries."}]

    ambiguous_or_failed_changes_log = []
    try:
        doc = Document(input_docx_path)
        log_debug(f"Successfully opened '{input_docx_path}'")
    except Exception as e:
        log_debug(f"FATAL: Error opening '{input_docx_path}': {e}")
        return False, None, [{"issue": f"FATAL: Error opening Word document '{input_docx_path}': {e}"}]

    edits_processed_in_current_run = 0
    
    for para_idx, paragraph_obj in enumerate(doc.paragraphs):
        current_paragraph_text_snapshot = paragraph_obj.text 
        log_debug(f"\n--- Processing Paragraph {para_idx+1} (Initial Text Snapshot: '{current_paragraph_text_snapshot[:50]}...') ---")

        for edit_item_idx, edit_item in enumerate(list(edits_to_make)): 
            log_debug(f"Attempting edit item {edit_item_idx+1} ('{edit_item.get('specific_old_text')}' -> '{edit_item.get('specific_new_text')}') in P{para_idx+1}")
            
            required_keys = ["contextual_old_text", "specific_old_text", "reason_for_change"]
            if not all(key in edit_item for key in required_keys):
                log_message = f"P{para_idx+1}, EditItem{edit_item_idx+1}: Invalid structure (missing keys): {edit_item}. Skipping."
                log_debug(log_message)
                ambiguous_or_failed_changes_log.append({"paragraph_index": para_idx, "edit_item_index": edit_item_idx +1, "issue": "Invalid edit item structure.", "edit_item_snippet": str(edit_item)[:100]})
                continue

            specific_new_text = edit_item.get("specific_new_text", "")
            status = "INIT" 
            try:
                status = replace_text_in_paragraph_with_tracked_change(
                    para_idx, paragraph_obj, 
                    edit_item["contextual_old_text"], edit_item["specific_old_text"],
                    specific_new_text, edit_item["reason_for_change"],
                    author_name, CASE_SENSITIVE_SEARCH, ADD_COMMENTS_TO_CHANGES,
                    ambiguous_or_failed_changes_log
                )
            except Exception as e_replace: 
                status = "UNHANDLED_EXCEPTION_IN_REPLACE_CALL" # Differentiate from internal catch
                log_message = f"P{para_idx+1}, EditItem{edit_item_idx+1}: Unhandled exception during *call* to replacement function for '{edit_item['specific_old_text']}'. Error: {e_replace}"
                log_debug(log_message)
                # log_debug(traceback.format_exc()) 
                ambiguous_or_failed_changes_log.append({
                    "paragraph_index": para_idx, 
                    "edit_item_index": edit_item_idx +1,
                    "issue": f"Unhandled exception in replacement call: {e_replace}", 
                    "contextual_old_text": edit_item["contextual_old_text"],
                    "specific_old_text": edit_item["specific_old_text"],
                    "llm_reason": edit_item["reason_for_change"],
                    "type": "CriticalError"
                })

            if status == "SUCCESS":
                success_msg = f"SUCCESS: P{para_idx+1}: Applied change for context '{edit_item['contextual_old_text'][:30]}...'. Reason: {edit_item['reason_for_change']}"
                print(success_msg) 
                log_debug(success_msg)
                edits_processed_in_current_run += 1
                break 
            elif status not in ["CONTEXT_NOT_FOUND", "REGEX_ERROR_IN_CONTEXT_SEARCH", "INIT"]: 
                info_msg = f"INFO: P{para_idx+1}: Edit for context '{edit_item['contextual_old_text'][:30]}...' status: {status}."
                print(info_msg)
                log_debug(info_msg)
                if status in ["XML_REMOVAL_ERROR", "XML_MAPPING_FAILED", 
                              "UNHANDLED_EXCEPTION_IN_REPLACE", "UNHANDLED_EXCEPTION_IN_REPLACE_CALL",
                              "XML_REMOVAL_ERROR_NO_INDICES"]:
                    log_debug(f"P{para_idx+1}: Breaking further edits for this paragraph due to critical error status: {status}")
                    break

    try:
        doc.save(output_docx_path)
        print(f"\nProcessed document saved to '{output_docx_path}'")
        log_debug(f"Processed document saved to '{output_docx_path}'")
    except Exception as e:
        log_debug(f"FATAL: Error saving document to '{output_docx_path}': {e}")
        ambiguous_or_failed_changes_log.append({"issue": f"FATAL: Error saving document to '{output_docx_path}': {e}"})
        return False, None, ambiguous_or_failed_changes_log

    error_log_file_path = None
    if ambiguous_or_failed_changes_log:
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        log_filename_with_ts = f"{ERROR_LOG_FILENAME_BASE}_{timestamp}.txt"
        output_dir = os.path.dirname(output_docx_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        error_log_file_path = os.path.join(output_dir, log_filename_with_ts) if output_dir else log_filename_with_ts
        try:
            with open(error_log_file_path, "w", encoding="utf-8") as f:
                f.write(f"--- LOG OF CHANGES NOT MADE, AMBIGUITIES, OR WARNINGS ({datetime.datetime.now()}) ---\n")
                f.write(f"Input DOCX: {os.path.basename(input_docx_path)}\n")
                f.write(f"Output DOCX: {os.path.basename(output_docx_path)}\n")
                f.write(f"Total Edit Instructions: {len(edits_to_make)}, Edits Successfully Processed This Run: {edits_processed_in_current_run}\n")
                f.write(f"Log Items (Failures/Warnings/Errors): {len(ambiguous_or_failed_changes_log)}\n\n")
                for log_entry in ambiguous_or_failed_changes_log:
                    f.write("-----------------------------------------\n")
                    # paragraph_index in log_entry is 0-based from script
                    para_display_index = log_entry.get('paragraph_index', -1)
                    if isinstance(para_display_index, int) and para_display_index >=0 :
                        para_display_index +=1 # Convert to 1-based for display
                    else:
                        para_display_index = 'N/A'

                    f.write(f"Paragraph Index (1-based): {para_display_index}\n")
                    f.write(f"Visible Text Snippet: {log_entry.get('visible_text_snippet', 'N/A')}\n")
                    f.write(f"Context Searched: '{log_entry.get('contextual_old_text', '')}'\n")
                    f.write(f"Specific Old Text: '{log_entry.get('specific_old_text', '')}'\n")
                    f.write(f"Specific New Text: '{log_entry.get('specific_new_text', 'N/A')}'\n")
                    f.write(f"LLM Reason for Change: {log_entry.get('llm_reason', 'N/A')}'\n")
                    f.write(f"Issue: {log_entry.get('issue', 'Unknown')}\n")
                    log_type = log_entry.get('type', 'Error') 
                    f.write(f"Type: {log_type}\n")
                    if 'edit_item_index' in log_entry: f.write(f"Edit Item Index (0-based): {log_entry['edit_item_index']}\n")
                    if 'edit_item_snippet' in log_entry: f.write(f"Edit Item Snippet: {log_entry['edit_item_snippet']}\n")

                f.write("-----------------------------------------\n")
            print(f"Log file with {len(ambiguous_or_failed_changes_log)} items saved to: '{error_log_file_path}'")
            log_debug(f"Log file saved to: '{error_log_file_path}'")
        except Exception as e_log:
            log_debug(f"ERROR: Could not write to log file '{error_log_file_path}': {e_log}")
            error_log_file_path = None 
    else:
        # Refined summary messages
        if not edits_to_make:
            print(f"No edits provided to process.")
        elif edits_processed_in_current_run == len(edits_to_make):
             print(f"All {edits_processed_in_current_run} targeted changes applied successfully. No issues logged.")
        elif edits_processed_in_current_run < len(edits_to_make) and edits_processed_in_current_run > 0:
            print(f"{edits_processed_in_current_run} changes applied. Some edits may not have found their context or were ambiguous. No critical errors logged.")
        else: # edits_processed_in_current_run == 0 and edits_to_make
            print(f"No changes applied. Edits may not have found their context or were ambiguous. No critical errors logged.")
        log_debug("No items for error log file.")
    return True, error_log_file_path, ambiguous_or_failed_changes_log

if __name__ == '__main__':
    import argparse
    DEFAULT_EDITS_FILE_PATH = "edits_to_apply.json"
    DEFAULT_INPUT_DOCX_PATH = "sample_input.docx" 
    DEFAULT_OUTPUT_DOCX_PATH = "sample_output_corrected.docx" # Changed output name

    parser = argparse.ArgumentParser(description="Apply tracked changes to a Word document.")
    # (argparse setup unchanged)
    parser.add_argument("--input", default=DEFAULT_INPUT_DOCX_PATH, help=f"Input DOCX (default: {DEFAULT_INPUT_DOCX_PATH})")
    parser.add_argument("--output", default=DEFAULT_OUTPUT_DOCX_PATH, help=f"Output DOCX (default: {DEFAULT_OUTPUT_DOCX_PATH})")
    parser.add_argument("--editsfile", default=DEFAULT_EDITS_FILE_PATH, help=f"JSON edits file (default: {DEFAULT_EDITS_FILE_PATH})")
    parser.add_argument("--editsjson", type=str, help="JSON string of edits.")
    parser.add_argument("--author", default=DEFAULT_AUTHOR_NAME, help=f"Author for changes (default: {DEFAULT_AUTHOR_NAME}).")
    parser.add_argument("--debug", action="store_true", help="Enable debug mode.")
    parser.add_argument("--caseinsensitive", action="store_false", dest="case_sensitive", help="Case-insensitive search.")
    parser.set_defaults(case_sensitive=True)
    parser.add_argument("--nocomments", action="store_false", dest="add_comments", help="Disable comments.")
    parser.set_defaults(add_comments=True)
    
    args = parser.parse_args()
    if args.debug: DEBUG_MODE = True

    edits_data = []
    # (edits data loading unchanged)
    if args.editsjson:
        try:
            edits_data = json.loads(args.editsjson)
            log_debug(f"Loaded edits from --editsjson argument.")
        except json.JSONDecodeError as e:
            print(f"FATAL: Error decoding JSON from --editsjson: {e}. Exiting.")
            exit(1)
    elif args.editsfile:
        try:
            with open(args.editsfile, 'r', encoding='utf-8') as f:
                edits_data = json.load(f)
            log_debug(f"Successfully loaded {len(edits_data)} edits from '{args.editsfile}'.")
        except FileNotFoundError:
            print(f"FATAL: Edits file '{args.editsfile}' not found. Exiting.")
            exit(1)
        except json.JSONDecodeError as e:
            print(f"FATAL: Error decoding JSON from '{args.editsfile}': {e}. Exiting.")
            exit(1)
        except Exception as e:
            print(f"FATAL: An unexpected error loading '{args.editsfile}': {e}. Exiting.")
            exit(1)
    else:
        if not DEBUG_MODE: DEBUG_MODE = True 
        log_debug("No edits provided via --editsfile or --editsjson. Will try to run with dummy edits if input/output are defaults.")
        if args.input == DEFAULT_INPUT_DOCX_PATH and args.output == DEFAULT_OUTPUT_DOCX_PATH :
            print("INFO: No edits specified. Using dummy edits for testing.")
            edits_data = [
                {"contextual_old_text": "cost would be $101 ,", "specific_old_text": "$101", "specific_new_text": "$777", "reason_for_change": "Test $101, expect highlight (due to comma)."},
                {"contextual_old_text": "MrArbor, but that name", "specific_old_text": "MrArbor", "specific_new_text": "DrArbor", "reason_for_change": "Test MrArbor, expect highlight (due to comma)."},
                {"contextual_old_text": "name can change", "specific_old_text": "name", "specific_new_text": "label", "reason_for_change": "Test 'name', no highlight expected."},
                {"contextual_old_text": "with 9 longs.", "specific_old_text": "longs.", "specific_new_text": "iterations!", "reason_for_change": "Test end 'longs.', expect highlight (due to period)."}
            ]
            if not os.path.exists(DEFAULT_INPUT_DOCX_PATH):
                try:
                    doc_dummy = Document()
                    doc_dummy.add_paragraph("A simple file. It should change this sentence from saying the cost would be $101 , to a new number.")
                    doc_dummy.add_paragraph("Another line. Lets count 1, 2, 3.")
                    doc_dummy.add_paragraph("Here is a long long long long long long long long long long repetitive line with 9 longs.")
                    doc_dummy.add_paragraph("The last line was last edited by MrArbor, but that name can change.")
                    doc_dummy.save(DEFAULT_INPUT_DOCX_PATH)
                    print(f"Created dummy input: {DEFAULT_INPUT_DOCX_PATH}")
                except Exception as e_doc:
                    print(f"Could not create dummy input: {e_doc}")
                    exit(1)
        else:
             print("FATAL: No edits provided. Use --editsfile or --editsjson. Exiting.")
             exit(1)

    if not os.path.exists(args.input):
        print(f"FATAL: Input file '{args.input}' not found. Exiting.")
        exit(1)

    process_document_with_edits(
        args.input, args.output, edits_data,
        args.author, args.debug, args.case_sensitive, args.add_comments
    )