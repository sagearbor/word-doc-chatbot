import os
os.environ['LITELLM_LOG'] = 'DEBUG' 
import shutil
import tempfile
import uuid
from typing import List, Dict, Optional
import traceback 
import json

from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.responses import FileResponse, JSONResponse
# docx.Document is used by extract_text_for_llm if that's kept, and by word_processor
from docx import Document 

# Corrected imports based on new llm_handler and word_processor structure
from .llm_handler import (
    get_llm_suggestions,
    get_llm_analysis_from_summary,   # For Approach 1
    get_llm_analysis_from_raw_xml,   # For Approach 2
    get_llm_suggestions_with_fallback,  # Phase 2.2 Advanced Merging
    get_merge_analysis,                 # Phase 2.2 Analysis
    get_advanced_legal_instructions     # Phase 2.2 Instructions
)
from .word_processor import (
    process_document_with_edits,
    DEFAULT_AUTHOR_NAME,
    _build_visible_text_map,         # Used by extract_text_for_llm
    extract_tracked_changes_as_text, # For Approach 1
    get_document_xml_raw_text        # For Approach 2
)
from .legal_document_processor import (
    parse_legal_document,
    extract_fallback_requirements,
    generate_instructions_from_fallback,
    LegalDocumentStructure,
    LegalRequirement
)
from .requirements_processor import (
    process_fallback_document_requirements,
    generate_enhanced_instructions,
    RequirementsProcessor,
    ProcessedRequirement
)

# Import Phase 4.1 workflow orchestrator
try:
    from .legal_workflow_orchestrator import (
        process_legal_document_workflow,
        LegalDocumentWorkflowResult,
        ProcessingStatus
    )
    WORKFLOW_ORCHESTRATOR_AVAILABLE = True
    print("Phase 4.1 Legal Workflow Orchestrator imported successfully")
except ImportError as e:
    print(f"Warning: Legal Workflow Orchestrator not available: {e}")
    WORKFLOW_ORCHESTRATOR_AVAILABLE = False

app = FastAPI(title="Word Document Processing API")

# Ensure TEMP_DIR_ROOT is defined at the module level
TEMP_DIR_ROOT = tempfile.mkdtemp(prefix="wordapp_root_")
print(f"Temporary root directory created at: {TEMP_DIR_ROOT}")


# This function is used by /process-document/ for LLM suggestions, keep it.
def extract_text_for_llm(path: str) -> str:
    try:
        doc = Document(path) # Requires python-docx Document
        all_paragraphs_text = []
        for p_obj in doc.paragraphs:
            map_text, _ = _build_visible_text_map(p_obj)
            all_paragraphs_text.append(map_text)
        return "\n".join(all_paragraphs_text)
    except Exception as e:
        print(f"Error in extract_text_for_llm: {e}")
        traceback.print_exc()
        try: # Fallback
            doc = Document(path)
            return "\n".join(p.text for p in doc.paragraphs if p.text)
        except Exception as e_fallback:
            print(f"Fallback text extraction also failed: {e_fallback}")
            return ""


@app.post("/analyze-document/")
async def analyze_document_endpoint(
    file: UploadFile = File(...),
    analysis_mode: str = Form("summary") # Default to "summary", options: "summary", "raw_xml"
):
    if not file.filename or not file.filename.lower().endswith(".docx"): # Check lower() for .DOCX
        raise HTTPException(status_code=400, detail="Invalid file type or missing filename. Only .docx files are accepted.")

    request_id = str(uuid.uuid4())
    original_filename = os.path.basename(file.filename) 
    input_docx_path = os.path.join(TEMP_DIR_ROOT, f"{request_id}_analysis_{original_filename}")

    print(f"[PID:{os.getpid()}] /analyze-document/ received for '{original_filename}', mode: '{analysis_mode}'")

    try:
        os.makedirs(TEMP_DIR_ROOT, exist_ok=True) # Ensure temp dir exists
        with open(input_docx_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        print(f"[PID:{os.getpid()}] File '{original_filename}' saved to '{input_docx_path}'")

        analysis_result_text: Optional[str] = None

        if analysis_mode == "summary":
            print(f"[PID:{os.getpid()}] Using 'summary' analysis mode for {original_filename}.")
            summary_of_changes = extract_tracked_changes_as_text(input_docx_path)
            
            print(f"[PID:{os.getpid()}] Summary of changes extracted (first 500 chars):\n{summary_of_changes[:500]}...")
            
            # The get_llm_analysis_from_summary function now handles "No changes" or "Error_Internal" itself
            analysis_result_text = get_llm_analysis_from_summary(summary_of_changes, original_filename)

        elif analysis_mode == "raw_xml":
            print(f"[PID:{os.getpid()}] Using 'raw_xml' analysis mode for {original_filename}.")
            raw_xml_content = get_document_xml_raw_text(input_docx_path)

            print(f"[PID:{os.getpid()}] Raw XML content extracted (first 500 chars):\n{raw_xml_content[:500]}...")

            if raw_xml_content.startswith("Error_Internal:"):
                 analysis_result_text = raw_xml_content # Pass through extraction error
            else:
                # Basic check for very large XML that might be problematic
                # This is a character count, not token count. Actual token count depends on model.
                MAX_RAW_XML_CHARS_FOR_LLM = 150 * 1024 # Approx 150KB as a very rough limit for the input string
                if len(raw_xml_content) > MAX_RAW_XML_CHARS_FOR_LLM:
                    warning_msg = (
                        f"Warning: The document's internal XML content is very large "
                        f"({len(raw_xml_content) // 1024}KB). Analysis might be slow, incomplete, or fail. "
                        "Consider using the 'Summarize Extracted Changes (Concise)' mode for very large/complex documents."
                    )
                    print(f"[PID:{os.getpid()}] {warning_msg}")
                    # We can still proceed but the LLM might struggle.
                    # Or, return this warning to the user:
                    # analysis_result_text = warning_msg 
                
                if analysis_result_text is None: # If not set by size warning
                    analysis_result_text = get_llm_analysis_from_raw_xml(raw_xml_content, original_filename)
        else:
            # This error indicates a problem with how analysis_mode is sent or handled, not user input usually.
            err_msg = f"Error_Server: Invalid analysis_mode '{analysis_mode}' specified. Must be 'summary' or 'raw_xml'."
            print(f"[PID:{os.getpid()}] {err_msg}")
            # Return as 200 so Streamlit can display the error message from the "analysis" field
            return JSONResponse(content={"analysis": err_msg}, status_code=200)

        # Handle cases where LLM calls might return None or a specific error string
        if analysis_result_text is None:
            analysis_result_text = "Error_Server: Analysis result was unexpectedly empty after LLM call."
            print(f"[PID:{os.getpid()}] [WARN] {analysis_result_text} for {original_filename}, mode {analysis_mode}")
        
        print(f"[PID:{os.getpid()}] Analysis result for LLM (first 300 chars): {analysis_result_text[:300]}...")
        return JSONResponse(content={"analysis": analysis_result_text})

    except HTTPException:
        raise # Re-raise HTTPExceptions to let FastAPI handle them as actual HTTP errors
    except Exception as e: 
        print(f"[PID:{os.getpid()}] Critical error in /analyze-document/ for '{original_filename}' (mode: {analysis_mode}): {e}")
        traceback.print_exc() 
        # Return a 200 with a server error message for Streamlit to display from the "analysis" field
        return JSONResponse(content={"analysis": f"Error_Server: An unexpected critical error occurred on the server: {str(e)}"}, status_code=200)
    finally:
        if os.path.exists(input_docx_path):
            try: 
                os.remove(input_docx_path)
                print(f"[PID:{os.getpid()}] Cleaned up temp file: {input_docx_path}")
            except Exception as e_clean: 
                print(f"[PID:{os.getpid()}] Error cleaning up analysis file {input_docx_path}: {e_clean}")


@app.post("/process-document/")
async def process_document(
    file: UploadFile = File(...),
    instructions: str = Form(...),
    author_name: Optional[str] = Form(None),
    case_sensitive: bool = Form(True),
    add_comments: bool = Form(True),
    debug_mode: bool = Form(False),          
    extended_debug_mode: bool = Form(False)  
):
    # ... (keep existing /process-document/ endpoint logic) ...
    if not file.filename or not file.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Only .docx files are supported.")
    request_id = uuid.uuid4().hex
    base_input_filename = os.path.basename(file.filename) 
    input_path = os.path.join(TEMP_DIR_ROOT, f"{request_id}_input_{base_input_filename}")
    output_path = os.path.join(TEMP_DIR_ROOT, f"{request_id}_output_{base_input_filename}")
    try:
        os.makedirs(TEMP_DIR_ROOT, exist_ok=True)
        with open(input_path, "wb") as buff: shutil.copyfileobj(file.file, buff)
        doc_text_for_llm = extract_text_for_llm(input_path) 
        print("\n[MAIN_PY_DEBUG] Text sent for /process-document/ LLM (get_llm_suggestions):")
        print(doc_text_for_llm[:1000] + "..." if len(doc_text_for_llm) > 1000 else doc_text_for_llm)
        print("[MAIN_PY_DEBUG] End of text for /process-document/ LLM.\n")
        edits: Optional[List[Dict]] = get_llm_suggestions(doc_text_for_llm, instructions, base_input_filename)
        print("\n[MAIN_PY_DEBUG] Raw edits from LLM (get_llm_suggestions output):") 
        if edits is not None: print(json.dumps(edits, indent=2))                                     
        else: print("LLM returned None for edits.")                                  
        print("[MAIN_PY_DEBUG] End of raw edits from LLM.\n")
        if edits is None: 
            raise HTTPException(status_code=500, detail="LLM failed to generate or return valid suggestions format (returned None).")
        processed_edits_count = 0
        log_file_path_returned: Optional[str] = None
        log_details_returned: List[Dict] = [] 
        if edits: 
            wp_success, log_file_path_returned, log_details_returned, processed_edits_count = process_document_with_edits(
                input_docx_path=input_path, output_docx_path=output_path, edits_to_make=edits,
                author_name=author_name if author_name else DEFAULT_AUTHOR_NAME,
                debug_mode_flag=debug_mode, extended_debug_mode_flag=extended_debug_mode, 
                case_sensitive_flag=case_sensitive, add_comments_param=add_comments,         
            )
            if not wp_success: 
                error_issue = "Word processing script encountered a critical failure."
                if log_details_returned and isinstance(log_details_returned, list) and log_details_returned[0].get("issue"):
                    error_issue = str(log_details_returned[0]["issue"])
                raise HTTPException(status_code=500, detail=error_issue)
        else:  # if not edits (LLM suggested no changes)
            shutil.copy(input_path, output_path) 
            log_details_returned = [{"issue": "LLM suggested no changes for the given instructions.", "type": "Info"}]
            processed_edits_count = 0 

        log_content_for_response: str # Make it a string type from the start

        if log_file_path_returned and os.path.exists(log_file_path_returned):
            with open(log_file_path_returned, "r", encoding="utf-8") as f_log:
                log_content_for_response = f_log.read()
        elif log_details_returned: # This is the List[Dict] from ambiguous_or_failed_changes_log
            log_content_for_response = "\n".join( # This will be an empty string if log_details_returned is empty
                f"Type: {d.get('type', 'Log')}, Issue: {d.get('issue', 'N/A')}, "
                f"Para: {d.get('paragraph_index', -1)+1 if isinstance(d.get('paragraph_index'), int) else 'N/A'}, "
                f"Old: '{d.get('specific_old_text', '')}'"
                for d in log_details_returned
            )
        else: # log_details_returned is None (shouldn't happen if initialized to []) OR was an empty list and caught by prior elif being false.
              # This 'else' means no log file AND log_details_returned was effectively empty or None.
            if edits: # If LLM suggested edits but they resulted in an empty log from word_processor
                log_content_for_response = "No specific processing issues, warnings, or errors were recorded in the detailed log for the suggested edits. They may have been skipped (e.g., 'context not found') without being logged as specific issues here."
            else: # LLM suggested no edits in the first place (this path is handled by the outer 'if edits:' else block mostly)
                 log_content_for_response = "The AI suggested no changes, so no processing log was generated for new edits."
        
        # If log_content_for_response became "" from an empty log_details_returned list, provide a default message.
        if not log_content_for_response and isinstance(log_details_returned, list) and not log_details_returned:
            log_content_for_response = "No specific processing issues, warnings, or errors were recorded in the detailed log for the suggested edits."
        
        # --- THIS LOGIC MUST BE OUTSIDE THE PREVIOUS IF ---
        total_suggested_edits = len(edits) 
        final_status_message = "Processing complete." # Default status

        if total_suggested_edits == 0:
            final_status_message = "Processing complete. The LLM suggested no changes."
        elif processed_edits_count == 0 and total_suggested_edits > 0:
            final_status_message = f"Processing complete. {total_suggested_edits} edit(s) were suggested by the LLM, but none were applied by the processor. Please check the log for details (e.g., context not found, boundary rule skips, ambiguities)."
        elif processed_edits_count < total_suggested_edits:
            final_status_message = f"Processing complete. {processed_edits_count} out of {total_suggested_edits} suggested changes were applied. Some edits may have been skipped or were ambiguous. Please check the log for details."
        elif processed_edits_count == total_suggested_edits and total_suggested_edits > 0: # Ensure we don't say "all 0"
             final_status_message = f"Processing complete. All {processed_edits_count} suggested changes were successfully applied."

        if not os.path.exists(output_path) and os.path.exists(input_path): # Ensure output file exists
             shutil.copy(input_path, output_path) 

        return JSONResponse(
            content={
                "processed_filename": os.path.basename(output_path),
                "download_url": f"/download/{os.path.basename(output_path)}", 
                "log_content": log_content_for_response, # This should now always be a non-None string
                "status_message": final_status_message,
                "issues_count": len(log_details_returned) if log_details_returned else 0, 
                "edits_applied_count": processed_edits_count,
                "edits_suggested_count": total_suggested_edits
            }
        )

    except HTTPException: raise
    except Exception as e: 
        print(f"Error in /process-document/ endpoint: {e}"); traceback.print_exc() 
        raise HTTPException(status_code=500, detail=f"An unexpected error occurred on the server: {str(e)}")
    finally:
        if os.path.exists(input_path):
            try: os.remove(input_path)
            except Exception as e_clean_in: print(f"Error cleaning up input file {input_path}: {e_clean_in}")


@app.get("/download/{filename}")
async def download_file_presumably_from_temp_dir(filename: str):
    # ... (keep existing /download/ endpoint logic) ...
    if ".." in filename or filename.startswith("/") or filename.startswith("\\"):
        raise HTTPException(status_code=400, detail="Invalid filename.")
    file_path = os.path.join(TEMP_DIR_ROOT, filename)
    if not os.path.isfile(file_path): 
        print(f"Download request for non-existent or non-file path: {file_path}")
        raise HTTPException(status_code=404, detail="File not found or is not accessible.")
    return FileResponse(file_path, filename=filename, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@app.on_event("shutdown")
def cleanup_temp_dir(): 
    # ... (keep existing shutdown event logic) ...
    print(f"Attempting to clean up temporary directory: {TEMP_DIR_ROOT}")
    try:
        if os.path.exists(TEMP_DIR_ROOT):
            shutil.rmtree(TEMP_DIR_ROOT)
            print(f"Successfully removed temporary directory: {TEMP_DIR_ROOT}")
    except Exception as e:
        print(f"Error during temporary directory cleanup: {e}"); traceback.print_exc()

# New endpoints for fallback document processing

@app.post("/upload-fallback-document/")
async def upload_fallback_document(
    file: UploadFile = File(...),
    test_case_id: str = Form("default"),
    description: str = Form("")
):
    """Upload and analyze fallback document for requirement extraction"""
    if not file.filename or not file.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Only .docx files are supported.")
    
    request_id = str(uuid.uuid4())
    original_filename = os.path.basename(file.filename)
    fallback_path = os.path.join(TEMP_DIR_ROOT, f"{request_id}_fallback_{original_filename}")
    
    print(f"[PID:{os.getpid()}] /upload-fallback-document/ received for '{original_filename}'")
    
    try:
        os.makedirs(TEMP_DIR_ROOT, exist_ok=True)
        with open(fallback_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        print(f"[PID:{os.getpid()}] Fallback document saved to '{fallback_path}'")
        
        # Parse legal document structure
        document_structure = parse_legal_document(fallback_path)
        
        # Extract requirements
        requirements = extract_fallback_requirements(fallback_path)
        
        # Generate summary
        summary = {
            "filename": original_filename,
            "test_case_id": test_case_id,
            "description": description,
            "document_title": document_structure.title,
            "sections_count": len(document_structure.sections),
            "requirements_count": len(requirements),
            "whereas_clauses_count": len(document_structure.whereas_clauses),
            "authors": document_structure.authors,
            "high_priority_requirements": len([r for r in requirements if r.priority <= 2]),
            "requirement_types": {
                req_type: len([r for r in requirements if r.requirement_type == req_type])
                for req_type in ["must", "shall", "required", "prohibited"]
            }
        }
        
        print(f"[PID:{os.getpid()}] Fallback document analysis complete: {len(requirements)} requirements found")
        
        return JSONResponse(content={
            "status": "success",
            "summary": summary,
            "fallback_id": request_id,
            "requirements_preview": [
                {"text": req.text[:100] + "..." if len(req.text) > 100 else req.text,
                 "type": req.requirement_type,
                 "priority": req.priority,
                 "section": req.section}
                for req in requirements[:10]  # First 10 requirements
            ]
        })
        
    except Exception as e:
        print(f"[PID:{os.getpid()}] Error processing fallback document: {e}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error processing fallback document: {str(e)}")
    
    finally:
        if os.path.exists(fallback_path):
            try:
                os.remove(fallback_path)
                print(f"[PID:{os.getpid()}] Cleaned up fallback temp file: {fallback_path}")
            except Exception as e:
                print(f"[PID:{os.getpid()}] Error cleaning up fallback file: {e}")

@app.post("/analyze-fallback-requirements/")
async def analyze_fallback_requirements(
    file: UploadFile = File(...),
    context: str = Form("")
):
    """Analyze fallback document and extract requirements for instruction generation"""
    if not file.filename or not file.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Only .docx files are supported.")
    
    request_id = str(uuid.uuid4())
    original_filename = os.path.basename(file.filename)
    fallback_path = os.path.join(TEMP_DIR_ROOT, f"{request_id}_analyze_{original_filename}")
    
    print(f"[PID:{os.getpid()}] /analyze-fallback-requirements/ received for '{original_filename}'")
    
    try:
        os.makedirs(TEMP_DIR_ROOT, exist_ok=True)
        with open(fallback_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # Generate enhanced instructions using requirements processor
        try:
            instructions = generate_enhanced_instructions(fallback_path, context)
            print(f"[PID:{os.getpid()}] Generated enhanced instructions using requirements processor")
        except Exception as e:
            print(f"[PID:{os.getpid()}] Enhanced processing failed, falling back to basic: {e}")
            instructions = generate_instructions_from_fallback(fallback_path, context)
        
        # Extract basic requirements (using the fixed extractor)
        try:
            requirements = extract_fallback_requirements(fallback_path)
            print(f"[PID:{os.getpid()}] Extracted {len(requirements)} basic requirements from fallback document")
            
            # Create categorized view for API response
            categorized_requirements = {}
            for req in requirements:
                key = f"{req.requirement_type}_priority_{req.priority}"
                if key not in categorized_requirements:
                    categorized_requirements[key] = []
                categorized_requirements[key].append({
                    "text": req.text,
                    "section": req.section,
                    "context": req.context[:200] + "..." if len(req.context) > 200 else req.context,
                    "priority": req.priority,
                    "requirement_type": req.requirement_type
                })
            
            requirements_count = len(requirements)
            
            # Try enhanced processing as secondary step (optional)
            try:
                processed_requirements = process_fallback_document_requirements(fallback_path)
                if processed_requirements and len(processed_requirements) > len(requirements):
                    print(f"[PID:{os.getpid()}] Enhanced processing found {len(processed_requirements)} requirements, using enhanced results")
                    # Only use enhanced if it found more requirements
                    enhanced_categorized = {}
                    for req in processed_requirements:
                        key = f"{req.category.value}_priority_{req.priority_level.value}"
                        if key not in enhanced_categorized:
                            enhanced_categorized[key] = []
                        enhanced_categorized[key].append({
                            "text": req.reformatted_text,
                            "section": req.original.section,
                            "context": req.original.context[:200] + "..." if len(req.original.context) > 200 else req.original.context,
                            "conflicts": req.conflicts,
                            "confidence_score": req.confidence_score,
                            "processing_notes": req.processing_notes[:2] if req.processing_notes else []
                        })
                    categorized_requirements = enhanced_categorized
                    requirements_count = len(processed_requirements)
            except Exception as e_enhanced:
                print(f"[PID:{os.getpid()}] Enhanced processing failed, using basic results: {e_enhanced}")
                
        except Exception as e:
            print(f"[PID:{os.getpid()}] Basic requirement extraction also failed: {e}")
            requirements_count = 0
            categorized_requirements = {}
        
        return JSONResponse(content={
            "status": "success",
            "instructions": instructions,
            "requirements_count": requirements_count,
            "categorized_requirements": categorized_requirements,
            "context_used": context,
            "filename": original_filename
        })
        
    except Exception as e:
        print(f"[PID:{os.getpid()}] Error analyzing fallback requirements: {e}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error analyzing fallback requirements: {str(e)}")
    
    finally:
        if os.path.exists(fallback_path):
            try:
                os.remove(fallback_path)
                print(f"[PID:{os.getpid()}] Cleaned up analyze temp file: {fallback_path}")
            except Exception as e:
                print(f"[PID:{os.getpid()}] Error cleaning up analyze file: {e}")

@app.post("/process-document-with-fallback/")
async def process_document_with_fallback(
    input_file: UploadFile = File(...),
    fallback_file: UploadFile = File(...),
    user_instructions: str = Form(""),
    author_name: Optional[str] = Form(None),
    case_sensitive: bool = Form(True),
    add_comments: bool = Form(True),
    debug_mode: bool = Form(False),
    extended_debug_mode: bool = Form(False),
    merge_strategy: str = Form("append")  # "append", "prepend", "priority"
):
    """Process document using both fallback requirements and user instructions"""
    # Validate file types
    for file, name in [(input_file, "input"), (fallback_file, "fallback")]:
        if not file.filename or not file.filename.lower().endswith(".docx"):
            raise HTTPException(status_code=400, detail=f"Only .docx files are supported for {name} file.")
    
    request_id = str(uuid.uuid4())
    input_filename = os.path.basename(input_file.filename)
    fallback_filename = os.path.basename(fallback_file.filename)
    
    input_path = os.path.join(TEMP_DIR_ROOT, f"{request_id}_input_{input_filename}")
    fallback_path = os.path.join(TEMP_DIR_ROOT, f"{request_id}_fallback_{fallback_filename}")
    output_path = os.path.join(TEMP_DIR_ROOT, f"{request_id}_output_{input_filename}")
    
    print(f"[PID:{os.getpid()}] /process-document-with-fallback/ - Input: '{input_filename}', Fallback: '{fallback_filename}'")
    
    try:
        os.makedirs(TEMP_DIR_ROOT, exist_ok=True)
        
        # Save both files
        with open(input_path, "wb") as buffer:
            shutil.copyfileobj(input_file.file, buffer)
        with open(fallback_path, "wb") as buffer:
            shutil.copyfileobj(fallback_file.file, buffer)
        
        # Extract text for LLM processing
        doc_text_for_llm = extract_text_for_llm(input_path)
        
        # TEMPORARILY DISABLE Phase 2.2 - it's not working properly
        # Force use of working Phase 2.1 Enhanced Instructions
        force_use_phase21 = True
        
        # Use Phase 2.2 Advanced Instruction Merging (or force Phase 2.1)
        try:
            if force_use_phase21:
                print(f"[PID:{os.getpid()}] FORCING Phase 2.1 Enhanced Instructions (Phase 2.2 disabled for debugging)...")
                raise Exception("Phase 2.2 disabled - using Phase 2.1")
            
            print(f"[PID:{os.getpid()}] Attempting Phase 2.2 Advanced Instruction Merging...")
            
            # Convert simple merge strategy to Phase 2.2 strategy
            phase22_strategy_map = {
                "append": "intelligent_merge",
                "prepend": "user_priority", 
                "priority": "fallback_priority"
            }
            phase22_strategy = phase22_strategy_map.get(merge_strategy, "intelligent_merge")
            
            # Get LLM suggestions using Phase 2.2 advanced merging
            edits = get_llm_suggestions_with_fallback(
                doc_text_for_llm, 
                user_instructions, 
                input_filename, 
                fallback_path
            )
            
            # For debug mode, capture additional fallback processing info
            if debug_mode or extended_debug_mode:
                try:
                    from .legal_document_processor import extract_fallback_requirements
                    fallback_requirements = extract_fallback_requirements(fallback_path)
                    debug_fallback_info = {
                        "fallback_requirements_count": len(fallback_requirements) if fallback_requirements else 0,
                        "requirement_types": list(set([req.requirement_type for req in fallback_requirements])) if fallback_requirements else [],
                        "sample_requirements": [req.text[:100] + "..." if len(req.text) > 100 else req.text for req in fallback_requirements[:3]] if fallback_requirements else []
                    }
                except Exception as e:
                    debug_fallback_info = {"error": f"Could not analyze fallback requirements: {str(e)}"}
            
            print(f"[PID:{os.getpid()}] Phase 2.2 Advanced Merging completed successfully")
            
        except Exception as e:
            print(f"[PID:{os.getpid()}] Phase 2.2 Advanced Merging failed: {e}")
            print(f"[PID:{os.getpid()}] Falling back to Phase 2.1 enhanced instructions...")
            
            # Fallback to Phase 2.1 method - use our FIXED function directly
            print(f"[PID:{os.getpid()}] Using fixed generate_instructions_from_fallback function...")
            fallback_instructions = generate_instructions_from_fallback(
                fallback_path, 
                context=f"Processing {input_filename} with fallback requirements"
            )
            print(f"[PID:{os.getpid()}] Generated {len(fallback_instructions)} characters of fallback instructions")
            
            # Merge instructions based on strategy
            if merge_strategy == "append":
                combined_instructions = f"{fallback_instructions}\n\nAdditional User Instructions:\n{user_instructions}" if user_instructions else fallback_instructions
            elif merge_strategy == "prepend":
                combined_instructions = f"{user_instructions}\n\nFallback Requirements:\n{fallback_instructions}" if user_instructions else fallback_instructions
            else:  # priority - fallback takes precedence
                combined_instructions = fallback_instructions
                if user_instructions:
                    combined_instructions += f"\n\nNote: Apply user instructions only if they do not conflict with above requirements:\n{user_instructions}"
            
            print(f"[PID:{os.getpid()}] Combined instructions length: {len(combined_instructions)} characters")
            
            # Get LLM suggestions using combined instructions (Phase 2.1 fallback)
            edits = get_llm_suggestions(doc_text_for_llm, combined_instructions, input_filename)
        
        if edits is None:
            raise HTTPException(status_code=500, detail="LLM failed to generate suggestions from fallback document.")
        
        # Process document with edits
        processed_edits_count = 0
        log_file_path_returned = None
        log_details_returned = []
        
        if edits:
            wp_success, log_file_path_returned, log_details_returned, processed_edits_count = process_document_with_edits(
                input_docx_path=input_path,
                output_docx_path=output_path,
                edits_to_make=edits,
                author_name=author_name if author_name else DEFAULT_AUTHOR_NAME,
                debug_mode_flag=debug_mode,
                extended_debug_mode_flag=extended_debug_mode,
                case_sensitive_flag=case_sensitive,
                add_comments_param=add_comments
            )
            
            if not wp_success:
                error_issue = "Word processing failed with fallback document."
                if log_details_returned and isinstance(log_details_returned, list) and log_details_returned[0].get("issue"):
                    error_issue = str(log_details_returned[0]["issue"])
                raise HTTPException(status_code=500, detail=error_issue)
        else:
            shutil.copy(input_path, output_path)
            log_details_returned = [{"issue": "No changes suggested based on fallback document requirements.", "type": "Info"}]
            processed_edits_count = 0
        
        # Prepare log content
        log_content_for_response = ""
        if log_file_path_returned and os.path.exists(log_file_path_returned):
            with open(log_file_path_returned, "r", encoding="utf-8") as f:
                log_content_for_response = f.read()
        elif log_details_returned:
            log_content_for_response = "\n".join(
                f"Type: {d.get('type', 'Log')}, Issue: {d.get('issue', 'N/A')}, "
                f"Para: {d.get('paragraph_index', -1)+1 if isinstance(d.get('paragraph_index'), int) else 'N/A'}, "
                f"Old: '{d.get('specific_old_text', '')}'"
                for d in log_details_returned
            )
        
        if not log_content_for_response:
            log_content_for_response = "No specific processing issues recorded for fallback document processing."
        
        # Generate status message
        total_suggested_edits = len(edits)
        if total_suggested_edits == 0:
            status_message = "Processing complete. No changes suggested based on fallback document."
        elif processed_edits_count == total_suggested_edits:
            status_message = f"Processing complete. All {processed_edits_count} fallback-based changes were applied."
        else:
            status_message = f"Processing complete. {processed_edits_count} out of {total_suggested_edits} fallback-based changes were applied."
        
        # Add debug information for troubleshooting
        debug_info = {}
        if debug_mode or extended_debug_mode:
            # Basic debug information (always included when debug is enabled)
            debug_info = {
                "debug_mode_enabled": debug_mode,
                "extended_debug_enabled": extended_debug_mode,
                "user_instructions_length": len(user_instructions) if user_instructions else 0,
                "raw_edits_from_llm": len(edits) if edits else 0,
                "processing_successful": processed_edits_count > 0,
                "edits_preview": [
                    {
                        "old": edit.get("specific_old_text", "")[:50],
                        "new": edit.get("specific_new_text", "")[:50],
                        "reason": edit.get("reason", "")[:100]
                    } for edit in (edits[:3] if edits else [])
                ],
                "user_friendly_summary": {
                    "requirements_found": "✅ Found requirements from fallback document",
                    "llm_processing": f"🤖 LLM suggested {len(edits) if edits else 0} edits",
                    "document_processing": f"📝 Successfully applied {processed_edits_count} out of {total_suggested_edits} edits",
                    "potential_issues": []
                }
            }
            
            # Add potential issue diagnostics
            if len(edits) == 0:
                debug_info["user_friendly_summary"]["potential_issues"].append("❌ LLM generated no edit suggestions - check if fallback requirements are clear")
            elif processed_edits_count == 0:
                debug_info["user_friendly_summary"]["potential_issues"].append("❌ No edits were applied - text matching may have failed")
            elif processed_edits_count < total_suggested_edits:
                debug_info["user_friendly_summary"]["potential_issues"].append("⚠️ Some edits failed to apply - check for exact text matches")
            
            # Extended debug information (only when extended debugging is enabled)
            if extended_debug_mode:
                debug_info["extended_details"] = {
                    "fallback_document_analysis": debug_fallback_info if 'debug_fallback_info' in locals() else "Fallback requirements extracted and processed",
                    "instruction_merging": f"User instructions ({len(user_instructions)} chars) merged with fallback requirements",
                    "llm_prompt_preview": "LLM was asked to modify the document based on merged requirements...",
                    "edit_details": [
                        {
                            "edit_number": i + 1,
                            "old_text": edit.get("specific_old_text", ""),
                            "new_text": edit.get("specific_new_text", ""),
                            "contextual_text": edit.get("contextual_old_text", "")[:200],
                            "reason": edit.get("reason", ""),
                            "applied_successfully": i < processed_edits_count
                        } for i, edit in enumerate(edits[:5] if edits else [])  # Show first 5 edits
                    ],
                    "processing_method": "Phase 2.2 Advanced Merging" if "Phase 2.2" in log_content_for_response else "Phase 2.1 Enhanced Instructions"
                }
        
        response_content = {
            "processed_filename": os.path.basename(output_path),
            "download_url": f"/download/{os.path.basename(output_path)}",
            "log_content": log_content_for_response,
            "status_message": status_message,
            "issues_count": len(log_details_returned) if log_details_returned else 0,
            "edits_applied_count": processed_edits_count,
            "edits_suggested_count": total_suggested_edits,
            "fallback_filename": fallback_filename,
            "merge_strategy": merge_strategy,
            "processing_method": "Phase 2.2 Advanced Merging" if "Phase 2.2" in log_content_for_response else "Phase 2.1 Enhanced Instructions"
        }
        
        if debug_info:
            response_content["debug_info"] = debug_info
            
        return JSONResponse(content=response_content)
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"[PID:{os.getpid()}] Error in /process-document-with-fallback/: {e}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error processing document with fallback: {str(e)}")
    
    finally:
        for path in [input_path, fallback_path]:
            if os.path.exists(path):
                try:
                    os.remove(path)
                    print(f"[PID:{os.getpid()}] Cleaned up temp file: {path}")
                except Exception as e:
                    print(f"[PID:{os.getpid()}] Error cleaning up {path}: {e}")

@app.post("/analyze-merge/")
async def analyze_merge_endpoint(
    fallback_file: UploadFile = File(...),
    user_instructions: str = Form(""),
    merge_strategy: str = Form("intelligent_merge")  # Phase 2.2 strategies
):
    """Phase 2.2: Analyze advanced instruction merging without processing document"""
    
    if not fallback_file.filename or not fallback_file.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Only .docx files are supported for fallback file.")
    
    request_id = str(uuid.uuid4())
    fallback_filename = os.path.basename(fallback_file.filename)
    fallback_path = os.path.join(TEMP_DIR_ROOT, f"{request_id}_analyze_fallback_{fallback_filename}")
    
    print(f"[PID:{os.getpid()}] /analyze-merge/ - Fallback: '{fallback_filename}', Strategy: '{merge_strategy}'")
    
    try:
        os.makedirs(TEMP_DIR_ROOT, exist_ok=True)
        
        # Save fallback file
        with open(fallback_path, "wb") as buffer:
            shutil.copyfileobj(fallback_file.file, buffer)
        
        # Get Phase 2.2 merge analysis
        merge_analysis = get_merge_analysis(fallback_path, user_instructions)
        
        # Add additional metadata
        merge_analysis.update({
            "fallback_filename": fallback_filename,
            "merge_strategy_requested": merge_strategy,
            "user_instructions_length": len(user_instructions),
            "user_instructions_preview": user_instructions[:200] + "..." if len(user_instructions) > 200 else user_instructions
        })
        
        print(f"[PID:{os.getpid()}] Phase 2.2 merge analysis complete for '{fallback_filename}'")
        return JSONResponse(content=merge_analysis)
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"[PID:{os.getpid()}] Error in /analyze-merge/: {e}")
        traceback.print_exc()
        return JSONResponse(
            content={
                "error": f"Merge analysis failed: {str(e)}",
                "available": False,
                "fallback_filename": fallback_filename
            },
            status_code=500
        )
    
    finally:
        if os.path.exists(fallback_path):
            try:
                os.remove(fallback_path)
                print(f"[PID:{os.getpid()}] Cleaned up analyze-merge temp file: {fallback_path}")
            except Exception as e:
                print(f"[PID:{os.getpid()}] Error cleaning up analyze-merge file: {e}")

@app.post("/process-legal-document/")
async def process_legal_document_endpoint(
    input_file: UploadFile = File(...),
    user_instructions: str = Form(""),
    fallback_file: Optional[UploadFile] = File(None),
    author_name: Optional[str] = Form(None),
    enable_audit_logging: bool = Form(True),
    enable_backup: bool = Form(True),
    enable_validation: bool = Form(True)
):
    """
    Phase 4.1 Legal Document Processing Workflow - End-to-end orchestrated processing
    
    This endpoint uses the complete workflow orchestrator that integrates:
    - Phase 1.1: Legal Document Parser
    - Phase 2.1: Requirements Extraction
    - Phase 2.2: Advanced Instruction Merging
    - Document Processing and Validation
    
    If no fallback document is provided, it falls back to the original simple workflow.
    """
    if not WORKFLOW_ORCHESTRATOR_AVAILABLE:
        raise HTTPException(
            status_code=503, 
            detail="Legal Workflow Orchestrator not available. Please use /process-document-with-fallback/ endpoint."
        )
    
    # Validate input file
    if not input_file.filename or not input_file.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Only .docx files are supported for input file.")
    
    # Validate fallback file if provided
    if fallback_file and fallback_file.filename:
        if not fallback_file.filename.lower().endswith(".docx"):
            raise HTTPException(status_code=400, detail="Only .docx files are supported for fallback file.")
    
    request_id = str(uuid.uuid4())
    input_filename = os.path.basename(input_file.filename)
    
    # Save files
    input_path = os.path.join(TEMP_DIR_ROOT, f"{request_id}_legal_input_{input_filename}")
    fallback_path = None
    
    print(f"[PID:{os.getpid()}] /process-legal-document/ - Starting Phase 4.1 workflow for '{input_filename}'")
    
    try:
        os.makedirs(TEMP_DIR_ROOT, exist_ok=True)
        
        # Save input file
        with open(input_path, "wb") as buffer:
            shutil.copyfileobj(input_file.file, buffer)
        
        # Save fallback file if provided
        if fallback_file and fallback_file.filename:
            fallback_filename = os.path.basename(fallback_file.filename)
            fallback_path = os.path.join(TEMP_DIR_ROOT, f"{request_id}_legal_fallback_{fallback_filename}")
            with open(fallback_path, "wb") as buffer:
                shutil.copyfileobj(fallback_file.file, buffer)
            print(f"[PID:{os.getpid()}] Using fallback document: '{fallback_filename}'")
        else:
            print(f"[PID:{os.getpid()}] No fallback document - using original simple workflow")
        
        # Process using workflow orchestrator
        workflow_result = process_legal_document_workflow(
            input_document_path=input_path,
            user_instructions=user_instructions,
            fallback_document_path=fallback_path,
            author_name=author_name,
            enable_audit_logging=enable_audit_logging,
            enable_backup=enable_backup,
            enable_validation=enable_validation
        )
        
        # Check if processing was successful
        if workflow_result.overall_status != ProcessingStatus.COMPLETED:
            error_details = []
            for stage in workflow_result.stage_results:
                if stage.errors:
                    error_details.extend(stage.errors)
            
            raise HTTPException(
                status_code=500,
                detail=f"Processing failed: {workflow_result.status_message}. Errors: {'; '.join(error_details[:3])}"
            )
        
        # Prepare response
        response_data = {
            "workflow_id": workflow_result.workflow_id,
            "processed_filename": workflow_result.processed_filename,
            "download_url": f"/download/{workflow_result.processed_filename}",
            "status_message": workflow_result.status_message,
            "overall_status": workflow_result.overall_status.value,
            "processing_duration_seconds": workflow_result.total_duration_seconds,
            
            # Metrics
            "requirements_extracted": workflow_result.requirements_extracted,
            "requirements_merged": workflow_result.requirements_merged,
            "edits_suggested": workflow_result.edits_suggested,
            "edits_applied": workflow_result.edits_applied,
            "legal_coherence_score": workflow_result.legal_coherence_score,
            "issues_count": workflow_result.issues_count,
            
            # Stage summary
            "stages_completed": len([s for s in workflow_result.stage_results if s.status == ProcessingStatus.COMPLETED]),
            "stages_total": len(workflow_result.stage_results),
            
            # Validation results
            "validation_results": workflow_result.validation_results,
            
            # Processing method
            "processing_method": "Phase 4.1 Complete Workflow" if fallback_path else "Original Simple Workflow",
            
            # Log content (first 5000 chars)
            "log_content": workflow_result.log_content[:5000] if workflow_result.log_content else "No log content available"
        }
        
        # Copy output file to download location if it exists
        if workflow_result.output_document_path and os.path.exists(workflow_result.output_document_path):
            download_path = os.path.join(TEMP_DIR_ROOT, workflow_result.processed_filename)
            if workflow_result.output_document_path != download_path:
                shutil.copy2(workflow_result.output_document_path, download_path)
        
        return JSONResponse(content=response_data)
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"[PID:{os.getpid()}] Error in /process-legal-document/: {e}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error in legal document processing: {str(e)}")
    
    finally:
        # Cleanup only the initial upload files
        for path in [input_path, fallback_path]:
            if path and os.path.exists(path):
                try:
                    os.remove(path)
                    print(f"[PID:{os.getpid()}] Cleaned up temp file: {path}")
                except Exception as e:
                    print(f"[PID:{os.getpid()}] Error cleaning up {path}: {e}")

@app.get("/llm-config/")
async def get_llm_config():
    """Get current LLM configuration status"""
    from .legal_document_processor import USE_LLM_EXTRACTION, USE_LLM_INSTRUCTIONS, get_current_mode
    
    return JSONResponse(content={
        "current_mode": get_current_mode(),
        "extraction_method": "LLM" if USE_LLM_EXTRACTION else "Regex",
        "instruction_method": "LLM" if USE_LLM_INSTRUCTIONS else "Hardcoded",
        "llm_extraction_enabled": USE_LLM_EXTRACTION,
        "llm_instructions_enabled": USE_LLM_INSTRUCTIONS,
        "description": "LLM mode uses AI to understand documents intelligently, while regex/hardcoded uses pattern matching"
    })

@app.post("/llm-config/")
async def set_llm_config(
    extraction_method: str = Form("regex"),  # "llm" or "regex"
    instruction_method: str = Form("hardcoded")  # "llm" or "hardcoded"
):
    """Configure LLM vs regex/hardcoded approaches"""
    from .legal_document_processor import (
        enable_llm_extraction, disable_llm_extraction,
        enable_llm_instructions, disable_llm_instructions,
        get_current_mode
    )
    
    # Configure extraction method
    if extraction_method.lower() == "llm":
        enable_llm_extraction()
    else:
        disable_llm_extraction()
    
    # Configure instruction method
    if instruction_method.lower() == "llm":
        enable_llm_instructions()
    else:
        disable_llm_instructions()
    
    return JSONResponse(content={
        "status": "Configuration updated",
        "new_mode": get_current_mode(),
        "extraction_method": extraction_method,
        "instruction_method": instruction_method
    })

@app.get("/")
async def root():
    endpoints = {
        "message": "Word Document Processing API with Phase 4.1 Legal Workflow Orchestrator",
        "version": "4.1.0",
        "endpoints": {
            "original": [
                "POST /process-document/ - Original document processing",
                "POST /analyze-document/ - Analyze document with tracked changes"
            ],
            "phase_1_2": [
                "POST /upload-fallback-document/ - Upload and analyze fallback document",
                "POST /analyze-fallback-requirements/ - Extract requirements from fallback",
                "POST /process-document-with-fallback/ - Process with fallback requirements"
            ],
            "phase_2_2": [
                "POST /analyze-merge/ - Analyze requirement merging strategies"
            ],
            "phase_4_1": [
                "POST /process-legal-document/ - Complete legal workflow with all phases integrated"
            ],
            "configuration": [
                "GET /llm-config/ - Get current LLM configuration",
                "POST /llm-config/ - Configure LLM vs regex approaches"
            ],
            "utility": [
                "GET /download/{filename} - Download processed document",
                "GET / - This help message"
            ]
        },
        "features": {
            "phase_1_1": "Legal Document Parser",
            "phase_2_1": "Requirements Extraction",
            "phase_2_2": "Advanced Instruction Merging",
            "phase_4_1": "Complete Workflow Orchestration"
        },
        "workflow_orchestrator_available": WORKFLOW_ORCHESTRATOR_AVAILABLE
    }
    return endpoints